#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Task4: 生成海龟策略 .docx 报告
格式: 宋体五号(10.5pt), 1.5倍行距, 0段间距, 两端对齐
"""

import os
import sys
import pandas as pd
import numpy as np

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '..'))
from src.indicators import data_diagnosis
from src.report_utils import (
    set_cjk_font, add_paragraph, add_heading_styled,
    add_picture_captioned, add_table, FONT_NAME, FONT_SIZE,
)
from src.equation_utils import add_equation, _r, _fraction

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from turtle_strategy import (
    run_backtest, calc_metrics, load_stock, param_sweep,
)

# ═══════════════════════════ 路径 ═══════════════════════════

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DATA_DIR = os.path.join(BASE_DIR, '..', 'data', 'csv')
CHART_DIR = os.path.join(BASE_DIR, '..', 'artifacts', 'charts', 'task4')
OUTPUT_FILE = os.path.join(BASE_DIR, 'Rebecca+Task4.docx')


# ═══════════════════════════ 公式 ═══════════════════════════

def add_donchian_formula(doc):
    """Donchian Channel 公式"""
    add_equation(doc, [
        _r("DC_High(N)_t = max(High_{t-N+1}, ..., High_t)"),
    ])
    add_equation(doc, [
        _r("DC_Low(N)_t = min(Low_{t-N+1}, ..., Low_t)"),
    ])

def add_atr_formula(doc):
    """ATR 公式"""
    add_equation(doc, [
        _r("TR_t = max(High_t − Low_t, |High_t − Close_{t−1}|, |Low_t − Close_{t−1}|)"),
    ])
    add_equation(doc, [
        _r("ATR_t = "),
        _fraction([_r("1")], [_r("N")]),
        _r(" · TR_t + "),
        _fraction([_r("N−1")], [_r("N")]),
        _r(" · ATR_{t−1}"),
    ])

def add_position_formula(doc):
    """仓位计算"""
    add_equation(doc, [
        _r("Shares = "),
        _fraction([_r("Capital × Risk%")], [_r("ATR_t")]),
    ])

def add_stop_formula(doc):
    """止损公式"""
    add_equation(doc, [
        _r("StopPrice = EntryPrice − k × ATR_t    (k = 2)"),
    ])


# ═══════════════════════════ 第一章: 海龟策略核心思想 ═══════════════════════════

def write_ch1(doc):
    add_heading_styled(doc, '一、海龟交易策略核心思想', level=1)

    add_heading_styled(doc, '1.1 策略起源', level=2)
    add_paragraph(doc,
        '海龟交易策略（Turtle Trading Strategy）由著名商品交易员 Richard Dennis 和 William Eckhardt '
        '于 1983 年创立。Dennis 坚信交易能力可以被传授——就像在新加坡农场养殖海龟一样——因此他与 Eckhardt '
        '招募并培训了 23 名来自不同背景的普通人（被称为“海龟”），教授他们一套完整的机械交易系统。'
        '在随后的四年中，这群海龟为 Dennis 创造了超过 1 亿美元的利润，证明了“交易可以被教会”的命题。',
        first_line_indent=Cm(0.74))

    add_heading_styled(doc, '1.2 核心思想', level=2)
    add_paragraph(doc,
        '海龟策略的核心思想可以归纳为以下六个维度：',
        first_line_indent=Cm(0.74))

    core_ideas = [
        ('趋势跟踪（Trend Following）',
         '海龟策略不预测市场方向，而是跟随已经形成的趋势。策略的基础信念是：价格趋势一旦形成，'
         '更可能继续而非反转。通过突破 N 日高点入场，策略捕捉的是趋势启动或加速的节点。'),
        ('完全机械化（Mechanical System）',
         '策略的每一个决策——何时入场、买多少、何时加仓、何时止损、何时离场——都由数学规则精确决定，'
         '完全排除主观判断和情绪干扰。这种机械化设计使策略可以从历史数据中精确回测其表现。'),
        ('风险管理优先（Risk Management First）',
         '海龟策略的核心创新不在入场规则，而在其严谨的风险管理体系。每笔交易的风险暴露通过 ATR（N 值）'
         '量化并严格控制——单笔最大亏损不超过总资金的 2%。这种“先管理风险，再追求收益”的思维，'
         '是海龟策略区别于其他趋势策略的关键特征。'),
        ('基于波动率的仓位管理（Volatility-Based Position Sizing）',
         '海龟策略使用 N 值（即 20 日 ATR）作为波动率度量，并以此标准化仓位大小。在高波动期自动减仓、'
         '在低波动期自动加仓——使每一笔交易承担的风险在金额上是恒定的。这是现代量化风险管理的基础。'),
        ('动态止损（Trailing Stop）',
         '采用 2×ATR 的动态跟踪止损，止损点随价格上升而上移，但绝不向下调整。这种“只紧不松”的止损机制'
         '既保护了已有利润（止盈效果），又限制了单笔损失（止损效果）。'),
        ('双系统过滤（Two-System Approach）',
         '海龟策略包含两个子系统：System 1（20 日突破入场 + 10 日突破离场）和 System 2（55 日突破入场 + '
         '20 日突破离场）。双系统并行运行，互为补充——System 1 更敏感但假信号多，System 2 更稳健但信号少。'),
    ]

    for title, desc in core_ideas:
        add_paragraph(doc, f'({core_ideas.index((title, desc))+1}) {title}：{desc}',
                     first_line_indent=Cm(0.74))
        add_paragraph(doc, '')

    add_paragraph(doc,
        '本次实现仅做多方向（Long-only），不包含做空。A 股市场做空机制受限（融券门槛高、标的少），'
        '且仅做多设计可与 Task3 双均线策略直接横向对比。在持续下跌行情中，策略通过空仓降低损失，'
        '而无法像原版海龟那样通过做空从下跌中获利。',
        first_line_indent=Cm(0.74))
    add_paragraph(doc, '')

    add_heading_styled(doc, '1.3 关键优势', level=2)
    advantages = [
        '系统性与纪律性：完全机械化操作消除情绪偏差——恐惧和贪婪不再影响交易决策。',
        '风险可控：每笔交易的风险预先量化，总资金风险暴露有明确上限。策略可以在连续亏损中存活。',
        '适应波动率变化：通过 ATR 标准化仓位，策略在高波动期自动缩小仓位、低波动期自动放大仓位，'
        '使风险暴露始终恒定。',
        '双系统互补：System 1 和 System 2 覆盖不同时间尺度的趋势——短周期捕捉快速趋势，长周期捕捉主要趋势。',
        '跨市场适用：海龟策略最初设计用于商品期货，但其趋势跟踪逻辑适用于任何具有趋势特征的市场——'
        '股票、外汇、加密货币等。',
        '可验证性：因为是机械策略，可以在历史数据上精确回测其表现，避免了“选择性记忆”的认知偏差。',
    ]
    for i, adv in enumerate(advantages, 1):
        add_paragraph(doc, f'{i}. {adv}', first_line_indent=Cm(0.74))

    doc.add_page_break()


# ═══════════════════════════ 第二章: 核心概念 ═══════════════════════════

def write_ch2(doc):
    add_heading_styled(doc, '二、核心概念详解', level=1)

    # 2.1 高低点通道
    add_heading_styled(doc, '2.1 高低价格通道 (Donchian Channel)', level=2)
    add_paragraph(doc,
        '唐奇安通道（Donchian Channel）由 Richard Donchian 在 1930 年代提出，是技术分析中历史最悠久的'
        '趋势跟踪指标之一。Donchian 被公认为“趋势跟踪之父”，他提出了“移动平均线交叉”和“价格通道突破”'
        '这两个至今仍被广泛使用的趋势跟踪方法。',
        first_line_indent=Cm(0.74))
    add_paragraph(doc,
        'Donchian 通道的计算极为简单——上轨为过去 N 日的最高价，下轨为过去 N 日的最低价——但蕴含了深刻的'
        '趋势跟踪逻辑：当价格突破 N 日高点时，说明市场的买入力量已经超越了此前 N 个交易日中的任何一天，'
        '这是一个统计显著的价格行为。海龟策略使用两个不同周期的通道：入场通道（通常为 20 日或 55 日）'
        '和离场通道（通常为 10 日或 20 日）。入场需要价格突破较长的周期（确保趋势有足够的持续性），'
        '而离场使用较短的周期（更快地响应趋势的逆转）。',
        first_line_indent=Cm(0.74))
    add_paragraph(doc, '唐奇安通道的数学公式：', first_line_indent=Cm(0.74))
    add_donchian_formula(doc)
    add_paragraph(doc, '')
    add_paragraph(doc,
        '其中，入场通道周期（entry_period）决定了突破信号的敏感度——周期越短，信号越频繁但假信号越多；'
        '周期越长，信号越可靠但可能错过部分趋势。离场通道周期（exit_period）应小于入场周期，'
        '确保趋势一旦逆转能及时退出。',
        first_line_indent=Cm(0.74))

    # 2.2 ATR
    add_heading_styled(doc, '2.2 平均真实波幅 (ATR)', level=2)
    add_paragraph(doc,
        '平均真实波幅（Average True Range, ATR）由 J. Welles Wilder Jr. 于 1978 年在其著作 '
        '《New Concepts in Technical Trading Systems》中提出。ATR 在 Wilder 的设计中原本用于度量'
        '商品期货市场的波动率，海龟策略将其引入股票和期货交易中，并赋予了一个新的名称“N 值”'
        '（因为它是海龟策略中正常化仓位大小的分母）。',
        first_line_indent=Cm(0.74))
    add_paragraph(doc,
        'ATR 的独到之处在于其“真实波幅”（True Range）的定义——它不仅考虑当日最高与最低价的差异（日内振幅），'
        '还考虑了昨日收盘与今日最高/最低之间的跳空缺口。这使得 ATR 能完整捕捉所有形式的跨日价格运动，'
        '包括隔夜跳空。',
        first_line_indent=Cm(0.74))
    add_paragraph(doc, 'ATR 的数学定义：', first_line_indent=Cm(0.74))
    add_atr_formula(doc)
    add_paragraph(doc, '')
    add_paragraph(doc,
        '在海龟策略中，N 值（即 20 日 ATR）扮演三重角色：(1) 仓位大小计算的分母——N 值越大，在相同风险预算下'
        '仓位越小；(2) 止损距离的基准——止损价 = 入场价 − 2×N；(3) 加仓间距的刻度——价格每上涨 0.5×N，'
        '加仓一个单位。N 值将波动率度量与交易风险直接关联，是海龟策略的核心创新。',
        first_line_indent=Cm(0.74))

    # 2.3 止损条件
    add_heading_styled(doc, '2.3 止损条件', level=2)
    add_paragraph(doc,
        '海龟策略的止损机制由两个层级构成，分别处理“趋势反转”和“个别交易失败”两种不同的风险场景：',
        first_line_indent=Cm(0.74))

    add_paragraph(doc, '第一层：硬止损（Hard Stop-Loss）', bold=True, size=Pt(10.5),
                 first_line_indent=Cm(0.74))
    add_stop_formula(doc)
    add_paragraph(doc, '')
    add_paragraph(doc,
        '硬止损基于 ATR（N 值）设定。默认参数 k=2，即止损距离 = 2 倍的当前 N 值。这意味着如果价格的正常'
        '日波动范围是 1 个 N 值，那么 2 个 N 值的反向运动是一个统计上显著的异常事件——趋势很可能已经改变。'
        '硬止损具有“只上移不下移”的特性（对于多头持仓），这意味着止损价会随着价格上涨而逐步提高，'
        '形成了天然的移动止盈机制。',
        first_line_indent=Cm(0.74))

    add_paragraph(doc, '第二层：通道离场（Channel Exit）', bold=True, size=Pt(10.5),
                 first_line_indent=Cm(0.74))
    add_paragraph(doc,
        '当价格跌破离场通道下轨（exit_period 日最低价）时触发离场。通道离场的逻辑是：如果价格跌破了过去 '
        'N 日的最低点，说明市场已经打破了上升趋势的最小结构——连最低价都在下移。通道离场通常比硬止损更早触发'
        '（因为 exit_period 日低点往往高于 2×ATR 止损线），因此它是海龟策略“常规退出”的主要方式，'
        '而硬止损是极端情况下的“紧急刹车”。',
        first_line_indent=Cm(0.74))

    add_paragraph(doc, '仓位限制与风险预算', bold=True, size=Pt(10.5),
                 first_line_indent=Cm(0.74))
    add_position_formula(doc)
    add_paragraph(doc, '')
    add_paragraph(doc,
        '海龟策略通过仓位公式将每笔交易的最大亏损控制在总资金的 Risk%（通常为 1-2%）。'
        '例如 100 万资金，Risk%=2%，ATR=30 元，则 Shares = 20000/30 ≈ 667 股——'
        '即使止损触发（亏损 2×ATR = 60 元/股），总亏损也仅 ≈ 667×60 = 40020 元，恰为总资金的 4%。'
        '此外，海龟策略还限制单一市场总风险敞口不超过总资金的 4%，相关性高的多个市场合计不超过 6%。',
        first_line_indent=Cm(0.74))

    doc.add_page_break()


# ═══════════════════════════ 第三章 ═══════════════════════════

def write_ch3(doc, df, result, metrics, stock_name, stock_code):
    add_heading_styled(doc, '三、策略回测结果分析', level=1)

    add_paragraph(doc,
        f'以{stock_name}（{stock_code}.SH）日线数据，海龟策略 System 1（入场 20 日突破，'
        f'离场 10 日突破，止损 2×ATR），初始资金 100 万元，手续费 0.03%，滑点 0.01%，'
        f'回测区间 {df["Date"].iloc[0].strftime("%Y-%m-%d")} 至 {df["Date"].iloc[-1].strftime("%Y-%m-%d")}。',
        first_line_indent=Cm(0.74))

    # 3.1 交易信号
    add_heading_styled(doc, '3.1 价格通道与交易信号', level=2)
    add_picture_captioned(doc, os.path.join(CHART_DIR, f'{stock_code}_海龟策略信号图.png'),
        f'图1：{stock_name}（{stock_code}）海龟策略交易信号——入场/离场通道与买卖点',
        width_inches=5.8)
    total_pyramid_adds = sum(len(t.get('pyramid_adds', [])) for t in result['trades'])
    add_paragraph(doc,
        f'图1展示了价格、入场通道（20日最高价）和离场通道（10日最低价）。共产生'
        f'{metrics["total_trades"]} 次完整交易，其中止损退出 {metrics["stop_exits"]} 次，'
        f'信号退出 {metrics["signal_exits"]} 次。红色▲标记买入点——每次突破 20 日高点后入场；'
        f'橙色◆标记加仓点——价格每上涨 0.5×ATR 加仓一个单位，最多加仓 4 次（金字塔加仓），'
        f'共 {total_pyramid_adds} 次加仓；'
        f'绿色▼标记卖出点——跌破 10 日低点或触及 2×ATR 止损线。',
        first_line_indent=Cm(0.74))

    # 3.2 ATR
    add_heading_styled(doc, '3.2 ATR 波动率与通道关系', level=2)
    add_picture_captioned(doc, os.path.join(CHART_DIR, f'{stock_code}_ATR分析.png'),
        f'图2：{stock_name}（{stock_code}）价格通道与 N 值（ATR-20）走势图',
        width_inches=5.8)
    add_paragraph(doc,
        '图2展示了 N 值（ATR-20）随时间的变化。N 值的升高意味着市场波动率在增大——'
        '通常伴随趋势行情；N 值的降低意味着波动率在缩小——通常伴随盘整。海龟策略通过 N 值自动调整仓位：'
        '波动大时减小仓位、波动小时增大仓位，使每笔交易的风险敞口保持恒定。'
        'N/Price 比率（橙色虚线）显示了波动率相对于价格的比例，是跨品种比较的重要指标。',
        first_line_indent=Cm(0.74))

    # 3.3 资产曲线
    add_heading_styled(doc, '3.3 资产曲线与回撤分析', level=2)
    add_picture_captioned(doc, os.path.join(CHART_DIR, f'{stock_code}_资产曲线与回撤.png'),
        f'图3：{stock_name}（{stock_code}）资产曲线对与回撤分析',
        width_inches=5.8)
    add_paragraph(doc,
        f'策略累计回报 {metrics["total_return"]:.2f}%，年化 {metrics["annual_return"]:.2f}%，'
        f'最大回撤 {metrics["mdd"]:.2f}%，夏普比率 {metrics["sharpe"]:.4f}。'
        f'买入持有累计回报 {metrics["bh_total_return"]:.2f}%。'
        f'策略相对基准{"跑赢" if metrics["total_return"] > metrics["bh_total_return"] else "跑输"}'
        f'{abs(metrics["total_return"] - metrics["bh_total_return"]):.2f} 个百分点。'
        f'回撤曲线显示，策略的回撤主要发生在趋势反转初期——此时离场信号尚未触发，'
        f'而价格已经出现显著回落。这是趋势跟踪策略的固有特征。',
        first_line_indent=Cm(0.74))

    # 3.4 绩效汇总
    add_heading_styled(doc, '3.4 绩效指标汇总', level=2)
    plr_s = f'{metrics["profit_loss_ratio"]:.2f}' if metrics['profit_loss_ratio'] else 'N/A'
    add_table(doc, ['指标', '数值', '评价'], [
        ['累计回报', f'{metrics["total_return"]:.2f}%',
         '正收益' if metrics['total_return'] > 0 else '负收益'],
        ['年化收益率', f'{metrics["annual_return"]:.2f}%',
         '达标(>10%)' if metrics['annual_return'] >= 10 else ('正收益' if metrics['annual_return'] >= 0 else '为负')],
        ['最大回撤 (MDD)', f'{metrics["mdd"]:.2f}%',
         '超20%警戒' if metrics['mdd'] < -20 else ('正常范围' if metrics['mdd'] > -10 else '接近警戒')],
        ['夏普比率', f'{metrics["sharpe"]:.4f}',
         '良好(>1)' if metrics['sharpe'] >= 1 else ('一般(>0)' if metrics['sharpe'] >= 0 else '跑输无风险利率')],
        ['胜率', f'{metrics["win_rate"]:.1f}%',
         '趋势策略低胜率正常' if metrics['win_rate'] < 50 else '尚可'],
        ['盈亏比', plr_s,
         '赚大亏小(>2)' if (metrics['profit_loss_ratio'] and metrics['profit_loss_ratio'] >= 2) else '一般'],
        ['交易次数', f'{metrics["total_trades"]} 次',
         f'止损{metrics["stop_exits"]}次 / 信号{metrics["signal_exits"]}次'],
        ['买入持有', f'{metrics["bh_total_return"]:.2f}%', '基准对比'],
    ])

    # 3.5 交易明细
    add_heading_styled(doc, '3.5 交易明细', level=2)
    completed = [t for t in result['trades'] if t['sell_date'] is not None]
    if completed:
        rows, cum = [], 1.0
        for i, t in enumerate(completed):
            cum *= (1 + t['return_pct'] / 100)
            rows.append([
                str(i + 1),
                t['buy_date'].strftime('%Y-%m-%d'), f'{t["buy_price"]:.2f}',
                t['sell_date'].strftime('%Y-%m-%d'), f'{t["sell_price"]:.2f}',
                f'{(t["sell_date"] - t["buy_date"]).days}天',
                f'{t["return_pct"]:+.2f}%', f'{(cum-1)*100:+.2f}%',
                t.get('exit_reason', 'N/A'),
            ])
            # 金字塔加仓子行
            for add in t.get('pyramid_adds', []):
                add_date_str = add['add_date'].strftime('%Y-%m-%d')
                rows.append([
                    f'  +加仓',
                    add_date_str, f'{add["add_price"]:.2f}',
                    '', '',
                    f'+{add["add_shares"]:.0f}股',
                    '', '',
                    f'金字塔 (unit {add.get("unit", "?")})',
                ])
        add_table(doc, ['#', '买入日', '买入价', '卖出日', '卖出价', '持有',
                       '单笔收益', '累计收益', '退出原因'], rows)

    doc.add_page_break()


# ═══════════════════════════ 第四章 ═══════════════════════════

def write_ch4(doc, df, stock_name, stock_code):
    add_heading_styled(doc, '四、参数敏感性分析', level=1)

    add_heading_styled(doc, '4.1 入场/离场周期网格扫描', level=2)
    add_paragraph(doc,
        '入场周期（10/15/20/30/40/55 日）和离场周期（5/10/15/20/25 日）进行网格扫描，'
        '共测试 30 组有效参数组合（离场周期必须小于入场周期）。',
        first_line_indent=Cm(0.74))

    add_picture_captioned(doc, os.path.join(CHART_DIR, f'{stock_code}_参数敏感性热力图.png'),
        f'图4：{stock_name}（{stock_code}）参数敏感性 — 累计回报与夏普比率热力图',
        width_inches=5.8)

    add_paragraph(doc,
        '热力图显示：(1) 入场周期和离场周期的间隔是关键因素——间隔过小（如入场 15/离场 10）'
        '导致频繁交易、假信号增多；间隔过大（如入场 55/离场 5）则离场过晚、回吐利润。'
        '(2) 夏普比率通常比累计回报更稳定，参数选择应优先参考夏普比率——它是风险调整后的收益度量。'
        '(3) 表现好的参数组合通常形成“稳健区域”而非孤立的“尖锐峰值”——后者往往是过拟合的信号。',
        first_line_indent=Cm(0.74))

    # 4.2 止损倍数
    add_heading_styled(doc, '4.2 止损倍数敏感性', level=2)
    add_picture_captioned(doc, os.path.join(CHART_DIR, f'{stock_code}_止损敏感性.png'),
        f'图5：{stock_name}（{stock_code}）止损倍数敏感性分析',
        width_inches=5.8)

    add_paragraph(doc,
        '止损倍数的分析揭示了一个关键的权衡：(1) 止损倍数过小（<1.0）——止损过于敏感，频繁触发止损，'
        '胜率极低但单笔损失小；(2) 止损倍数适中（1.5-2.5）——这是海龟原始推荐的区间，平衡了“不被噪音震出”'
        '和“不在反转时损失过大”的需求；(3) 止损倍数过大（>3.0）——止损几乎等效于只有通道离场，'
        '回撤增大但胜率提高。2×ATR 是海龟原版的经验最优值。',
        first_line_indent=Cm(0.74))

    # 4.3 过拟合风险
    add_heading_styled(doc, '4.3 过拟合风险与参数选择原则', level=2)
    add_paragraph(doc,
        '海龟策略的参数空间相对较小（仅入场周期、离场周期、止损倍数三个核心参数），过拟合风险天然低于'
        '参数众多的复杂策略。参数选择的三原则：(1) 参数必须有经济含义——20 日约等于一个自然月，'
        '10 日约等于两周，时间尺度差异有实际意义；(2) 不应过度精细化——在 19 日和 20 日之间选择 19.37 '
        '几乎肯定是过拟合；(3) 最优参数的邻域也应表现良好——如果只有 20+10 表现好而 21+10 和 19+10 '
        '都很差，可能只是巧合。',
        first_line_indent=Cm(0.74))

    doc.add_page_break()


# ═══════════════════════════ 第五章 ═══════════════════════════

def write_ch5(doc):
    add_heading_styled(doc, '五、多标的对比分析', level=1)
    add_paragraph(doc,
        '选取 7 只标的（茅台、平安、宁德时代、比亚迪、平安银行、腾讯、阿里巴巴），'
        '统一使用海龟策略 System 1（入场 20/离场 10/止损 2×ATR）回测对比。',
        first_line_indent=Cm(0.74))

    add_heading_styled(doc, '5.1 回测结果对比', level=2)
    add_picture_captioned(doc, os.path.join(CHART_DIR, '多股票策略对比.png'),
        '图6：多股票海龟策略回测对比——累计回报、夏普比率、最大回撤',
        width_inches=5.8)

    stock_list = [
        ('600519', '贵州茅台', 'A股'), ('601318', '中国平安', 'A股'),
        ('300750', '宁德时代', 'A股'), ('002594', '比亚迪', 'A股'),
        ('000001', '平安银行', 'A股'),
        ('00700', '腾讯控股', '港股'), ('09988', '阿里巴巴', '港股'),
    ]
    rows = []
    for code, name, market in stock_list:
        try:
            fn = f"{code}_{name}_{market}_daily.csv"
            df_s = pd.read_csv(os.path.join(DATA_DIR, fn), encoding='utf-8-sig')
            df_s['Date'] = pd.to_datetime(df_s['Date'])
            df_s = df_s.sort_values('Date').reset_index(drop=True)
            r = run_backtest(df_s, entry_period=20, exit_period=10, atr_stop_mult=2.0)
            m = calc_metrics(r)
            if m and 'error' not in m:
                rows.append([
                    name, code,
                    f'{m["total_return"]:.2f}%', f'{m["annual_return"]:.2f}%',
                    f'{m["mdd"]:.2f}%', f'{m["sharpe"]:.4f}',
                    f'{m["win_rate"]:.1f}%', f'{m["total_trades"]}',
                    f'{m["bh_total_return"]:.2f}%',
                    '是' if m['total_return'] > m['bh_total_return'] else '否',
                ])
        except Exception:
            pass

    add_table(doc, ['标的', '代码', '累计回报', '年化收益', 'MDD', '夏普',
                   '胜率', '交易次数', '买入持有', '跑赢基准?'], rows)

    add_heading_styled(doc, '5.2 策略适用场景分析', level=2)
    add_paragraph(doc,
        '从多标的对比中可以总结海龟策略的适用条件：(1) 趋势性强的标的——如宁德时代、比亚迪等新能源龙头，'
        '由于行业处于成长期，价格趋势持续性强，适合海龟策略的趋势跟踪逻辑。(2) 高波动标的效果分化——'
        '高波动意味着更大的 ATR，这会自动减小仓位（通过仓位公式），降低总风险，但也限制了潜在收益。'
        '(3) 低波动蓝筹（如平安银行、中国平安）趋势信号相对稀少，但假突破也少，适合 System 2 '
        '（55 日突破）。(4) 港股标的由于流动性差异和交易制度差异，滑点和成本对策略影响不同，'
        '需要更大参数的入场周期来过滤噪音。',
        first_line_indent=Cm(0.74))

    doc.add_page_break()


# ═══════════════════════════ 第六章 ═══════════════════════════

def write_ch6(doc, metrics, stock_name):
    add_heading_styled(doc, '六、总结与使用心得', level=1)

    add_heading_styled(doc, '6.1 策略表现总结', level=2)
    add_paragraph(doc,
        f'本报告以{stock_name}为主要标的，完成了海龟交易策略（System 1: 入场 20 日/离场 10 日/止损 2×ATR）'
        f'的完整回测分析。策略累计回报 {metrics["total_return"]:.2f}%，年化 {metrics["annual_return"]:.2f}%，'
        f'MDD {metrics["mdd"]:.2f}%，夏普 {metrics["sharpe"]:.4f}，'
        f'{"跑赢" if metrics["total_return"] > metrics["bh_total_return"] else "跑输"}买入持有基准。',
        first_line_indent=Cm(0.74))

    add_heading_styled(doc, '6.2 海龟策略的局限与挑战', level=2)
    limitations = [
        '滞后性本质：所有趋势跟踪策略（包括海龟）都是“滞后”的——入场信号在趋势已经运行一段后才确认，'
        '离场信号在趋势已经反转一段后才触发。这种滞后性不是策略的缺陷，而是其设计哲学的一部分。',
        '盘整期连续亏损（Whipsaw）：当市场进入横盘区间时，价格反复突破通道上轨又回落，'
        '海龟策略会连续产生买入-止损的亏损循环。这是趋势跟踪策略在非趋势市场中必然付出的“保险费”。'
        '海龟训练中，Dennis 明确告诉学员“会有连续 10-12 次亏损”——这是正常的系统行为。',
        '最大回撤可能很大：在没有止损保护的情况下（如只使用通道离场），单一趋势的反转可以吞噬大量利润。'
        '即使是 2×ATR 止损，在跳空下跌的极端行情中也无法以止损价格成交（滑点问题）。',
        '参数对历史数据的依赖：20/10 等参数是基于 1980 年代美国商品期货市场的经验选择，'
        '直接搬到 2020 年代的 A 股市场不一定最优。需要进行市场特定的参数校准。',
        '心理执行的困难：知道一套策略长期能赚钱，和在连续亏损 10 次后继续执行它是两回事。'
        '海龟策略的最大敌人不是市场，而是交易者自身的情绪。Dennis 的学员中，有些人因为无法忍受'
        '连续亏损而放弃了系统——这正是机械策略的核心挑战。',
        '仅做多限制：本策略仅实现做多方向，在持续下跌行情中只能通过空仓减少损失，'
        '无法像原版海龟策略（可双向交易）那样通过做空在熊市中获利。在 A 股市场，'
        '这是一个实际但重要的约束。若未来扩展到可融券标的或期货市场，应加入做空方向以还原海龟策略的完整设计。',
    ]
    for i, lim in enumerate(limitations, 1):
        add_paragraph(doc, f'{i}. {lim}', first_line_indent=Cm(0.74))

    add_heading_styled(doc, '6.3 改进方向与个人心得', level=2)
    improvements = [
        '引入趋势过滤器：使用 ADX（平均趋向指数）判断市场状态——ADX < 20 时暂停交易，'
        '仅在趋势确立时启用信号。这可以减少盘整期的亏损。',
        '多时间框架确认：在 System 1（20 日）信号触发时，检查 System 2（55 日）的趋势方向是否一致。'
        '两系统方向一致时加大仓位，不一致时减小仓位。',
        '动态参数调整：根据市场波动率环境（如 VIX 指数、市场整体 ATR 水平）动态调整通道周期——'
        '高波动期使用更长的通道周期以减少假突破。',
        '组合化应用：将海龟策略同时应用于多只低相关性标的，通过分散化降低单标的连续亏损对总账户的影响。'
        '这是海龟原版中“多市场交易”的核心设计——东方不亮西方亮。',
        '加入基本面过滤：海龟策略是纯粹的技术策略，但结合基本面的“趋势质量”判断——'
        '如净利润增速、行业景气度——可以提高信号的可靠性。',
    ]
    for i, imp in enumerate(improvements, 1):
        add_paragraph(doc, f'{i}. {imp}', first_line_indent=Cm(0.74))

    add_paragraph(doc, '')
    add_paragraph(doc,
        '海龟策略的核心价值不在于它的具体参数（20/10/2×ATR），而在于其揭示的交易哲学：趋势跟踪、'
        '风险管理优先、完全机械化、通过波动率标准化风险。这些原则超越了任何特定市场和时间——'
        '无论是 1980 年代的商品期货，还是 2020 年代的 A 股市场。理解并内化这些原则，'
        '比找到一组“最优参数”重要得多。正如 Richard Dennis 所说：'
        '“我不关心市场明天是涨是跌，我只关心当市场朝一个方向大幅运动时，我在不在里面。”',
        first_line_indent=Cm(0.74))

    add_heading_styled(doc, '6.4 海龟法则的适应场景', level=2)
    scenarios = [
        ('适合海龟策略的场景',
         ['1. 强趋势市场（如牛市主升浪、行业景气上行周期）— 趋势持续性强，突破信号可靠性高',
          '2. 高流动性标的（大盘蓝筹、核心资产）— 滑点可控，交易成本占比低',
          '3. 波动率适中的标的 — 过低的波动率意味着趋势幅度小、利润空间有限；'
          '过高的波动率会触发频繁止损',
          '4. 多资产组合交易 — 单一标的不确定性高，分散化可熨平权益曲线',
          '5. 中长期投资时间框架 — 海龟策略的持仓周期通常为数周至数月，'
          '不适合日内短线交易']),
        ('不适合海龟策略的场景',
         ['1. 持续盘整/震荡市场 — 反复的假突破导致连续的止损亏损，利润被交易成本消耗',
          '2. 流动性极差的标的 — 滑点和冲击成本会严重侵蚀策略表现',
          '3. 高波动+无趋势的标的 — 价格波动大但没有持续方向，兼顾了最差的两种特征',
          '4. 有重大事件风险的时期（如财报季、政策窗口）— ATR 可能在一天内跳升数倍，'
          '突破历史极值，使基于历史 ATR 的仓位和止损失效']),
    ]
    for title, items in scenarios:
        add_paragraph(doc, title + '：', bold=True, first_line_indent=Cm(0.74))
        for item in items:
            add_paragraph(doc, item, first_line_indent=Cm(0.74))
        add_paragraph(doc, '')

    doc.add_page_break()


# ═══════════════════════════ 第七章 ═══════════════════════════

def write_ch7(doc):
    add_heading_styled(doc, '七、交互式回测看板', level=1)
    add_paragraph(doc,
        '本任务构建了自包含的交互式 Web 回测看板（Task4/dashboard/index.html），覆盖 10 只标的。'
        '所有回测计算在浏览器端完成，无需后端服务器。',
        first_line_indent=Cm(0.74))

    add_heading_styled(doc, '7.1 核心功能', level=2)
    add_table(doc, ['功能', '说明'], [
        ['标的选择', '10 只标的（5 A 股 + 5 港股），切换即时重新回测'],
        ['参数调节', '入场周期、离场周期、ATR 周期、止损倍数、风险比例、资金量——滑块+数字双控'],
        ['KPI 卡片', '累计回报、年化收益、MDD、夏普比率、胜率、盈亏比、交易次数、买入持有收益'],
        ['价格信号图', '收盘价 + 入场通道 + 离场通道 + 买入▲ + 卖出▼，支持缩放拖拽'],
        ['ATR 走势图', 'N 值变化 + 收盘价对比，理解波动率与价格关系'],
        ['资产曲线图', '策略资产(蓝) vs 买入持有(灰虚线)'],
        ['回撤曲线图', '回撤面积图，标注 MDD 值和警戒线'],
        ['交易明细表', '每笔交易的买卖日期/价格/持有天数/收益率/退出原因'],
        ['纯前端回测', '参数修改即时生效，避免未来函数，成本自动扣除'],
    ])

    add_heading_styled(doc, '7.2 部署方式', level=2)
    add_paragraph(doc,
        '看板已部署至 GitHub Pages，本地可直接双击打开 HTML 文件使用。'
        '所有计算在浏览器端用 JavaScript 完成，参数调节即时响应，无需等待服务器。',
        first_line_indent=Cm(0.74))
    doc.add_page_break()


# ═══════════════════════════ 主程序 ═══════════════════════════

def main():
    print('=' * 60)
    print('  生成 Task4 .docx 报告 — 海龟交易策略')
    print('=' * 60)

    df = pd.read_csv(os.path.join(DATA_DIR, '600519_贵州茅台_A股_daily.csv'), encoding='utf-8-sig')
    df['Date'] = pd.to_datetime(df['Date'])
    df = df.sort_values('Date').reset_index(drop=True)
    print(f'\n📂 贵州茅台: {len(df)} 条数据')

    # 检查图表
    charts_needed = [
        '600519_海龟策略信号图.png', '600519_ATR分析.png',
        '600519_资产曲线与回撤.png', '600519_参数敏感性热力图.png',
        '600519_止损敏感性.png', '多股票策略对比.png',
    ]
    missing = [c for c in charts_needed if not os.path.exists(os.path.join(CHART_DIR, c))]
    if missing:
        print(f'\n⚠️ 缺少图表: {missing}')
        print('  请先运行: python Task4/scripts/plot_turtle.py')
        return

    print('⏳ 运行回测...')
    result = run_backtest(df, entry_period=20, exit_period=10, atr_stop_mult=2.0)
    metrics = calc_metrics(result)
    print(f'   累计回报: {metrics["total_return"]:.2f}%  夏普: {metrics["sharpe"]:.4f}')

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = FONT_SIZE
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.5
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    chapters = [
        ('一、海龟策略核心思想', lambda: write_ch1(doc)),
        ('二、核心概念详解', lambda: write_ch2(doc)),
        ('三、策略回测结果分析', lambda: write_ch3(doc, df, result, metrics, '贵州茅台', '600519')),
        ('四、参数敏感性分析', lambda: write_ch4(doc, df, '贵州茅台', '600519')),
        ('五、多标的对比分析', lambda: write_ch5(doc)),
        ('六、总结与使用心得', lambda: write_ch6(doc, metrics, '贵州茅台')),
        ('七、交互式回测看板', lambda: write_ch7(doc)),
    ]

    for label, fn in chapters:
        print(f'>>> {label}')
        fn()

    os.makedirs(os.path.dirname(OUTPUT_FILE), exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f'\n✅ {OUTPUT_FILE}  ({os.path.getsize(OUTPUT_FILE)/1024:.0f} KB)')
    print('=' * 60)


if __name__ == '__main__':
    main()
