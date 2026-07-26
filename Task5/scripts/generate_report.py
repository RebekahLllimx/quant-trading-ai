#!/usr/bin/env python3
"""Generate the editable TASK5 Word draft in the same style as TASK4."""

from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
TEMPLATE = ROOT / "Task4" / "Rebecca+Task4.docx"
OUTPUT = ROOT / "Task5" / "Rebecca+Task5.docx"
DATA_DIR = ROOT / "data" / "task5"
CHART_DIR = ROOT / "artifacts" / "charts" / "task5"

BODY_FONT = "宋体"
BODY_SIZE = 10.5
USABLE_DXA = 8640


def clear_document_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def clear_page_furniture(doc: Document) -> None:
    for section in doc.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                paragraph._element.getparent().remove(paragraph._element)
            container.add_paragraph("")


def set_run(run, *, size=BODY_SIZE, bold=False, font=BODY_FONT, italic=False):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    rfonts.set(qn("w:eastAsia"), font)
    return run


def set_paragraph(
    paragraph,
    *,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    before=0,
    after=0,
    line=1.5,
    first_line=True,
    keep_next=False,
):
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.first_line_indent = Inches(0.292) if first_line else None
    fmt.keep_with_next = keep_next
    return paragraph


def add_h1(doc: Document, text: str):
    p = doc.add_paragraph(style="Normal")
    set_paragraph(p, before=12, after=6, first_line=False, keep_next=True)
    set_run(p.add_run(text), size=14, bold=True)
    return p


def add_h2(doc: Document, text: str):
    p = doc.add_paragraph(style="Normal")
    set_paragraph(p, before=8, after=4, first_line=False, keep_next=True)
    set_run(p.add_run(text), size=12, bold=True)
    return p


def add_body(doc: Document, text: str, *, first_line=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph(style="Normal")
    set_paragraph(p, first_line=first_line, align=align)
    set_run(p.add_run(text))
    return p


def add_equation(doc: Document, text: str):
    p = doc.add_paragraph(style="Normal")
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, before=2, after=2)
    set_run(p.add_run(text), size=11.5, font="Times New Roman", italic=True)
    return p


def add_code(doc: Document, lines: list[str]):
    for line in lines:
        p = doc.add_paragraph(style="Normal")
        set_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False, line=1.0)
        p.paragraph_format.left_indent = Inches(0.25)
        set_run(p.add_run(line if line else " "), size=9, font="Courier New")


def shade_cell(cell, fill="D9D9D9"):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tcpr = tc.get_or_add_tcPr()
    tcmar = tcpr.first_child_found_in("w:tcMar")
    if tcmar is None:
        tcmar = OxmlElement("w:tcMar")
        tcpr.append(tcmar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcmar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tcmar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    trpr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trpr.append(tbl_header)


def set_table_geometry(table, widths: list[int]):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblpr = table._tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(sum(widths)))
    tblw.set(qn("w:type"), "dxa")
    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), "0")
    tblind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(widths[idx]))
            tcw.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_table(
    doc: Document,
    number: int,
    title: str,
    headers: list[str],
    rows: list[list[str]],
    widths: list[int],
    *,
    left_columns: set[int] | None = None,
):
    caption = doc.add_paragraph(style="Normal")
    set_paragraph(caption, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, before=4, after=3, keep_next=True)
    set_run(caption.add_run(f"表 {number}：{title}"), bold=True)

    table = doc.add_table(rows=1, cols=len(headers), style="Table Grid")
    set_repeat_header(table.rows[0])
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, line=1.15)
        set_run(p.add_run(str(text)), bold=True)

    left_columns = left_columns or set()
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            alignment = WD_ALIGN_PARAGRAPH.LEFT if idx in left_columns else WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph(p, align=alignment, first_line=False, line=1.15)
            set_run(p.add_run(str(value)))
    set_table_geometry(table, widths)
    spacer = doc.add_paragraph(style="Normal")
    set_paragraph(spacer, first_line=False, line=1.0, after=2)
    return table


def add_figure(doc: Document, path: Path, number: int, title: str, explanation: str, width=6.0):
    p = doc.add_paragraph(style="Normal")
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, before=4, line=1.0, keep_next=True)
    p.add_run().add_picture(str(path), width=Inches(width))
    caption = doc.add_paragraph(style="Normal")
    set_paragraph(caption, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, after=3, keep_next=True)
    set_run(caption.add_run(f"图 {number}：{title}"), bold=True)
    add_body(doc, explanation)


def fmt_pct(value: float, digits=2) -> str:
    return f"{value * 100:.{digits}f}%"


def main() -> None:
    metrics = pd.read_csv(DATA_DIR / "processed" / "task5_model_metrics.csv")
    baseline = metrics[metrics["model"] == "majority_baseline"].iloc[0]
    metrics = metrics[metrics["model"].isin(["logistic_regression", "decision_tree", "random_forest"])].copy()
    candidates = pd.read_csv(DATA_DIR / "processed" / "task5_candidate_metrics.csv")
    yearly = pd.read_csv(DATA_DIR / "processed" / "task5_yearly_metrics.csv")
    splits = pd.read_csv(DATA_DIR / "metadata" / "split_summary.csv")
    quality = json.loads((DATA_DIR / "metadata" / "data_quality_report.json").read_text(encoding="utf-8"))

    doc = Document(TEMPLATE)
    clear_document_body(doc)
    clear_page_furniture(doc)

    add_h1(doc, "一、分类型机器学习算法")
    add_body(doc, "本任务预测一只股票在未来5个交易日内是否上涨。上涨记为1，不上涨记为0，因此属于二分类问题。模型读取的是观察日及以前可以计算出的价格、波动率和成交量指标，未来5日收益只用于生成标签。")

    add_h2(doc, "1.1 逻辑回归")
    add_body(doc, "逻辑回归先把特征加权求和，再用Sigmoid函数把线性得分转为0到1之间的概率。系数为正表示其他条件不变时，该特征增大会提高模型给出的上涨概率；系数为负则相反。本次先对特征做标准化，所以系数的绝对值可以用来比较模型对不同变量的敏感程度。")
    add_equation(doc, "P(y = 1 | X) = 1 / (1 + exp[-(β0 + β1x1 + ... + βpxp)])")
    add_body(doc, "逻辑回归的概率输出与最终分类是两件事。例如可以用0.5作为分类阈值，但ROC和AUC会考察所有可能阈值下的排序。逻辑回归训练快、概率与系数易于解释，但基本决策边界是线性的，难以直接表示复杂的非线性与特征交互。")

    add_h2(doc, "1.2 决策树")
    add_body(doc, "决策树反复选择一个特征和切分阈值，使切分后的子节点尽量包含同一类样本。例如，树可以先判断价格是否明显高于20日均线，再根据波动率将样本继续细分。叶节点中上涨样本的比例可以作为上涨概率。这种规则能表达阈值效应和特征交互，也不要求标准化。")
    add_body(doc, "决策树使用贪心切分，一棵过深的树容易把训练期噪声当成规律。本次用max_depth限制树深，并用min_samples_leaf保证每个叶节点有足够样本，目的是牺牲部分样本内拟合能力，换取更稳定的样本外表现。")

    add_h2(doc, "1.3 随机森林")
    add_body(doc, "随机森林训练多棵有差异的决策树。每棵树使用Bootstrap抽样得到的训练样本，并在节点切分时只考察随机抽取的一部分特征。这两层随机性用来降低各棵树之间的相关性，最终对概率取平均，通常可以降低单棵树的方差。")
    add_body(doc, "随机森林能处理非线性和交互关系，但它不会自动创造信号。如果原始特征对未来涨跌只有很弱的信息，增加树数只能使这种弱关系的估计更稳定，不保证样本外AUC显著提高。")

    doc.add_page_break()
    add_table(
        doc,
        1,
        "三种分类模型的比较",
        ["模型", "基本方法", "优点", "局限"],
        [
            ["逻辑回归", "线性得分转为上涨概率", "训练快，概率和系数容易解释", "难以直接表示复杂的非线性关系"],
            ["决策树", "按特征阈值逐层切分", "规则直观，不需要标准化", "容易过拟合，对样本变化较敏感"],
            ["随机森林", "多棵随机决策树取平均", "比单树稳定，能表示变量交互", "计算量较大，整体规则不如单树直观"],
        ],
        [1200, 2300, 2450, 2690],
        left_columns={1, 2, 3},
    )

    add_h1(doc, "二、机器学习模型评价指标")
    add_h2(doc, "2.1 混淆矩阵")
    add_body(doc, "混淆矩阵把实际标签与预测标签放在一起统计。本任务把上涨作为正类。TP表示实际上涨且模型也预测上涨，TN表示实际不涨且模型预测不涨。FP是模型给出上涨信号但实际未上涨，FN则是漏掉的上涨样本。在真实交易中，FP可能对应一次无效交易和成本，FN对应错过机会，两种错误的经济代价并不相同。")
    add_table(
        doc,
        2,
        "二分类混淆矩阵",
        ["实际标签", "预测不涨（0）", "预测上涨（1）"],
        [["实际不涨（0）", "TN：判断正确", "FP：错误预测上涨"], ["实际上涨（1）", "FN：漏掉上涨", "TP：识别出上涨"]],
        [2100, 3270, 3270],
    )

    add_h2(doc, "2.2 Accuracy、Precision、Recall与F1")
    add_body(doc, "Accuracy是全部样本中预测正确的比例。Precision回答“模型给出上涨信号时，其中有多少真的上涨”；Recall回答“全部实际上涨样本中，模型找出了多少”。F1是Precision和Recall的调和平均，任何一项很低都会拉低F1。这四项指标都依赖分类阈值，阈值从0.5调低后，模型通常会预测更多上涨，Recall提高，但FP也可能增加。")
    add_equation(doc, "Accuracy = (TP + TN) / (TP + TN + FP + FN)")
    add_equation(doc, "Precision = TP / (TP + FP)    Recall = TP / (TP + FN)")
    add_equation(doc, "F1 = 2 × Precision × Recall / (Precision + Recall)")

    add_h2(doc, "2.3 ROC曲线与AUC")
    add_body(doc, "分类阈值改变后，真阳性率TPR和假阳性率FPR也会变化。ROC曲线把不同阈值下的FPR放在横轴，把TPR放在纵轴。曲线越靠近左上角，说明模型在保持较低误报的同时能找到更多上涨样本。AUC是ROC曲线下面积，它不固定在某一个阈值上。")
    add_body(doc, "AUC还可以理解为：随机抽取一个上涨样本和一个不涨样本，模型把上涨样本给出更高分数的概率。AUC等于0.5时接近随机排序，等于1才是完全正确排序。因此AUC为0.521并不表示分类正确率为52.1%，它表示上述成对排序的正确率约为52.1%。计算时应输入概率或连续得分，而不是0和1的硬分类结果。")

    add_h1(doc, "三、数据、标签与特征设计")
    add_h2(doc, "3.1 数据来源与股票池")
    add_body(doc, f"日线数据通过Tushare Pro获取，少量缺口使用AKShare补充。股票池按2018年1月日成交额中位数排序后冻结为{quality['symbols']}只沪深A股。原始数据从2017年开始，其中2017年主要用于计算滚动指标；模型样本从{quality['date_start']}开始，到{quality['date_end']}结束。最终数据共有{quality['final_rows']:,}行，股票与日期组合没有重复，建模字段没有剩余缺失值。")

    add_h2(doc, "3.2 标签定义")
    add_body(doc, "标签是模型需要预测的结果，特征是做出预测时已经可知的信息。本任务用未来5个交易日的前复权收益生成标签：收益大于0记为1，否则记为0。未来价格只参与标签计算，不进入模型特征。这条边界非常重要：如果把t+1日之后的行情用于构造X(t)，即使结果很好也属于未来信息泄漏。")
    add_equation(doc, "future_return_5d(t) = Close(t + 5) / Close(t) - 1")
    add_equation(doc, "y(t) = 1, if future_return_5d(t) > 0; otherwise y(t) = 0")
    add_body(doc, "5日方向标签符合本次二分类要求，但也会丢失收益幅度。上涨0.1%与上涨10%都记为1，下跌的两种情况也同样被合并。此外，相邻日期的5日窗口大量重叠，样本数虽多，但实际独立信息量小于表面行数。")

    add_h2(doc, "3.3 特征构造")
    add_body(doc, "本次使用15个仅依赖历史价格和成交信息的比例型特征。用比例而不是价格绝对值，是为了让不同价格水平的股票可以放在同一模型中比较。这些特征不是15个完全独立的信息源，RSI、MACD、均线偏离和多期收益都由历史价格变换得到，彼此可能高度相关。")
    add_body(doc, "特征设计的思路是覆盖几种可检验的市场状态。短中期收益同时允许模型识别动量或反转；均线偏离、RSI和MACD描述价格所处的趋势位置；ATR、历史波动率与日内振幅表示当前风险状态；成交量比和成交额比用来识别异常活跃度；开收盘收益表示当日方向。所有变量均在观察日收盘后计算。")
    add_body(doc, "这套指标可以作为基础候选特征，但不能因为常见就视为已被验证的交易因子。真正的判断标准是时间样本外表现。本次AUC只略高于0.5，因此应当得出“存在很弱的排序信息”，而不是“这些指标已经验证有效”。")
    add_table(
        doc,
        3,
        "特征分组与变量",
        ["特征组", "变量", "用途"],
        [
            ["收益与趋势", "1日、5日、10日和20日收益；MA5、MA20、MA60偏离度", "描述近期涨跌和价格相对均线的位置"],
            ["趋势强弱", "RSI14；MACD相对价格", "描述趋势方向和强弱"],
            ["波动状态", "ATR14相对价格；20日波动率；日内振幅", "描述风险和日内价格范围"],
            ["交易活跃度", "20日成交量比；20日成交额比；开收盘收益", "描述成交活跃度和当日买卖压力"],
        ],
        [1500, 3900, 3240],
        left_columns={1, 2},
    )

    add_h2(doc, "3.4 训练集、验证集与测试集")
    split_names = {"train": "训练集", "validation": "验证集", "test": "测试集"}
    split_rows = []
    for _, row in splits.iterrows():
        split_rows.append([
            split_names[row["Split"]],
            f"{row['start']} 至 {row['end']}",
            f"{int(row['rows']):,}",
            str(int(row["symbols"])),
            fmt_pct(row["positive_rate"]),
        ])
    add_table(doc, 4, "样本的时间划分", ["数据段", "日期", "样本行", "股票数", "上涨比例"], split_rows, [1200, 3000, 1500, 1200, 1740])
    add_body(doc, "金融数据不能随机打乱。本次用2018年至2022年的数据训练模型，用2023年验证集选择参数，再用2024年至2025年的测试集评估一次。由于标签需要未来5个交易日，训练和验证边界分别删除495行和485行，避免前一数据段的标签跨入后一数据段。")
    add_figure(doc, CHART_DIR / "time_split.png", 1, "训练集、验证集与测试集的时间划分", "图1按月份展示三个数据段。2023年只用于选参数，2024年至2025年的测试数据不参与模型选择。这样得到的AUC比随机切分更接近实际使用场景。", width=6.0)

    add_h1(doc, "四、Python建模流程")
    add_h2(doc, "4.1 数据预处理与参数选择")
    add_body(doc, "模型训练前先用训练数据的0.5%和99.5%分位数对特征截尾。逻辑回归还需要StandardScaler标准化。截尾边界和标准化参数都只从训练数据估计。每类模型预先设置少量候选参数，按2023年验证集AUC选择其中一组。完成选择后，模型在训练集和验证集的合并数据上重新拟合，再生成测试集概率。")
    candidate_lookup = metrics.set_index("model")["selected_candidate"].to_dict()
    add_body(doc, f"验证集最终选择逻辑回归{candidate_lookup['logistic_regression']}，决策树{candidate_lookup['decision_tree']}，随机森林{candidate_lookup['random_forest']}。测试结果生成后没有再调整参数。")
    selected_candidates = candidates.sort_values(["model", "validation_auc"], ascending=[True, False]).groupby("model", as_index=False).first()
    selected_candidates = selected_candidates.set_index("model")
    add_table(
        doc,
        5,
        "验证集选定的模型参数",
        ["模型", "选定参数", "验证集AUC", "设计意图"],
        [
            ["逻辑回归", candidate_lookup["logistic_regression"], f"{selected_candidates.loc['logistic_regression', 'validation_auc']:.4f}", "在预设正则化强度中选验证AUC最高者"],
            ["决策树", candidate_lookup["decision_tree"], f"{selected_candidates.loc['decision_tree', 'validation_auc']:.4f}", "限制树深和叶节点样本数，抑制过拟合"],
            ["随机森林", candidate_lookup["random_forest"], f"{selected_candidates.loc['random_forest', 'validation_auc']:.4f}", "250棵树并限制复杂度，降低单树方差"],
        ],
        [1350, 2550, 1450, 3290],
        left_columns={1, 3},
    )
    add_body(doc, "三类模型的验证集AUC介于0.5173与0.5209之间，差距本来就很小。随机森林在2023年验证集上最高，但这只决定该模型采用哪组参数，不意味着它必然在2024至2025年测试期继续排名第一。")

    add_h2(doc, "4.2 主要代码")
    add_code(doc, [
        "# 按时间划分，并清除标签跨界样本",
        "train = df[(df.Date <= '2022-12-31') &",
        "           (df.label_end_date <= '2022-12-31')]",
        "valid = df[(df.Date > '2022-12-31') &",
        "           (df.label_end_date <= '2023-12-31')]",
        "test = df[df.Date > '2023-12-31']",
        "",
        "# 训练模型并用上涨概率计算AUC",
        "model.fit(X_train, y_train)",
        "prob_up = model.predict_proba(X_test)[:, 1]",
        "auc = roc_auc_score(y_test, prob_up)",
        "fpr, tpr, thresholds = roc_curve(y_test, prob_up)",
    ])

    add_h1(doc, "五、模型评估结果")
    add_h2(doc, "5.1 AUC与分类指标")
    result_rows = []
    for _, row in metrics.iterrows():
        result_rows.append([
            row["model_label"],
            f"{row['auc']:.4f}",
            f"[{row['auc_ci_low']:.4f}, {row['auc_ci_high']:.4f}]",
            f"{row['accuracy']:.4f}",
            f"{row['precision']:.4f}",
            f"{row['recall']:.4f}",
            f"{row['f1']:.4f}",
        ])
    add_table(doc, 6, "三种模型在测试集上的表现", ["模型", "AUC", "95%区间", "Accuracy", "Precision", "Recall", "F1"], result_rows, [1200, 950, 1800, 1170, 1170, 1170, 1180])
    best = metrics.sort_values("auc", ascending=False).iloc[0]
    test_positive_rate = (best["tp"] + best["fn"]) / (best["tp"] + best["fn"] + best["tn"] + best["fp"])
    add_body(doc, f"逻辑回归的测试集AUC最高，为{best['auc']:.4f}；随机森林为0.5189，决策树为0.5068。以0.5为随机排序基准，三者分别高2.10、1.89和0.68个百分点。这说明特征中不是完全没有信息，但信息量很弱，很多上涨样本与不涨样本仍会被排反。")
    add_body(doc, f"“略高于随机”不等于“与随机完全没有统计差异”。逻辑回归按交易日聚类重采样的95%区间为[{best['auc_ci_low']:.4f}, {best['auc_ci_high']:.4f}]，下界略高于0.5。但样本量较大时，很小的差异也可能在统计上被检测到。该区间仍紧贴0.5，且相邻5日标签存在时间重叠，所以不应将它解读为强预测能力。")
    add_body(doc, f"再看硬分类结果，逻辑回归的Accuracy为{best['accuracy']:.2%}，测试期上涨样本本身占{test_positive_rate:.2%}。如果事后始终猜测试期的多数类“上涨”，正确率为{test_positive_rate:.2%}，反而略高于逻辑回归。若严格按训练期多数类“不涨”作为事前基线，Accuracy为{baseline['accuracy']:.2%}，逻辑回归也只高约0.56个百分点。这是说“0.5阈值下和随机差别不大”的直接依据。")
    add_figure(doc, CHART_DIR / "roc_curves.png", 2, "三种分类模型的测试集ROC曲线", "图2中三条ROC曲线都贴近随机分类的对角线。逻辑回归略高于另外两个模型，但差距不大。这个结果不支持把模型直接当成可盈利策略。", width=5.6)

    add_h2(doc, "5.2 混淆矩阵")
    add_figure(doc, CHART_DIR / "confusion_matrices.png", 3, "分类阈值为0.5时的测试集混淆矩阵", "图3显示，三个模型都较少预测上涨。逻辑回归识别出3,793个上涨样本，但漏掉19,642个，上涨Recall为0.1619。其预测概率的中位数为0.4754，只有15.74%的样本超过0.5，因此低Recall很大程度上是弱信号、训练期基准率与固定阈值共同作用的结果。0.5阈值只是分类演示，不等于经过交易成本优化的买入阈值。", width=6.0)

    add_h2(doc, "5.3 特征解释")
    add_figure(doc, CHART_DIR / "feature_importance.png", 4, "逻辑回归系数与随机森林特征重要性", "图4左侧显示，逻辑回归对20日成交额比给出较大正系数，对20日成交量比给出较大负系数。两者在经济含义上接近且相关，系数一正一负提醒我们：单个系数可能受共线性影响，不宜独立解读为稳定交易规律。右侧随机森林较依赖20日均线偏离、10日收益、60日均线偏离和RSI14，说明模型主要在价格趋势变量中切分样本。不纯度重要性只表示模型使用程度，不表示因果作用。", width=6.0)

    add_h2(doc, "5.4 为什么与随机差别不大")
    add_body(doc, "第一，5日股价方向的噪声很大。公司公告、宏观消息、行业变化与市场情绪都可能在观察日之后出现，而模型只看历史价量指标。一组信息有限的特征面对高噪声标签，可以学到的稳定关系本来就很小。")
    add_body(doc, "第二，15个变量大多是同一份历史价量数据的不同变换，增加指标数量不等于增加了同等数量的新信息。逻辑回归、决策树和随机森林虽然学习方式不同，但都受到相同的信息上限约束。")
    add_body(doc, "第三，训练期、验证期和测试期的市场分布在变化。三段上涨比例分别为48.35%、44.97%和50.37%，测试期内也从2024年的48.62%变为2025年的52.17%。这种基准率与市场状态的变化会使训练期学到的阈值关系衰减。复杂模型更容易学到某一阶段特有的细小结构，这可以解释随机森林在验证集略好，测试集却没有超过逻辑回归。")
    logit_yearly = yearly[yearly["model"] == "logistic_regression"].set_index("year")
    forest_yearly = yearly[yearly["model"] == "random_forest"].set_index("year")
    add_body(doc, f"从年度结果看，逻辑回归在2024和2025年的AUC分别为{logit_yearly.loc[2024, 'auc']:.4f}和{logit_yearly.loc[2025, 'auc']:.4f}，随机森林分别为{forest_yearly.loc[2024, 'auc']:.4f}和{forest_yearly.loc[2025, 'auc']:.4f}。弱但连续两年略高于0.5，比只在某一年突然出现高分更值得关注；但这种稳定性仍然只能支持“弱排序信号”的判断。")
    add_body(doc, "AUC只衡量涨跌排序，不衡量收益幅度，也没有计入手续费、滑点、停牌和涨跌停。因此本次学习的正确结论是：建模流程已经完成，并且捕捉到了很弱的样本外排序信息；目前证据还不支持将它直接视为可盈利的交易引擎。")

    add_h1(doc, "六、总结")
    add_h2(doc, "6.1 任务完成情况")
    add_body(doc, "本报告解释了逻辑回归、决策树和随机森林，也说明了混淆矩阵、ROC和AUC的含义。Python部分使用沪深A股API数据完成标签、特征、时间划分、模型训练、AUC计算和ROC绘制。独立校验脚本共检查18项数据与指标，全部通过。")

    add_h2(doc, "6.2 结果的限制")
    add_body(doc, "股票池按2018年初成交活跃度冻结，不等同于完整A股或历史指数成分股。模型只使用价格和成交量衍生指标，没有加入行业、估值和财务信息。5日标签在相邻日期高度重叠，0.5阈值也没有经过交易成本约束。上述限制会影响结果的外推范围。")

    add_h2(doc, "6.3 交互式Dashboard")
    add_body(doc, "项目同时生成了Task5/dashboard/index.html。页面可以切换模型，查看AUC、混淆矩阵、ROC曲线和特征排序，打开本地HTML即可运行。Dashboard用于作品集展示，不影响课程报告的计算结果。")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Saved {OUTPUT}")


if __name__ == "__main__":
    main()
