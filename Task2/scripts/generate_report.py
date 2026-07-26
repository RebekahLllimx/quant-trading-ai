#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Phase 6: 生成 Task2 .docx 报告
格式: 宋体五号(10.5pt), 1.5倍行距, 0段间距, 两端对齐
"""

import os
import sys
import pandas as pd
import numpy as np
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '..'))
from src.indicators import (
    calc_all_indicators, calc_rsi, calc_macd, calc_bollinger,
    calc_atr, calc_kdj, calc_ma, calc_cci, calc_adx,
    data_diagnosis,
)
from src.report_utils import (
    set_cjk_font, add_paragraph, add_heading_styled,
    add_picture_captioned, add_table, FONT_NAME, FONT_SIZE,
)
from src.equation_utils import (
    add_equation, rsi_formula, macd_formulas, bollinger_formulas,
    atr_formula, kdj_formulas, ma_formulas, cci_formula, adx_formulas,
)

# ═══════════════════════════════════════════════════════════════
# 路径
# ═══════════════════════════════════════════════════════════════
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DATA_DIR = os.path.join(BASE_DIR, '..', 'data', 'csv')
CHART_DIR = os.path.join(BASE_DIR, '..', 'artifacts', 'charts', 'task2')
OUTPUT_FILE = os.path.join(BASE_DIR, '..', 'Task2', 'Rebecca+Task2.docx')

# ═══════════════════════════════════════════════════════════════
# 内容章节
# ═══════════════════════════════════════════════════════════════

def write_data_diagnosis(doc, df):
    """第一章: 数据诊断分析"""
    add_heading_styled(doc, '一、数据诊断分析', level=1)

    diag = data_diagnosis(df, '贵州茅台 (600519)')

    # 1.1 数据概况
    add_heading_styled(doc, '1.1 数据概况', level=2)
    add_paragraph(doc, f'本报告以贵州茅台（600519.SH）为主要分析标的，数据来源于AKShare前复权（qfq）日线数据。数据时间范围为{diag["date_start"]}至{diag["date_end"]}，共计{diag["rows"]}个交易日，时间跨度约{diag["date_range_days"]}个自然日。',
                 first_line_indent=Cm(0.74))

    # 1.2 缺失值检查
    add_heading_styled(doc, '1.2 缺失值检查', level=2)
    add_paragraph(doc, f'对开盘价（Open）、最高价（High）、最低价（Low）、收盘价（Close）、成交量（Volume）五个核心字段进行缺失值检查。检查结果表明，{diag["rows"]}条记录中不存在缺失值，数据完整性良好。')

    # 1.3 描述性统计
    add_heading_styled(doc, '1.3 描述性统计分析', level=2)

    cs = diag['close_stats']
    rs = diag['return_stats']

    add_paragraph(doc, f'收盘价的描述性统计如表1所示。贵州茅台在过去一年间（约400个自然日、267个交易日）的收盘价均值为{cs["mean"]:.2f}元，标准差为{cs["std"]:.2f}元。价格区间为[{cs["min"]:.2f}, {cs["max"]:.2f}]元。',
                 first_line_indent=Cm(0.74))

    add_table(doc,
              ['统计量', '数值', '含义'],
              [
                  ['均值 (Mean)', f'{cs["mean"]:.2f} 元', '过去267个交易日收盘价的算术平均值'],
                  ['标准差 (Std)', f'{cs["std"]:.2f} 元', '每日收盘价偏离均值的程度'],
                  ['最小值 (Min)', f'{cs["min"]:.2f} 元', '区间内最低收盘价'],
                  ['最大值 (Max)', f'{cs["max"]:.2f} 元', '区间内最高收盘价'],
                  ['偏度 (Skewness)', f'{cs["skewness"]:.4f}',
                   f'{"右偏（正偏），右侧尾部更长" if cs["skewness"]>0 else "左偏（负偏），左侧尾部更长"}'],
                  ['峰度 (Kurtosis)', f'{cs["kurtosis"]:.4f}',
                   '大于0表示分布比正态分布更尖峭，具有"尖峰厚尾"特征'],
              ])

    add_paragraph(doc, f'日收益率统计如表2所示。日均收益率为{rs["daily_mean_pct"]:.4f}%，日波动率为{rs["daily_std_pct"]:.4f}%，年化波动率为{rs["annual_vol_pct"]:.2f}%（基于252个交易日年化）。最大回撤为{diag["max_drawdown_pct"]:.2f}%。',
                 first_line_indent=Cm(0.74))

    add_table(doc,
              ['指标', '数值', '说明'],
              [
                  ['日均收益率', f'{rs["daily_mean_pct"]:.4f}%', '日收益率的算术平均值'],
                  ['日波动率', f'{rs["daily_std_pct"]:.4f}%', '日收益率的标准差'],
                  ['年化波动率', f'{rs["annual_vol_pct"]:.2f}%', '日波动率 × √252'],
                  ['最大回撤', f'{diag["max_drawdown_pct"]:.2f}%', '区间内从最高点到后续最低点的最大跌幅'],
                  ['异常值数量(3σ)', f'{diag["outliers_3sigma"]}', '收盘价超出均值±3倍标准差的天数'],
              ])

    # 1.4 正态性检验
    add_heading_styled(doc, '1.4 收益率分布与正态性检验', level=2)

    jb = diag['jarque_bera']
    add_paragraph(doc,
        f'Jarque-Bera正态性检验结果为：JB统计量 = {jb["statistic"]:.4f}，p值 = {jb["p_value"]:.6f}。'
        f'p值远小于显著性水平0.05，拒绝收益率服从正态分布的原假设。结合偏度（{cs["skewness"]:.4f}）和峰度（{cs["kurtosis"]:.4f}）的数值，'
        f'贵州茅台的日收益率分布呈现"左偏、尖峰厚尾"的特征，与多数金融时间序列的典型特征一致。'
        f'这一发现提示：基于正态分布假设的风险度量（如VaR）可能低估极端风险。',
        first_line_indent=Cm(0.74))

    doc.add_page_break()


def write_indicator_section(doc, df, stock_info):
    """第二章: 技术指标分析 (8个指标)"""
    add_heading_styled(doc, '二、技术指标分析', level=1)

    add_paragraph(doc,
        '本章对八个技术指标逐一进行分析。每个指标的论述包括四个部分：'
        '（1）历史演化——指标的发明背景、原始动机和后续发展；'
        '（2）计算公式——精确的数学定义；'
        '（3）参数讨论——参数的来源和不同取值的影响；'
        '（4）信号含义——指标读数所传达的市场信息。'
        '分析标的为贵州茅台（600519.SH），图表基于2025年5月至2026年7月的日线数据（前复权）绘制。',
        first_line_indent=Cm(0.74))

    sections = [
        ('RSI', '相对强弱指数', '2.1', 'RSI是由J. Welles Wilder Jr.于1978年在《New Concepts in Technical Trading Systems》一书中提出的动量振荡指标。Wilder曾是一名机械工程师，他的核心动机是构建一个取值固定在[0, 100]的有界振荡器——在RSI之前，常见的动量指标（如简单动量=今日收盘−N日前收盘）是无界的，无法在不同品种和不同时间段之间比较读数的高低。RSI通过将涨跌幅比例归一化到[0, 100]区间，解决了动量度量的可比性问题。\n\nWilder选择参数N=14并非统计优化的结果，而是基于他对期货市场交易周期的观察：14个交易日约等于3个自然周，覆盖一个短周期波动的完整起落。他使用Wilder平滑（α=1/N而非EMA的α=2/(N+1)）是因为他希望RSI过滤短期噪音而非追逐每一个微小的价格变化。\n\n1990年代，Andrew Cardwell等人发展了RSI的背离分析方法，将RSI从单一的超买超卖判断扩展为趋势强度确认工具。Cardwell发现，在牛市中RSI倾向于在40–80区间运行（而非30–70），在熊市中倾向于在20–60区间运行——这一观察修正了固定30/70阈值的过度简化。'),
        ('MACD', '指数平滑异同移动平均线', '2.2', 'MACD由Gerald Appel于1979年提出，发表于《Technical Analysis of Stocks & Commodities》杂志。Appel是一位管理数亿美元资金的投资顾问，他的设计目标是将趋势方向判断（跟踪型指标的功能）和趋势力度度量（动量型指标的功能）融合到一个指标体系中。\n\nAppel选择EMA而非SMA，是因为EMA对近期价格赋予更高权重（α=2/(N+1)），能更快反映趋势的变化。参数12（约2.5个自然周）、26（约1个自然月）和9（约2个自然周）的选取映射到自然时间单位，反映了短期交易成本与中期交易成本的关系。\n\n1986年，Thomas Aspray发现MACD柱状图（BAR=DIF−DEA）比金叉死叉更早反映趋势变化——柱状图的缩短通常在交叉发生前1–3根K线就已开始。这一发现使交易者可以在金叉死叉之前获得趋势动能衰减的预警。'),
        ('布林带', 'Bollinger Bands', '2.3', '布林带由John Bollinger在1980年代早期提出，最初发表于1992年的《Stocks & Commodities》杂志。Bollinger曾任职于Financial News Network（后并入CNBC），负责技术分析内容的制作。\n\nBollinger的核心创新是：用标准差（σ）取代传统通道指标中使用的固定百分比。固定百分比通道的基本缺陷是假设市场波动率恒定——实际上金融市场的波动率随时间变化剧烈。在市场波动率低时固定通道过宽，价格永远触不到边界；波动率高时通道过窄，价格频繁突破导致信号失效。标准差是一个天然的时变波动率度量——波动大时σ增大、通道自动变宽，波动小时σ减小、通道自动收窄，使得布林带成为一个自适应波动率的通道系统。\n\nBollinger选择参数N=20（约一个自然月交易日）和k=2（正态分布下覆盖约95.4%的样本），但他本人强调正态假设只是近似——实际金融数据具有肥尾特征，价格触及轨道的概率高于5%。他后续引入了%b（价格在带内的相对位置）和BandWidth（带宽百分比）两个辅助指标，用于量化轨道挤压（Squeeze）和价格相对定位。'),
        ('ATR', '平均真实波幅', '2.4', 'ATR由J. Welles Wilder Jr.于1978年与RSI一同在《New Concepts in Technical Trading Systems》中提出。Wilder最初设计ATR是为商品期货市场服务——商品期货相比股票更容易出现隔夜跳空缺口（因天气、地缘、供需消息等）。\n\n传统波动度量方式是计算当日最高价与最低价的差值（即日内振幅）。但日内振幅无法捕捉跨交易日的价格跳跃——如果今日因隔夜利好消息直接高开，日内振幅可能很小，但实际的价格运动（从昨日收盘到今天高位）是巨大的。Wilder定义了真实波幅（True Range）：取"日内振幅""今高与昨收的缺口""今低与昨收的缺口"三者中的最大值，确保任何形式的跨日价格运动都被纳入波动度量。\n\nWilder对ATR使用1/N系数的Wilder平滑（而非EMA的2/(N+1)），因为他希望ATR反映市场波动率的背景水平而非追踪每日波动的瞬时变化。1998年，Van K. Tharp在《Trade Your Way to Financial Freedom》中将ATR引入仓位管理系统——用ATR标准化每笔交易的风险暴露，使ATR从单纯的波动度量工具扩展为风险管理的核心输入。'),
        ('KDJ', '随机指标', '2.5', '随机指标的原始形式（%K和%D）由George C. Lane在1950年代提出。Lane是一名期货交易员和投资教育家，他通过观察发现：在上涨趋势中，收盘价倾向于收在当日价格区间的上端（接近最高价）；在下跌趋势中，收盘价倾向于收在当日价格区间的下端（接近最低价）；在趋势即将反转时，收盘价的位置会先于价格本身开始偏移——即使绝对价格还在创新高，收盘价相对于当日区间的位置可能已经开始下移。\n\nLane将这一观察形式化为随机值RSV（Raw Stochastic Value）：收盘价在过去N日最高价和最低价之间的百分位位置。RSV=80意味着收盘价处于区间的上80%位置，RSV=20意味着收盘价处于区间的下20%位置。\n\nKDJ在传统KD指标（%K=RSV的EMA平滑，%D=%K的EMA平滑）基础上增加了一条J线（J=3K−2D）。J线的引入主要来自中国A股市场技术分析软件的实践（如通达信、同花顺），在国际市场上并不常见。J线通过数学放大K与D的差异，比K线和D线更早、更剧烈地反映方向变化。J>100或J<0的状态在数学上是不稳定的，因为J对K的放大效应在K的增速放缓时会迅速衰减。'),
        ('MA', '移动平均线', '2.6', '移动平均的概念可追溯到19世纪末的统计学领域，但它在金融市场的系统应用始于20世纪初。最著名的早期实践者是Richard Donchian——他在1930年代使用5日和20日移动平均线的交叉来生成商品期货交易信号，被公认为"趋势跟踪之父"。Donchian的核心理念是：让价格本身通过移动平均线告知趋势何时发生变化，而非依赖主观预测。\n\nSMA（简单移动平均）是最早的形式——每期权重相等。1960年代，P.N. Haurlan将指数平滑引入股票市场分析，而后EMA（指数移动平均）因对近期数据赋予更高权重而逐渐流行。EMA解决了SMA的两个数学问题：(1) SMA的"幽灵效应"——当旧数据被移出窗口时，SMA产生一次性的跳跃（即使当前价格不变）；(2) SMA对所有N期内数据等同对待，忽略了信息的时效衰减——昨天的价格和N天前的价格被赋予相同的影响力。\n\n1970年代，Robert Levy和Jack Schwager发展了多周期均线分析框架——用不同时间尺度的MA同时作用，通过观察均线的排列关系（多头排列：短期MA全部在长期MA之上）来判断趋势的强度和多周期一致性。这一分析方法至今仍是趋势跟踪策略的基础。'),
        ('CCI', '商品通道指数', '2.7', 'CCI由Donald Lambert于1980年在《Stocks & Commodities》杂志上发表的《Commodity Channel Index: Tool for Trading Cyclic Trends》一文中提出。Lambert是期货交易员和技术分析师，CCI最初是为商品期货市场设计的——商品价格具有较强的周期性特征（如季节性的供需周期），Lambert希望量化价格在其周期性波动中的"异常偏离程度"。\n\nLambert在度量偏离时选择了平均绝对偏差（Mean Deviation）而非标准差。他的理由是：标准差的平方运算会导致极端值被不成比例地放大——一根极端K线足以显著拉高标准差，进而使CCI的偏离度量失真。平均绝对偏差使用绝对值而非平方，对极端值的敏感度较低，能更稳健地反映价格的"正常偏离程度"。\n\nCCI公式中的常数0.015是Lambert通过历史回测选择的，目的是让约70–80%的CCI读数落在[−100, +100]区间内。这一常数没有统计学上的必然性——它是经验校准的结果。CCI虽名为"商品通道指数"，但其度量统计异常偏离的特性使得它在股票市场上同样适用。CCI的最大特点是**无界**——与RSI的[0, 100]不同，CCI理论值可以趋向任意大的正值或负值，因此在强趋势中不会像RSI那样钝化在极值区域。'),
        ('ADX', '平均趋向指数', '2.8', 'ADX由J. Welles Wilder Jr.于1978年与RSI和ATR一同提出，是Wilder"趋向运动系统"（Directional Movement System, DMS）的核心组成部分。Wilder可能是第一个将"趋势是否存在"作为一个可量化问题来研究的人。\n\nADX最大的概念创新在于：它将"方向"（价格上涨还是下跌）和"趋势强度"（无论方向如何，价格是否在朝一个方向持续运动）进行了彻底的分离。+DI（正向趋向指标）度量向上的方向运动力度，−DI（负向趋向指标）度量向下的方向运动力度，而ADX度量的是趋向运动的总体强度——不区分向上还是向下。\n\n这一分离意味着ADX可以回答一个此前未被系统回答的问题："当前市场是否适合使用趋势跟踪策略？"——当ADX<20时，无论±DI如何交叉，趋势跟踪策略的统计表现都会下降，因为市场处于无方向的震荡状态。ADX的计算涉及多个步骤：先计算真实波幅（ATR）和趋向运动（±DM），再计算趋向指标（±DI=Wilder(±DM, N)/ATR×100），然后求方向运动指数（DX=|+DI−-DI|/(+DI+-DI)×100），最后通过Wilder平滑得到ADX。多层平滑使得ADX成为一个延迟确认指标——这既是其可靠性来源（不容易被单日噪音误导），也是其局限（趋势可能已被确认时已运行了一段距离）。'),
    ]

    for key, name_cn, num, desc in sections:
        # 提取指标简短标题 (如 RSI, MACD)
        short_key = key.upper()

        add_heading_styled(doc, f'{num} {name_cn}（{key}）', level=2)

        # (a) 历史演化
        add_heading_styled(doc, f'{num}.1 历史演化', level=3)
        for para_text in desc.split('\n\n'):
            if para_text.strip():
                add_paragraph(doc, para_text.strip(), first_line_indent=Cm(0.74))

        # (b) 计算公式 (OMML 专业公式)
        add_heading_styled(doc, f'{num}.2 计算公式与参数', level=3)

        # 参数说明 (简短文字)
        param_notes = {
            'RSI': '使用 Wilder 平滑（α = 1/N），初始种子值取前 N 期简单移动平均。默认参数：N = 14。',
            'MACD': 'EMA 使用 α = 2/(N+1) 的指数平滑系数。默认参数：fast = 12, slow = 26, signal = 9。',
            '布林带': '标准差 σ 用于自适应波动率。默认参数：period = 20, k = 2.0。',
            'ATR': 'Wilder 平滑系数 α = 1/N。默认参数：N = 14。',
            'KDJ': 'RSV 为收盘价在 N 日区间内的百分位位置。默认参数：period = 9, smooth_k = 3, smooth_d = 3。',
            'MA': '常用周期：5日（1周）、10日（2周）、20日（1月）、60日（1季度）、120日（半年）、250日（年线）。',
            'CCI': '常数 0.015 为 Lambert 经验校准值，使 70–80% 读数落在 [−100, +100]。默认参数：N = 20。',
            'ADX': 'Wilder 平滑用于 +DI/−DI/DX/ADX 各层。默认参数：N = 14。',
        }

        # 插入 OMML 公式
        formula_funcs = {
            'RSI': rsi_formula,
            'MACD': macd_formulas,
            '布林带': bollinger_formulas,
            'ATR': atr_formula,
            'KDJ': kdj_formulas,
            'MA': ma_formulas,
            'CCI': cci_formula,
            'ADX': adx_formulas,
        }

        if key in formula_funcs:
            for line in formula_funcs[key]():
                add_equation(doc, line)

        if key in param_notes:
            add_paragraph(doc, param_notes[key], first_line_indent=Cm(0.74))

        # (c) 信号含义
        add_heading_styled(doc, f'{num}.3 信号含义', level=3)
        signals = {
            'RSI': 'RSI > 70：过去N日内涨幅的Wilder平滑值达到跌幅平滑值的2.33倍以上（RS > 2.33），近期价格上涨的持续性显著高于下跌。RSI < 30：跌幅平滑值达到涨幅平滑值的2.33倍以上，近期下跌的持续性显著高于上涨。RSI从下方突破50：涨幅平滑值重新超过跌幅平滑值，涨跌力量对比从跌方占优转为涨方占优。RSI从上方跌破50：跌幅平滑值重新超过涨幅平滑值，涨跌力量对比从涨方占优转为跌方占优。\n\n顶背离（价格创新高，RSI未创新高）：最新一波上涨中，涨幅与跌幅的比值（RS）未能超过前一波上涨时的水平。尽管价格更高，但内部涨跌力量比在衰减——推动新高的上涨力度弱于前一次。底背离（价格创新低，RSI未创新低）：同理，打压新低的下跌力度弱于前一次。\n\nRSI在强单边趋势中可能出现"钝化"——长期停留在超买区（>70）或超卖区（<30）而不回落。这不是反转信号，而是趋势极强的确认。RSI在震荡市中最有效，在强趋势市中使用区间信号时需要结合趋势强度指标（如ADX）共同判断。',
            'MACD': 'DIF上穿DEA（金叉）：DIF正以快于其近期平均速率的速度增长，短期EMA与长期EMA的间距在加速扩大。DIF下穿DEA（死叉）：DIF正以快于其近期平均速率的速度减小，短期EMA与长期EMA的间距在加速收缩。\n\nDIF > 0（零轴上方）：EMA(Close, 12) > EMA(Close, 26)，近12日加权平均成本高于近26日加权平均成本。当前市场价格高于约一个月来的平均成本中枢。DIF < 0（零轴下方）：近12日加权平均成本低于近26日加权平均成本。\n\nBAR绝对值连续增大：|DIF−DEA|在扩大，趋势的瞬时加速度在增大。BAR绝对值连续减小：趋势的瞬时加速度在减小——这通常比DIF与DEA的交叉提前1–3根K线出现，是动量衰减的早期信号。\n\nMACD在趋势明确的市场中最有效。在震荡市中，快慢EMA反复交叉产生连续的假金叉和假死叉——称为"拉锯效应"（whipsaw）。此时需要结合ADX或其他趋势确认指标来过滤信号。',
            '布林带': '价格突破上轨（%b > 1）：收盘价超过MID + 2σ。在当前N日样本中，这个价位与均值的偏差超过了2个标准差，属于统计意义的极端值。价格跌破下轨（%b < 0）：收盘价低于MID − 2σ，处于统计低位。\n\n带宽收窄（Squeeze）：标准差σ减小导致上下轨间距收缩。低波动率环境不会无限持续——历史数据显示波动率存在聚集效应（volatility clustering），极端低波动后通常跟随波动率的均值回归（放大）。布林带将这种状态称为Squeeze——带宽收缩至极低值后，往往伴随波动率的突然放大和方向性行情的出现。\n\n价格连续≥5根K线沿上轨运行（%b持续>0.8）：价格持续在MID+1.6σ以上的高位运行而没有回归中轨。这不是一次性的统计异常值，而是趋势在持续向上推进——同时中轨自身也在跟随上移。中轨斜率由平转正：SMA(Close, 20)的方向从横向转为向上，近期价格的重心开始上移。\n\n布林带的核心假设（价格在一定时间尺度上围绕均值波动）在非趋势市场中成立得较好。在强趋势市场中，价格可以持续沿轨道运行而不回归中轨——此时的正确解读是"价格沿轨道持续运行意味着趋势强劲"而非"突破轨道意味着即将反转"。',
            'ATR': 'ATR上升：近N日的平均真实波幅在增大。无论价格上涨还是下跌，每根K线的价格运动幅度在扩大。市场参与者的交易活跃度或分歧程度在上升。ATR下降：近N日的平均真实波幅在减小，K线运动幅度在缩小。\n\nATR突然跳升（当前ATR/过去20日均值 > 1.5）：真实波幅出现了超出近期正常范围的放大。通常与重大消息、财报发布、政策变动等事件有关。ATR持续低位（连续多日低于长期均值一半）：市场进入低波动环境——由于波动率有均值回归倾向，持续低ATR往往预示着后续波动放大。\n\nATR/Close（归一化ATR）：将ATR除以收盘价得到百分比形式，使不同价格水平的品种的波动烈度可以直接比较。例如ATR=5元、Close=100元的品种（ATR%=5%）比ATR=5元、Close=500元的品种（ATR%=1%）的波动烈度高5倍。\n\nATR是八个指标中唯一完全不做方向判断的指标——它只回答"市场波动有多大"而不回答"市场在涨还是跌"。ATR的核心价值在于风险管理和仓位调整中的波动率标准化。',
            'KDJ': 'K > 80且D > 80：RSV经过平滑后持续处于高位，收盘价在近期区间内持续接近最高价。K < 20且D < 20：收盘价在近期区间内持续接近最低价。\n\nK上穿D（低位金叉）：RSV的短期平滑线向上突破其自身的长期平滑线，收盘价在区间内的位置正在上移。如果发生在K/D<20的极低区域，说明这一上移是从区间极低位置开始的。K下穿D（高位死叉）：RSV的短期平滑线正在加速下降。\n\nJ > 100：K的偏移量2×(K−D)与K自身之和超过了100。K向上偏离D的程度较大，且两者的方向均为向上。J>100在数学上是不稳定的——J对K的放大效应在K增速放缓时会迅速衰减。同理J<0也是数学上不稳定的状态。\n\nKDJ的高敏感度使其在震荡市和短期波段中较为有效。在强趋势中，KDJ会"钝化"——K、D、J在80以上或20以下长期徘徊且反复缠绕，无法通过极值回归来区分"趋势延续"和"趋势反转"。此外，KDJ的区间依赖特性意味着：如果区间因单根长K线而突然扩大，后续数日的RSV值都会受到该极值的"锚定效应"影响。',
            'MA': '多头排列（MA5 > MA10 > MA20 > MA60）：各时间尺度的平均持仓成本从短到长依次升高——所有时间框架的交易者（短期到长期）在当前价格下均处于盈利状态。不同时间尺度的市场参与者在此价格水平上形成了一致的方向共识。空头排列（MA5 < MA10 < MA20 < MA60）：各时间尺度的平均持仓成本从短到长依次降低，市场整体的价格共识向下。\n\n短期MA上穿长期MA：近期参与者的持仓成本超过了较早参与者的成本——意味着市场愿意以更高的价格入场，短期方向共识从空转多。短期MA下穿长期MA：短期方向共识从多转空。\n\n多周期MA间距收敛（粘合）：不同时间尺度的平均持仓成本趋于一致。各时间框架的市场参与者对当前价格的共识高度集中。当多种时间尺度的参与者成本几乎相同时，任何方向的价格突破都会导致部分参与者开始盈利、另一部分开始亏损——这种结构的变化往往伴随趋势的启动。多周期MA间距发散：趋势正在加速。\n\nMA的核心假设是"价格趋势具有持续性"——趋势一旦形成倾向于延续。这一假设在趋势市中成立，但在震荡市中均线频繁缠绕交叉。年线（MA250）被广泛视为牛熊分界线——价格在年线上方运行时，过去一年内入场的投资者平均处于盈利状态。',
            'CCI': 'CCI > +100：典型价格（TP）高于其20期均值超过1.5倍的平均绝对偏差——当前TP偏离其近期均值的幅度已超过正常波动范围，处于统计强势区域。CCI < −100：TP处于统计弱势区域。\n\nCCI > +200：TP偏离均值超过3倍的平均绝对偏差。在当前波动率水平下，这种幅度的正向偏离属于罕见事件。CCI < −200：同等程度的极端负偏离。\n\nCCI从下方突破+100：TP从正常波动范围进入统计强势区域。CCI从上方跌破+100：TP从强势区域回归正常波动范围，统计显著性降低。\n\nCCI的最大特征是无界——与RSI[0,100]不同，CCI理论上可以趋向任意大的正值或负值。这使得CCI在强趋势中不会像RSI那样钝化在极值区域：当RSI已经在85附近失去辨别力时，CCI可能已经到了+300，继续度量趋势的"极端程度"。但CCI的无界性也意味着它不提供固定的超买超卖阈值——不同品种的CCI的正常波动范围可能不同。',
            'ADX': 'ADX > 25：过去14日的趋向运动强度经过平滑后处于中等偏上水平。市场存在可识别的方向性运动——价格在朝某个方向持续运动，而非在区间内来回摆动。ADX > 40：趋向运动强度处于高位，市场正在经历强趋势。ADX < 20：趋向运动强度不足，±DI的差值相对于它们的总和较小，多空方向力量交替占优。市场缺乏持续的方向性驱动力。\n\nADX从<20向上突破25：趋向运动强度从低位进入中等水平。市场的方向性从"不可识别"变为"可识别"——某种方向的力量正在从震荡中浮现并建立持续性。ADX从高位（>40）回落：趋势强度正在减弱。这不直接等于趋势反转——可能是趋势暂停（横盘整理后继续原方向）或趋势逆转的前奏。\n\n+DI > −DI且ADX > 25：在可识别的趋势中，向上力量高于向下力量。−DI > +DI且ADX > 25：向下力量高于向上力量。+DI上穿−DI：方向的主导权从空方转移到多方。如果此时ADX较低（<20），这一交叉可能短暂且不可靠。\n\nADX是一个延迟确认指标——当前的ADX值反映的是过去N期平均的趋势强度，当ADX确认趋势存在时，趋势可能已运行了一段时间。它解决的核心问题是"趋势是否存在"，最适用于帮助判断当前市场更适合趋势策略还是震荡策略。',
        }
        if key in signals:
            add_paragraph(doc, signals[key], first_line_indent=Cm(0.74))

    doc.add_page_break()


def write_synergy(doc):
    """第三章: 指标互补关系"""
    add_heading_styled(doc, '三、指标互补关系分析', level=1)

    add_paragraph(doc,
        '八个技术指标从不同的数据源和数学变换出发，度量市场的不同维度。单独使用任何一个指标都存在局限性——每个指标有其特定的适用条件和失效场景。将多个指标按观察维度进行分组，有助于理解指标之间的信息重叠和互补关系。',
        first_line_indent=Cm(0.74))

    add_heading_styled(doc, '3.1 指标维度分类', level=2)

    add_table(doc,
              ['观察维度', '回答的问题', '核心指标', '辅助指标'],
              [
                  ['趋势方向', '价格重心在向哪个方向移动？', 'MA排列、MACD(DIF正负)', '布林带(中轨斜率)'],
                  ['趋势强度', '方向性运动有多强？', 'ADX', 'MACD(BAR柱大小)'],
                  ['动量/速度', '涨跌的力有多大？', 'RSI(涨跌力量比)', 'KDJ(收盘价位置变化速度)'],
                  ['波动烈度', '价格波动有多大？', 'ATR', '布林带(带宽)'],
                  ['极端位置', '价格是否处于统计边缘？', '布林带(%b, 轨道)', 'CCI(±100/200)'],
                  ['短期拐点', '动量即将减速还是加速？', 'KDJ(K/D/J交叉)', 'MACD(BAR缩短)、CCI(±100穿越)'],
              ])

    add_heading_styled(doc, '3.2 不同市场状态下的有效指标组合', level=2)

    add_paragraph(doc,
        '震荡市（ADX<20，布林带宽窄）：价格在一定区间内来回运动，无单边趋势。此时最有效的信息维度是极端位置和短期动量拐点——识别价格何时接近区间边界以及动量何时开始向反方向偏移。RSI的30/70区域信号、KDJ的20/80交叉信号、布林带的轨道触碰信号在此类市场中提供的信息量最大。MACD在震荡市中的金叉死叉假信号率高，因为快慢EMA反复小幅交叉。',
        first_line_indent=Cm(0.74))

    add_paragraph(doc,
        '趋势市（ADX>25，布林带宽扩）：价格朝一个方向持续运动。此时趋势方向和强度的度量最为关键，而超买超卖类指标（RSI、KDJ）会发生钝化，极端位置信号需要区分处理。MA排列（多头/空头）和斜率、MACD的DIF正负和BAR方向、ADX的强度读数、布林带中轨斜率在此类市场中提供的信息量最大。ATR的升降可用于判断趋势是否伴随真实波动——上涨+ATR放大意味着趋势有资金参与支撑。',
        first_line_indent=Cm(0.74))

    add_heading_styled(doc, '3.3 指标互补矩阵', level=2)

    add_table(doc,
              ['市场状态', '主指标', '辅助确认', '指标关系'],
              [
                  ['ADX<20 震荡', 'RSI(14) + KDJ(9,3,3)', '布林带(20,2)上下轨', '超卖/超买共振确认'],
                  ['ADX>25 上升趋势', 'MACD(12,26,9) + MA多头排列', 'ATR(14)度量波动烈度', '方向+力度+波动率同步'],
                  ['ADX>25 下降趋势', 'MACD死叉 + MA空头排列', 'ATR(14)度量波动烈度', '方向+力度+波动率同步'],
                  ['布林Squeeze', '布林带(20,2) + ADX(14)', 'CCI(20)突破确认', '低波动→突破→趋势启动'],
              ])

    doc.add_page_break()


def write_param_evolution(doc, df, stock_info):
    """第四章: 参数演化对比分析 (含图表)"""
    add_heading_styled(doc, '四、参数演化与图表分析', level=1)

    add_paragraph(doc,
        '本章以贵州茅台（600519.SH）的日线数据为基础，通过图表展示各指标在不同参数配置下的行为差异，验证前文所述的参数演化逻辑。以下图表均使用mplfinance专业金融绘图库生成。',
        first_line_indent=Cm(0.74))

    # 获取当前图表编号
    add_heading_styled(doc, '4.1 综合指标分析', level=2)
    add_picture_captioned(doc,
        os.path.join(CHART_DIR, '600519_综合指标分析.png'),
        '图1：贵州茅台（600519）技术指标综合分析——K线（布林带）+ 成交量 + RSI + MACD + KDJ （日线，前复权，2025.05–2026.07）',
        width_inches=6.0)
    add_paragraph(doc,
        '图1显示了贵州茅台过去一年的日线走势及五个核心指标的同步变化。贵州茅台在过去一年中经历了大幅波动：'
        '2025年6月至7月价格从1500元附近大幅下跌，在2025年7月下旬创下1168.63元的区间最低点，随后进入约半年的震荡整理区间（约1200–1300元），'
        '2026年1月下旬价格突破上涨至1526.98元区间最高点，此后再次回落至1200元附近。'
        '各指标在不同阶段反映了不同的市场状态——RSI和KDJ在震荡区间（2025年9–12月）多次产生极值信号，'
        'MACD在2026年1月的上涨行情中产生了明确的趋势信号（DIF和DEA上穿零轴并维持高位），'
        '布林带的带宽在各阶段呈现明显的收缩-扩张周期。',
        first_line_indent=Cm(0.74))

    add_heading_styled(doc, '4.2 补充指标分析', level=2)
    add_picture_captioned(doc,
        os.path.join(CHART_DIR, '600519_补充指标_CCI_ATR_ADX.png'),
        '图2：贵州茅台（600519）补充技术指标——CCI + ATR + ADX/+DI/−DI（日线，前复权，2025.05–2026.07）',
        width_inches=6.0)
    add_paragraph(doc,
        '图2展示了CCI、ATR和ADX三个指标的变化。CCI在2025年7月的下跌行情中深度下探至−200以下，反映了统计极端位置的出现；'
        'ATR在2025年7月和2026年2月出现了两次显著的波动率放大——第一次伴随价格暴跌（恐慌性抛售），第二次伴随价格快速反弹后的震荡。'
        'ADX在2025年7–8月和2026年2–3月两次突破25阈值，确认了这两个时间段内趋势的存在。'
        '2025年9–12月期间ADX长期低于20，对应价格的横盘震荡区间，此时趋势跟踪策略的效率下降，'
        'RSI和KDJ等震荡指标更为适用。',
        first_line_indent=Cm(0.74))

    add_heading_styled(doc, '4.3 RSI参数演化对比', level=2)
    add_picture_captioned(doc,
        os.path.join(CHART_DIR, '600519_RSI参数对比.png'),
        '图3：贵州茅台（600519）RSI参数演化——N=7（敏感）、N=14（默认）、N=21（平滑）对比',
        width_inches=6.0)
    add_paragraph(doc,
        '图3对比了RSI在三个不同周期参数下的行为差异：N=7（敏感，红色反应最快）、N=14（默认，蓝色处于中等敏感度）、'
        'N=21（平滑，绿色过滤了大量噪音）。图中棕色背景标注了RSI>70和RSI<30区域，便于观察超买超卖信号。'
        '从图中可以观察到：N越小，RSI触碰30/70线的频率越高（N=7产生了更多的超买超卖信号），但信号的可信度受到噪音干扰；'
        'N越大，RSI曲线越平滑，信号数量减少但每次穿越的可靠性提高。参数N的选择本质上是在"信号频率"和"信号可靠性"之间做权衡。',
        first_line_indent=Cm(0.74))

    add_heading_styled(doc, '4.4 MACD参数演化对比', level=2)
    add_picture_captioned(doc,
        os.path.join(CHART_DIR, '600519_MACD参数对比.png'),
        '图4：贵州茅台（600519）MACD参数演化——(6,13,5)短线、(12,26,9)经典、(21,55,13)长线对比',
        width_inches=6.0)
    add_paragraph(doc,
        '图4展示了MACD在三组不同参数下的表现：(6,13,5)短线参数（上）产生更多、更快的金叉死叉信号，适合捕捉短期趋势转折，但假信号率较高；'
        '(12,26,9)经典参数（中）在信号频率和可靠性之间取得平衡，是日线分析的标准选择；'
        '(21,55,13)长线参数（下）产生的信号数量最少但每次信号的质量更高，适合周线级别和长线趋势确认。'
        '统计显示，短线参数在267个交易日中产生了最多的金叉/死叉次数，长线参数所产生的交叉最少但最接近价格的主要转折点。',
        first_line_indent=Cm(0.74))

    add_heading_styled(doc, '4.5 布林带参数演化对比', level=2)
    add_picture_captioned(doc,
        os.path.join(CHART_DIR, '600519_布林带参数对比.png'),
        '图5：贵州茅台（600519）布林带参数演化——(10,1.5)短线、(20,2.0)经典、(50,2.5)长线对比',
        width_inches=6.0)
    add_paragraph(doc,
        '图5展示了布林带在三组参数下的对比：(10,1.5)短线参数（上）产生窄通道，价格频繁触轨，适合短线波动交易；'
        '(20,2.0)经典参数（中）提供合理的波动范围，价格在统计上约95%的时间应在带内（理论值）；'
        '(50,2.5)长线参数（下）产生宽通道，价格很少触轨——触轨时意味着达到了季度级别的统计极端位置。'
        '参数组的突破上轨/跌破下轨次数统计显示了参数对信号密度的影响：短线参数触轨次数约为经典参数的2倍，'
        '长线参数触轨次数最少。k值（标准差倍数）直接影响信号密度。k越小，'
        '通道越窄，信号密度越高但假信号比例也越高。',
        first_line_indent=Cm(0.74))

    add_heading_styled(doc, '4.6 KDJ参数演化对比', level=2)
    add_picture_captioned(doc,
        os.path.join(CHART_DIR, '600519_KDJ参数对比.png'),
        '图6：贵州茅台（600519）KDJ参数演化——(5,2,2)极短线、(9,3,3)标准、(14,5,5)长线对比',
        width_inches=6.0)
    add_paragraph(doc,
        '图6对比了KDJ在三组参数下的差异：(5,2,2)极短线参数（上）的J线频繁触及>100和<0的极值区域，信号密度最高；'
        '(9,3,3)标准参数（中）是日线分析的标准选择，J线极值的出现频率适中；'
        '(14,5,5)长线参数（下）的J线波动范围收窄，极值出现频率最低。'
        '极短线参数的J线在262个交易日中J>100天数和J<0天数最为频繁，这些信号中有相当比例是噪音驱动的假信号。'
        '在实际使用中需要根据期望的交易频率和市场波动特征选择合适的参数组合。',
        first_line_indent=Cm(0.74))

    add_heading_styled(doc, '4.7 MA多周期共振分析', level=2)
    add_picture_captioned(doc,
        os.path.join(CHART_DIR, '600519_MA多周期共振.png'),
        '图7：贵州茅台（600519）多周期均线共振——MA5/10/20/60/120/250（日线，前复权，2025.05–2026.07）',
        width_inches=6.0)
    add_paragraph(doc,
        '图7显示了六条不同周期的移动平均线：(1) 2025年7–8月，短期均线（MA5、MA10、MA20）全部位于长期均线（MA60、MA120、MA250）'
        '下方，形成空头排列——各时间尺度交易者的持仓成本从上到下递减，市场整体处于亏损状态；'
        '(2) 2025年10–12月，六条均线在1200–1300元区间高度粘合——各时间尺度交易者的持仓成本几乎一致，'
        '随后的价格突破导致了趋势的启动；(3) 2026年2月之后，短期均线再次下穿长期均线，空头排列重新确立。'
        '均线从粘合到发散的过程是趋势启动的典型特征——当多周期交易者的成本趋同时，任何方向性的力量都可能导致集体性的方向选择。',
        first_line_indent=Cm(0.74))

    doc.add_page_break()


def write_dashboard(doc):
    """第五章: 交互看板"""
    add_heading_styled(doc, '五、交互式技术指标看板', level=1)

    add_heading_styled(doc, '5.1 看板概述', level=2)
    add_paragraph(doc,
        '为直观展示和分析技术指标，本任务构建了一个自包含的交互式Web看板'
        '（Task2/dashboard/index.html）。看板覆盖10只标的（5只A股+5只港股），'
        '集成K线图（含布林带+MA均线叠加）、成交量图和8个技术指标的独立图表，'
        '所有图表通过ECharts 5.5渲染，CSV数据序列化为JSON嵌入页面，无需后端服务器即可在浏览器中运行。',
        first_line_indent=Cm(0.74))

    add_heading_styled(doc, '5.2 核心功能', level=2)

    add_table(doc,
        ['功能', '说明'],
        [
            ['标的切换', '左侧边栏下拉选择或快捷按钮切换10只标的（A股/港股分组）'],
            ['参数调节', '每个指标独立调节参数（周期、平滑系数等），滑块拖动即时生效'],
            ['显示切换', '8个指标可独立勾选显示/隐藏，MA均线也可单独控制'],
            ['交互K线图', '蜡烛图 + 布林带上/中/下轨 + MA20/MA60 + 成交量柱状图 + 十字光标数据提示'],
            ['技术指标图', 'RSI / MACD / 布林带 / ATR / KDJ / CCI / ADX 各独立面板，含参考线'],
            ['Tooltip联动', '所有图表通过echarts.connect同步十字光标，鼠标悬浮任一图表即可查看同日全部指标数据'],
            ['缩放拖拽', 'K线图支持Ctrl+滚轮缩放和拖拽平移，底部滑块可调节时间范围'],
            ['回看天数', '支持60天/120天/250天/全部四档切换，指标基于全量数据计算后截取显示'],
            ['响应式布局', '桌面端左侧边栏+右侧图表双栏布局，移动端自动切换为上下堆叠'],
        ])

    add_heading_styled(doc, '5.3 技术架构', level=2)
    add_paragraph(doc,
        '看板采用纯前端架构：数据层由Python脚本（build_dashboard.py）读取CSV文件并序列化为嵌入式JSON；'
        '计算层将所有8个指标的计算逻辑（EMA、WilderSmooth、SMA、StdDev等）翻译为JavaScript，'
        '在浏览器端实时运算，支持参数动态调整；渲染层使用ECharts 5.5（CDN加载），'
        '所有图表共享同一个实例管理器，切换标的时自动释放旧实例。'
        '全量计算+显示截取的架构设计确保了回看天数切换时指标的准确性。',
        first_line_indent=Cm(0.74))

    add_heading_styled(doc, '5.4 部署方式', level=2)
    add_paragraph(doc,
        '看板已部署到GitHub Pages：https://rebekahlllimx.github.io/quant-trading-ai/Task2/dashboard/index.html。'
        '本地使用可直接双击打开HTML文件，或通过任意HTTP服务器托管。每次push后GitHub Pages自动同步部署。',
        first_line_indent=Cm(0.74))


def write_conclusion(doc):
    """第六章: 总结"""
    add_heading_styled(doc, '六、总结', level=1)

    add_paragraph(doc,
        '本报告完成了以下工作：',
        first_line_indent=Cm(0.74))

    items = [
        '对贵州茅台（600519.SH）的日线数据进行了系统的数据诊断分析，包括缺失值检查、描述性统计、正态性检验和异常值检测。'
        '结果表明：数据完整性良好，日收益率分布呈现"左偏、尖峰厚尾"的非正态特征（JB检验p≈0），年化波动率为20.35%，最大回撤为−23.47%。',

        '从历史演化、计算公式、参数讨论和信号含义四个维度，系统分析了八个技术指标（RSI、MACD、布林带、ATR、KDJ、MA、CCI、ADX）。'
        '每个指标的分析都回溯了其发明人的原始动机和后续的学术/实践演化，强调了参数选择的底层逻辑——参数值映射的是时间尺度和信号敏感度之间的权衡。',

        '通过mplfinance和ECharts实现了两种形式的可视化：静态matplotlib图表（7张，用于报告）和交互式HTML看板（用于动态分析）。'
        '交互看板支持10只A+H股票的切换、8个指标的独立参数调节和实时图表更新，可作为独立的分析工具使用。',

        '建立了指标互补性分析框架：按观察维度将8个指标分为趋势方向、趋势强度、动量速度、波动烈度、极端位置和短期拐点六类，'
        '并明确了不同市场状态（震荡市vs趋势市）下各指标的适用性和局限。ADX的值（>25或<20）是切换分析框架的关键分界线。',
    ]

    for i, item in enumerate(items, 1):
        add_paragraph(doc, f'({i}) {item}', first_line_indent=Cm(0.74))

    add_paragraph(doc, '')
    add_paragraph(doc,
        '本研究的主要局限包括：(1) 分析标的以大盘蓝筹股为主，小盘股和成长股的指标行为可能不同；'
        '(2) 参数分析仅覆盖了有限的范围，未做全参数空间的网格扫描优化；'
        '(3) 指标的历史表现不保证未来信号的有效性——市场结构和参与者行为会随时间演化。'
        '后续研究可以考虑：引入机器学习方法对多指标信号进行组合优化，对不同行业和市值规模的标的进行分组对比，'
        '以及将本报告中的指标分析框架扩展到更细粒度的时间周期（如小时线和分钟线）。',
        first_line_indent=Cm(0.74))


# ═══════════════════════════════════════════════════════════════
# 主程序
# ═══════════════════════════════════════════════════════════════

def main():
    print("=" * 60)
    print("  生成 Task2 .docx 报告")
    print("=" * 60)

    # 加载茅台数据
    df = pd.read_csv(
        os.path.join(DATA_DIR, '600519_贵州茅台_A股_daily.csv'),
        encoding='utf-8-sig'
    )
    df['Date'] = pd.to_datetime(df['Date'])
    print(f"\n📂 加载数据: 贵州茅台 {len(df)}条")

    stock_info = {"code": "600519", "name": "贵州茅台", "market": "上交所主板"}

    # 创建文档
    doc = Document()

    # 设置默认样式
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = FONT_SIZE
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.5
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # 写各章节
    print(">>> 写入数据诊断...")
    write_data_diagnosis(doc, df)

    print(">>> 写入指标分析...")
    write_indicator_section(doc, df, stock_info)

    print(">>> 写入指标互补关系...")
    write_synergy(doc)

    print(">>> 写入参数演化与图表...")
    write_param_evolution(doc, df, stock_info)

    print(">>> 写入交互看板...")
    write_dashboard(doc)

    print(">>> 写入总结...")
    write_conclusion(doc)

    # 保存
    os.makedirs(os.path.dirname(OUTPUT_FILE), exist_ok=True)
    doc.save(OUTPUT_FILE)
    file_size_mb = os.path.getsize(OUTPUT_FILE) / (1024 * 1024)
    print(f"\n✅ 报告已保存: {OUTPUT_FILE}")
    print(f"   文件大小: {file_size_mb:.1f} MB")
    print("=" * 60)


if __name__ == "__main__":
    main()
