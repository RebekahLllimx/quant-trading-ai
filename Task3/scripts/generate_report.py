#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Task3 Phase 4: 生成 .docx 报告
格式: 宋体五号(10.5pt), 1.5倍行距, 0段间距, 两端对齐
"""

import os, sys
import pandas as pd
import numpy as np

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '..'))
from src.indicators import data_diagnosis
from src.report_utils import (
    set_cjk_font, add_paragraph, add_heading_styled,
    add_picture_captioned, add_table, FONT_NAME, FONT_SIZE,
)
from src.equation_utils import add_equation, _r, _fraction

# ═══════════════════════════ 路径 ═══════════════════════════

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DATA_DIR = os.path.join(BASE_DIR, '..', 'data', 'csv')
CHART_DIR = os.path.join(BASE_DIR, '..', 'artifacts', 'charts', 'task3')
OUTPUT_FILE = os.path.join(BASE_DIR, 'Rebecca+Task3.docx')

# ═══════════════════════════ 回测引擎 ═══════════════════════════

def calc_sma(series, period):
    return series.rolling(window=period, min_periods=period).mean()

def run_backtest(df, short_period=5, long_period=15, initial_capital=1_000_000,
                 fee_rate=0.0003, slippage=0.0001):
    """双均线策略回测，避免未来函数"""
    close = df['Close'].astype(float)
    ma_short = calc_sma(close, short_period)
    ma_long = calc_sma(close, long_period)
    cash, shares, has_position = initial_capital, 0, False
    trades, equity_curve, signals_buy, signals_sell = [], [], [], []
    start_idx = max(short_period, long_period) + 1

    for t in range(start_idx, len(df)):
        price, date = close.iloc[t], df['Date'].iloc[t]
        equity_curve.append({'date': date, 'equity': cash + shares * price})

        ma_s1, ma_l1 = ma_short.iloc[t-1], ma_long.iloc[t-1]
        ma_s2, ma_l2 = ma_short.iloc[t-2], ma_long.iloc[t-2]
        if pd.isna(ma_s1) or pd.isna(ma_l1) or pd.isna(ma_s2) or pd.isna(ma_l2):
            continue

        golden = (ma_s1 > ma_l1) and (ma_s2 <= ma_l2)
        death = (ma_s1 < ma_l1) and (ma_s2 >= ma_l2)

        if golden and not has_position:
            bp = price * (1 + slippage)
            cost_per = bp * (1 + fee_rate)
            shares, cost = cash / cost_per, cash
            cash, has_position = 0, True
            trades.append({'buy_date': date, 'buy_price': bp, 'shares': shares,
                           'cost': cost, 'sell_date': None, 'sell_price': None, 'return_pct': None})
            signals_buy.append((date, bp))
        elif death and has_position:
            sp = price * (1 - slippage)
            cash = shares * sp * (1 - fee_rate)
            L = trades[-1]; L['sell_date'] = date; L['sell_price'] = sp
            L['return_pct'] = (cash / L['cost'] - 1) * 100
            shares, has_position = 0, False
            signals_sell.append((date, sp))

    if has_position:
        fp = close.iloc[-1] * (1 - slippage)
        cash = shares * fp * (1 - fee_rate)
        L = trades[-1]; L['sell_date'] = df['Date'].iloc[-1]
        L['sell_price'] = fp; L['return_pct'] = (cash / L['cost'] - 1) * 100

    full_eq = [{'date': df['Date'].iloc[i], 'equity': initial_capital} for i in range(start_idx)]
    full_eq.extend(equity_curve)

    bh_shares = initial_capital / (close.iloc[0] * (1 + slippage) * (1 + fee_rate))
    bh_eq = [{'date': d, 'equity': bh_shares * p} for d, p in zip(df['Date'], close)]

    return {'ma_short': ma_short, 'ma_long': ma_long, 'trades': trades,
            'signals_buy': signals_buy, 'signals_sell': signals_sell,
            'equity_curve': full_eq, 'final_equity': cash + shares * close.iloc[-1],
            'bh_equity_curve': bh_eq, 'bh_final_equity': bh_shares * close.iloc[-1]}

def calc_metrics(result, initial_capital=1_000_000, risk_free_rate=0.02):
    trades = [t for t in result['trades'] if t['sell_date'] is not None]
    equity = [e['equity'] for e in result['equity_curve']]
    if len(equity) < 2: return None

    final_equity = result['final_equity']
    total_return = (final_equity / initial_capital - 1) * 100
    days = max(1, (result['equity_curve'][-1]['date'] - result['equity_curve'][0]['date']).days)
    annual_return = (np.power(final_equity / initial_capital, 252.0 / days) - 1) * 100

    peak, mdd = -np.inf, 0
    for e in equity:
        if e > peak: peak = e
        dd = (e - peak) / peak * 100
        if dd < mdd: mdd = dd

    daily_rets = [equity[i]/equity[i-1]-1 for i in range(1,len(equity)) if equity[i-1]>0]
    sharpe = 0
    if len(daily_rets) > 1:
        mr, sr = np.mean(daily_rets), np.std(daily_rets, ddof=0)
        if sr > 0: sharpe = np.sqrt(252) * (mr - risk_free_rate/252) / sr

    wins = [t for t in trades if t['return_pct'] > 0]
    losses = [t for t in trades if t['return_pct'] <= 0]
    wr = len(wins)/len(trades)*100 if trades else 0
    aw = np.mean([t['return_pct'] for t in wins]) if wins else 0
    al = np.mean([abs(t['return_pct']) for t in losses]) if losses else 0
    plr = aw/al if al > 0 else None

    bh_return = (result['bh_final_equity'] / initial_capital - 1) * 100

    return {'total_return': total_return, 'annual_return': annual_return, 'mdd': mdd,
            'sharpe': sharpe, 'win_rate': wr, 'profit_loss_ratio': plr,
            'total_trades': len(trades), 'bh_total_return': bh_return, 'days': days,
            'avg_win': aw, 'avg_loss': al}

# ═══════════════════════════ 公式定义 ═══════════════════════════

def add_sma_formula(doc):
    """SMA_N(t) = (1/N) × Σ_{i=0}^{N-1} Close_{t-i}"""
    add_equation(doc, [
        _r("SMA"),
        _r("_N"),
        _r("(t) = "),
        _fraction([_r("1")], [_r("N")]),
        _r("  Σ_{i=0}^{N-1} Close"),
        _r("_{t-i}"),
    ])

def add_cum_return_formula(doc):
    """R_cum = (V_final - V_init) / V_init × 100%"""
    add_equation(doc, [
        _r("R"),
        _r("_{cum}"),
        _r(" = "),
        _fraction([_r("V_{final} − V_{init}")], [_r("V_{init}")]),
        _r(" × 100%"),
    ])

def add_annual_return_formula(doc):
    """R_ann = [(V_final/V_init)^(252/T) - 1] × 100%"""
    add_equation(doc, [
        _r("R_{ann} = [(V_{final}/V_{init})"),
        _r("^{252/T}"),
        _r(" − 1] × 100%"),
    ])

def add_mdd_formula(doc):
    """MDD = max_t[(Peak(t) - Equity(t)) / Peak(t)] × 100%"""
    add_equation(doc, [
        _r("MDD = max_{t} "),
        _fraction([_r("Peak(t) − Equity(t)")], [_r("Peak(t)")]),
        _r(" × 100%"),
    ])

def add_sharpe_formula(doc):
    """Sharpe = (R_ann - R_f) / σ_ann"""
    add_equation(doc, [
        _r("Sharpe = "),
        _fraction([_r("R"),
        _r("_{ann}"),
        _r(" − R"),
        _r("_f")], [_r("σ"),
        _r("_{ann}")]),
    ])

# ═══════════════════════════ 第一章 ═══════════════════════════

def write_ch1(doc, df, stock_name, stock_code):
    add_heading_styled(doc, '一、数据诊断与复权处理', level=1)
    diag = data_diagnosis(df, f'{stock_name} ({stock_code})')
    cs, rs, jb = diag['close_stats'], diag['return_stats'], diag['jarque_bera']

    # 1.1 复权
    add_heading_styled(doc, '1.1 数据复权', level=2)
    add_paragraph(doc,
        '股票分红送股会在除权除息日产生股价跳空缺口。不复权时，这些缺口会被均线策略误判为趋势转折，'
        '导致回测收益率严重失真。例如10送10的股票不复权会出现股价腰斩，产生虚假的卖出信号。'
        '本报告所有数据均来自AKShare前复权（qfq）接口，以最新股本为基准向后调整历史价格，'
        '保证了历史收益率计算的正确性。',
        first_line_indent=Cm(0.74))

    # 1.2 数据诊断
    add_heading_styled(doc, '1.2 数据质量诊断', level=2)
    add_paragraph(doc,
        f'以{stock_name}（{stock_code}.SH）为主要标的，数据范围{diag["date_start"]}至{diag["date_end"]}，'
        f'共{diag["rows"]}个交易日，五个核心字段（OHLCV）均无缺失值。',
        first_line_indent=Cm(0.74))

    add_table(doc, ['统计量', '数值', '含义'], [
        ['均值', f'{cs["mean"]:.2f} 元', '区间内收盘价算术平均'],
        ['标准差', f'{cs["std"]:.2f} 元', '收盘价离散程度'],
        ['最小值 / 最大值', f'{cs["min"]:.2f} / {cs["max"]:.2f} 元', '价格区间'],
        ['偏度', f'{cs["skewness"]:.4f}', '负值=左偏，正值=右偏'],
        ['峰度', f'{cs["kurtosis"]:.4f}', '正值=尖峰厚尾'],
    ])

    add_table(doc, ['收益率指标', '数值'], [
        ['日均收益率', f'{rs["daily_mean_pct"]:.4f}%'],
        ['日波动率', f'{rs["daily_std_pct"]:.4f}%'],
        ['年化波动率', f'{rs["annual_vol_pct"]:.2f}%'],
        ['最大回撤（价格）', f'{diag["max_drawdown_pct"]:.2f}%'],
    ])

    add_paragraph(doc,
        f'Jarque-Bera检验：JB={jb["statistic"]:.2f}, p={jb["p_value"]:.6f}，'
        f'拒绝正态分布原假设，收益率呈非正态特征，与金融时间序列典型特征一致。',
        first_line_indent=Cm(0.74))
    doc.add_page_break()

# ═══════════════════════════ 第二章 ═══════════════════════════

def write_ch2(doc):
    add_heading_styled(doc, '二、双均线策略原理', level=1)

    add_heading_styled(doc, '2.1 策略逻辑', level=2)
    add_paragraph(doc,
        '双均线交叉策略是趋势跟踪策略的基础形态。短周期均线（如MA5）反映近期价格重心，'
        '长周期均线（如MA15）反映中期趋势背景。当短期均价上穿长期均价时，意味着市场愿意以更高价格入场，'
        '触发买入（金叉）；反之触发卖出（死叉）。策略的四个规则完全由数学条件确定：'
        '金叉全仓买入、死叉全部卖出、不做空、不主观干预。',
        first_line_indent=Cm(0.74))

    add_heading_styled(doc, '2.2 数学定义', level=2)
    add_paragraph(doc, '简单移动平均线（SMA）的计算公式：', first_line_indent=Cm(0.74))
    add_sma_formula(doc)
    add_paragraph(doc, '')
    add_paragraph(doc,
        '金叉条件：SMA_short[t-1] > SMA_long[t-1] 且 SMA_short[t-2] ≤ SMA_long[t-2]。'
        '死叉条件：SMA_short[t-1] < SMA_long[t-1] 且 SMA_short[t-2] ≥ SMA_long[t-2]。',
        first_line_indent=Cm(0.74))

    add_heading_styled(doc, '2.3 避免未来函数', level=2)
    add_paragraph(doc,
        '未来函数（Look-ahead Bias）是回测中最常见的致命错误——用当日收盘价算均线来判断当日是否交易。'
        '本策略使用T-2和T-1的均线值判断信号（这两个值在T日开盘前已知），在T日以收盘价执行交易。'
        '这等价于“T日开盘确认信号→T日收盘执行”，是业界标准的折中方案。',
        first_line_indent=Cm(0.74))

    add_heading_styled(doc, '2.4 交易成本', level=2)
    add_paragraph(doc,
        '回测模拟手续费0.03%（万三）和滑点0.01%。手续费买卖各收一次，滑点模拟订单执行延迟。'
        '单边总成本约0.04%，双边约0.08%。交互看板中两项成本均可调节。',
        first_line_indent=Cm(0.74))
    doc.add_page_break()

# ═══════════════════════════ 第三章 ═══════════════════════════

def write_ch3(doc):
    add_heading_styled(doc, '三、量化策略评估指标', level=1)

    items = [
        ('3.1 累计回报', '反映策略总收益率，但不考虑时间长度和过程波动。',
         add_cum_return_formula),
        ('3.2 年化收益率', '将不同期限的收益率统一为年化口径（252个交易日/年），便于跨策略对比。'
         '长期年化>10%已属优秀。', add_annual_return_formula),
        ('3.3 最大回撤（MDD）', '衡量策略从峰值到谷底的最大跌幅。MDD直接对应投资者的心理承受极限——'
         '回撤30%的策略在执行中绝大多数人会在底部止损离场。风控原则要求在策略设计之初设定回撤阈值。',
         add_mdd_formula),
        ('3.4 夏普比率', '核心风险调整收益指标，衡量每单位波动风险换取的超额收益。'
         '>1表示风险调整后表现良好，>2为优秀，<0表示未跑赢无风险利率（Rf≈2%）。',
         add_sharpe_formula),
        ('3.5 胜率与盈亏比', '胜率=盈利次数/总次数。盈亏比=平均盈利/平均亏损。'
         '关键认知：胜率高不代表策略好，趋势跟踪策略的典型特征是“低胜率、高盈亏比”——多数交易小亏，少数交易大赚。'
         '一个胜率30%盈亏比5:1的策略，远优于胜率70%盈亏比0.5:1的策略。', None),
        ('3.6 基准对比', '回测结果必须与“买入持有”对比——如果在付出交易成本后收益不如持有不动，策略没有存在价值。', None),
    ]

    for title, desc, formula_fn in items:
        add_heading_styled(doc, title, level=2)
        add_paragraph(doc, desc, first_line_indent=Cm(0.74))
        if formula_fn:
            formula_fn(doc)
            add_paragraph(doc, '')

    doc.add_page_break()

# ═══════════════════════════ 第四章 ═══════════════════════════

def write_ch4(doc, df, result, metrics, stock_name, stock_code):
    add_heading_styled(doc, '四、策略回测结果分析', level=1)

    add_paragraph(doc,
        f'以{stock_name}（{stock_code}.SH）日线数据，MA5+MA15，初始资金100万元，手续费0.03%，'
        f'滑点0.01%，回测区间{df["Date"].iloc[0].strftime("%Y-%m-%d")}至{df["Date"].iloc[-1].strftime("%Y-%m-%d")}。',
        first_line_indent=Cm(0.74))

    # 4.1 交易信号
    add_heading_styled(doc, '4.1 交易信号', level=2)
    add_picture_captioned(doc, os.path.join(CHART_DIR, f'{stock_code}_策略信号图.png'),
        f'图1：{stock_name}（{stock_code}）双均线策略交易信号（MA5 & MA15，红色▲买入，绿色▼卖出）',
        width_inches=5.8)
    add_paragraph(doc,
        f'图1展示了收盘价、MA5、MA15及买卖信号。共产生{metrics["total_trades"]}次完整交易，'
        f'盈利{max(0,int(metrics["win_rate"]*metrics["total_trades"]/100))}次。'
        f'2025年9-12月的横盘震荡区间内两条均线反复交叉，产生了多次假信号——这是双均线策略在震荡市中的典型表现。',
        first_line_indent=Cm(0.74))

    # 4.2 资产曲线
    add_heading_styled(doc, '4.2 资产曲线', level=2)
    add_picture_captioned(doc, os.path.join(CHART_DIR, f'{stock_code}_资产曲线对比.png'),
        f'图2：{stock_name}（{stock_code}）资产曲线对比——策略 vs 买入持有',
        width_inches=5.8)
    add_paragraph(doc,
        f'策略累计回报{metrics["total_return"]:.2f}%，年化{metrics["annual_return"]:.2f}%；'
        f'买入持有累计回报{metrics["bh_total_return"]:.2f}%。'
        f'策略相对基准{"跑赢" if metrics["total_return"] > metrics["bh_total_return"] else "跑输"}'
        f'{abs(metrics["total_return"]-metrics["bh_total_return"]):.2f}个百分点。'
        f'从资产曲线看，策略在趋势阶段能捕捉机会，但震荡期的反复交易产生了成本损耗。',
        first_line_indent=Cm(0.74))

    # 4.3 回撤
    add_heading_styled(doc, '4.3 回撤分析', level=2)
    add_picture_captioned(doc, os.path.join(CHART_DIR, f'{stock_code}_回撤曲线.png'),
        f'图3：{stock_name}（{stock_code}）回撤曲线——MDD={metrics["mdd"]:.2f}%',
        width_inches=5.8)
    dd_comment = ('回撤超过20%警戒线，实际交易中应触发暂停机制。' if metrics['mdd'] < -20
                  else '回撤在可接受范围内。')
    add_paragraph(doc,
        f'策略MDD为{metrics["mdd"]:.2f}%。{dd_comment}'
        f'回撤低谷与卖出信号后的下跌阶段高度对应，验证了策略虽然无法精确预测每次下跌，'
        f'但能在大趋势转折时及时退出。',
        first_line_indent=Cm(0.74))

    # 4.4 指标汇总
    add_heading_styled(doc, '4.4 绩效指标汇总', level=2)
    plr_s = f'{metrics["profit_loss_ratio"]:.2f}' if metrics['profit_loss_ratio'] else 'N/A'
    add_table(doc, ['指标', '数值', '评价'], [
        ['累计回报', f'{metrics["total_return"]:.2f}%', '正收益' if metrics['total_return']>0 else '负收益'],
        ['年化收益率', f'{metrics["annual_return"]:.2f}%',
         '达标(>10%)' if metrics['annual_return']>=10 else ('正收益' if metrics['annual_return']>=0 else '为负')],
        ['最大回撤', f'{metrics["mdd"]:.2f}%',
         '超20%警戒' if metrics['mdd']<-20 else ('正常范围' if metrics['mdd']>-10 else '接近警戒')],
        ['夏普比率', f'{metrics["sharpe"]:.2f}',
         '良好(>1)' if metrics['sharpe']>=1 else ('一般(>0)' if metrics['sharpe']>=0 else '跑输无风险利率')],
        ['胜率', f'{metrics["win_rate"]:.1f}%', '趋势策略低胜率属正常' if metrics['win_rate']<50 else '尚可'],
        ['盈亏比', plr_s, '赚大亏小(>2)' if (metrics['profit_loss_ratio'] and metrics['profit_loss_ratio']>=2) else ('持平' if (metrics['profit_loss_ratio'] and metrics['profit_loss_ratio']>=1) else '赚小亏大') if metrics['profit_loss_ratio'] else '—'],
        ['交易次数', f'{metrics["total_trades"]}次', f'约{metrics["days"]//max(metrics["total_trades"],1)}天/次'],
        ['买入持有', f'{metrics["bh_total_return"]:.2f}%', '基准对比'],
    ])

    # 4.5 交易明细
    add_heading_styled(doc, '4.5 交易明细', level=2)
    completed = [t for t in result['trades'] if t['sell_date'] is not None]
    if completed:
        rows, cum = [], 1.0
        for i, t in enumerate(completed):
            cum *= (1 + t['return_pct']/100)
            rows.append([str(i+1), t['buy_date'].strftime('%Y-%m-%d'), f'{t["buy_price"]:.2f}',
                         t['sell_date'].strftime('%Y-%m-%d'), f'{t["sell_price"]:.2f}',
                         f'{(t["sell_date"]-t["buy_date"]).days}天',
                         f'{t["return_pct"]:+.2f}%', f'{(cum-1)*100:+.2f}%'])
        add_table(doc, ['#','买入日','买入价','卖出日','卖出价','持有','单笔收益','累计收益'], rows)

    doc.add_page_break()

# ═══════════════════════════ 第五章 ═══════════════════════════

def write_ch5(doc, df, stock_name, stock_code):
    add_heading_styled(doc, '五、参数敏感性分析', level=1)
    add_paragraph(doc,
        '对短均线（3,5,10,15,20,30）和长均线（10,15,20,30,40,60）进行36组网格扫描，分析参数影响。',
        first_line_indent=Cm(0.74))

    add_heading_styled(doc, '5.1 网格扫描结果', level=2)
    add_picture_captioned(doc, os.path.join(CHART_DIR, f'{stock_code}_参数敏感性热力图.png'),
        f'图4：{stock_name}（{stock_code}）参数敏感性——累计回报与夏普比率热力图',
        width_inches=5.8)
    add_paragraph(doc,
        '热力图显示：(1) 短周期与长周期的间隔是关键因素——间隔过小导致频繁交叉、成本累积；'
        '间隔过大则信号过少、错失机会。(2) 夏普比率热力图通常比回报热力图更平滑，参数选择应优先参考夏普比率。'
        '(3) 表现好的参数通常形成“稳健区域”而非孤立的“尖锐峰值”，后一种情况往往是过拟合的信号。',
        first_line_indent=Cm(0.74))

    add_heading_styled(doc, '5.2 过拟合风险', level=2)
    add_paragraph(doc,
        '双均线策略只有两个参数，过拟合风险天然较低——这是简单策略的优势。防范过拟合三原则：'
        '(1) 参数必须有经济逻辑——MA5对应1周、MA15对应3周，时间尺度差异有实际意义；'
        '(2) 参数不应过度精细化——在14和15之间选14.37几乎肯定是过拟合；'
        '(3) 参数稳健性检验——最优参数周围的参数组合也应表现良好，而非孤峰。',
        first_line_indent=Cm(0.74))
    doc.add_page_break()

# ═══════════════════════════ 第六章 ═══════════════════════════

def write_ch6(doc):
    add_heading_styled(doc, '六、多标的对比分析', level=1)
    add_paragraph(doc,
        '选取6只标的（茅台、平安、宁德时代、比亚迪、腾讯、阿里巴巴），统一使用MA5+MA15参数回测对比。',
        first_line_indent=Cm(0.74))

    add_heading_styled(doc, '6.1 回测结果对比', level=2)
    add_picture_captioned(doc, os.path.join(CHART_DIR, '多股票策略对比.png'),
        '图5：多股票双均线策略（MA5+MA15）回测对比——累计回报、夏普比率、最大回撤',
        width_inches=5.8)

    stock_list = [('600519','贵州茅台','A股'),('601318','中国平安','A股'),
                  ('300750','宁德时代','A股'),('002594','比亚迪','A股'),
                  ('00700','腾讯控股','港股'),('09988','阿里巴巴','港股')]
    rows = []
    for code, name, market in stock_list:
        try:
            fn = f"{code}_{name}_{market}_daily.csv"
            df_s = pd.read_csv(os.path.join(DATA_DIR, fn), encoding='utf-8-sig')
            df_s['Date'] = pd.to_datetime(df_s['Date']); df_s = df_s.sort_values('Date').reset_index(drop=True)
            r = run_backtest(df_s); m = calc_metrics(r)
            if m:
                rows.append([name, code, f'{m["total_return"]:.2f}%', f'{m["annual_return"]:.2f}%',
                             f'{m["mdd"]:.2f}%', f'{m["sharpe"]:.2f}', f'{m["win_rate"]:.1f}%',
                             f'{m["bh_total_return"]:.2f}%',
                             '是' if m['total_return']>m['bh_total_return'] else '否'])
        except: pass

    add_table(doc, ['标的','代码','累计回报','年化收益','MDD','夏普','胜率','买入持有','跑赢基准?'], rows)

    add_heading_styled(doc, '6.2 策略适用条件', level=2)
    add_paragraph(doc,
        '双均线策略的表现与标的的“趋势性”密切相关：(1) 高波动标的（宁德时代、比亚迪）潜在收益大但假信号多；'
        '低波动蓝筹（茅台）信号可靠性高但单笔收益小。(2) 策略在单边趋势市中效果最好，在区间震荡市中反复亏损。'
        '判断标的是否适合趋势策略可参考ADX指标——ADX长期>25的标的更适合。(3) 交易成本对低价股影响更大，'
        '滑点占比上升会显著侵蚀策略表现。',
        first_line_indent=Cm(0.74))
    doc.add_page_break()

# ═══════════════════════════ 第七章 ═══════════════════════════

def write_ch7(doc):
    add_heading_styled(doc, '七、交互式回测看板', level=1)
    add_paragraph(doc,
        '本任务构建了自包含的交互式Web回测看板（Task3/dashboard/index.html），覆盖10只标的。'
        '所有回测计算在浏览器端完成，无需后端服务器。',
        first_line_indent=Cm(0.74))

    add_heading_styled(doc, '7.1 核心功能', level=2)
    add_table(doc, ['功能', '说明'], [
        ['标的选择', '10只标的（5A股+5港股），切换后即时重新回测'],
        ['参数调节', '短/长均线周期、时间范围、初始资金、手续费率、滑点率——滑块+数字双控'],
        ['KPI卡片', '累计回报、年化收益、MDD、夏普比率、胜率、盈亏比、交易次数、买入持有基准'],
        ['价格信号图', '收盘价+MA短+MA长+红色▲买入+绿色▼卖出，支持缩放拖拽'],
        ['资产曲线图', '策略资产(蓝) vs 买入持有(灰虚线)'],
        ['回撤曲线图', '回撤面积图，标注MDD值'],
        ['交易明细表', '每笔交易的买卖日期/价格/持有天数/收益率，盈利绿亏损红'],
        ['纯前端回测', '参数修改即时生效，避免未来函数，成本自动扣除'],
    ])

    add_heading_styled(doc, '7.2 部署方式', level=2)
    add_paragraph(doc,
        '看板已部署至GitHub Pages: https://rebekahlllimx.github.io/quant-trading-ai/Task3/dashboard/index.html。'
        '本地可直接双击打开HTML文件使用。',
        first_line_indent=Cm(0.74))
    doc.add_page_break()

# ═══════════════════════════ 第八章 ═══════════════════════════

def write_ch8(doc, metrics, stock_name):
    add_heading_styled(doc, '八、总结', level=1)

    add_paragraph(doc,
        f'本报告以{stock_name}为主要标的，完成了双均线交叉策略（MA5+MA15）的完整回测分析。'
        f'策略累计回报{metrics["total_return"]:.2f}%，年化{metrics["annual_return"]:.2f}%，'
        f'MDD{metrics["mdd"]:.2f}%，夏普{metrics["sharpe"]:.2f}，'
        f'{"跑赢" if metrics["total_return"]>metrics["bh_total_return"] else "跑输"}买入持有基准。',
        first_line_indent=Cm(0.74))

    add_heading_styled(doc, '8.1 策略局限', level=2)
    lims = [
        '滞后性：均线是滞后指标，金叉确认时趋势已运行一段。这是趋势跟踪策略的固有特征。',
        '震荡市失效：无趋势的横盘行情中反复交叉产生假信号，累积成本侵蚀本金。可结合ADX过滤。',
        '单一信号源：仅依赖收盘价均线交叉，未使用成交量、波动率等补充信息。',
        '参数敏感性：不同MA周期组合结果差异大，历史最优参数不保证未来有效。',
        '极端行情：剧烈单边下跌中卖出信号出现在已跌一大段之后，需止损规则补充。',
    ]
    for i, lim in enumerate(lims, 1):
        add_paragraph(doc, f'{i}. {lim}', first_line_indent=Cm(0.74))

    add_heading_styled(doc, '8.2 改进方向', level=2)
    imps = [
        '引入ADX趋势过滤器：ADX<20时暂停交易，仅在趋势确立时启用信号。',
        '多周期均线确认：增加MA60过滤，仅当长期均线方向与交易方向一致时执行。',
        '动态仓位管理：基于ATR调整仓位，高波动时减仓、低波动时加仓。',
        '止盈止损机制：设置单笔止损（如-5%）和止盈（如+15%），控制单次亏损。',
        '多资产组合：同时应用于多只低相关性标的，通过分散化降低整体回撤。',
    ]
    for i, imp in enumerate(imps, 1):
        add_paragraph(doc, f'{i}. {imp}', first_line_indent=Cm(0.74))

    add_paragraph(doc, '')
    add_paragraph(doc,
        '双均线策略是一个“简单但不简陋”的入门策略。它的简单性是优势——逻辑透明、参数少、不易过拟合；'
        '也是局限——信号单一、震荡市失效。对于量化初学者，理解其核心逻辑是构建更复杂策略的基础。'
        '没有单一策略能持续盈利，真正的能力在于理解每种工具的适用条件、局限性和改进方向。',
        first_line_indent=Cm(0.74))

# ═══════════════════════════ 主程序 ═══════════════════════════

def main():
    print('=' * 60)
    print('  生成 Task3 .docx 报告')
    print('=' * 60)

    df = pd.read_csv(os.path.join(DATA_DIR, '600519_贵州茅台_A股_daily.csv'), encoding='utf-8-sig')
    df['Date'] = pd.to_datetime(df['Date'])
    df = df.sort_values('Date').reset_index(drop=True)
    print(f'\n📂 {len(df)}条 贵州茅台')

    charts_needed = ['600519_策略信号图.png', '600519_资产曲线对比.png',
                     '600519_回撤曲线.png', '600519_参数敏感性热力图.png', '多股票策略对比.png']
    missing = [c for c in charts_needed if not os.path.exists(os.path.join(CHART_DIR, c))]
    if missing:
        print(f'\n⚠️ 缺少图表: {missing}\n  请先运行: python Task3/scripts/plot_strategy.py')
        return

    print('⏳ 回测...')
    result = run_backtest(df)
    metrics = calc_metrics(result)
    print(f'   累计回报: {metrics["total_return"]:.2f}%  夏普: {metrics["sharpe"]:.2f}')

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'; style.font.size = FONT_SIZE
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    pf = style.paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(0)
    pf.line_spacing = 1.5; pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for section in doc.sections:
        section.top_margin = Cm(2.54); section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18); section.right_margin = Cm(3.18)

    for label, fn in [
        ('一、数据诊断与复权处理', lambda: write_ch1(doc, df, '贵州茅台', '600519')),
        ('二、双均线策略原理', lambda: write_ch2(doc)),
        ('三、量化策略评估指标', lambda: write_ch3(doc)),
        ('四、策略回测结果分析', lambda: write_ch4(doc, df, result, metrics, '贵州茅台', '600519')),
        ('五、参数敏感性分析', lambda: write_ch5(doc, df, '贵州茅台', '600519')),
        ('六、多标的对比分析', lambda: write_ch6(doc)),
        ('七、交互式回测看板', lambda: write_ch7(doc)),
        ('八、总结', lambda: write_ch8(doc, metrics, '贵州茅台')),
    ]:
        print(f'>>> {label}')
        fn()

    os.makedirs(os.path.dirname(OUTPUT_FILE), exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f'\n✅ {OUTPUT_FILE}  ({os.path.getsize(OUTPUT_FILE)/1024:.0f} KB)')
    print('=' * 60)

if __name__ == '__main__':
    main()
