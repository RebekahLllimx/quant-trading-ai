#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate the TASK7 Word report from validated local and platform evidence."""

from __future__ import annotations

import json
import os
import tempfile
import zipfile
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from lxml import etree


ROOT = Path(__file__).resolve().parents[2]
TASK_DIR = ROOT / "Task7"
DATA_DIR = ROOT / "data" / "task7"
CHART_DIR = ROOT / "artifacts" / "charts" / "task7"
OUTPUT = TASK_DIR / "Rebecca+Task7.docx"
DASHBOARD_URL = "https://rebekahlllimx.github.io/quant-trading-ai/Task7/dashboard/"
JOINQUANT_DIR = TASK_DIR / "inputs" / "joinquant"
PLATFORM_EVIDENCE_DIR = TASK_DIR / "output"

# Match the restrained black/gray typography used in the earlier reports.
# Use the installed macOS PostScript family name so LibreOffice can render
# Chinese reliably while preserving the Song-style appearance.
FONT = "Songti SC"
BODY_SIZE = Pt(10.5)
INK = RGBColor(0x00, 0x00, 0x00)
MUTED = RGBColor(0x00, 0x00, 0x00)
LIGHT_GRAY = "D9D9D9"
TABLE_WIDTH_DXA = 8520
TABLE_INDENT_DXA = 120


def set_font(run, size=BODY_SIZE, bold=None, color=INK, name=FONT):
    run.font.name = name
    run.font.size = size
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), name)


def ensure_font_table_entry(docx_path, font_name=FONT):
    """Declare the CJK font so LibreOffice does not substitute a Latin-only face."""
    docx_path = Path(docx_path)
    with zipfile.ZipFile(docx_path, "r") as source:
        font_table = etree.fromstring(source.read("word/fontTable.xml"))
        existing = {
            node.get(qn("w:name"))
            for node in font_table.findall(qn("w:font"))
        }
        if font_name in existing:
            return

        font = etree.SubElement(font_table, qn("w:font"))
        font.set(qn("w:name"), font_name)
        charset = etree.SubElement(font, qn("w:charset"))
        charset.set(qn("w:val"), "86")
        family = etree.SubElement(font, qn("w:family"))
        family.set(qn("w:val"), "roman")
        pitch = etree.SubElement(font, qn("w:pitch"))
        pitch.set(qn("w:val"), "variable")
        updated_font_table = etree.tostring(
            font_table,
            xml_declaration=True,
            encoding="UTF-8",
            standalone=True,
        )

        fd, temp_name = tempfile.mkstemp(
            prefix=f"{docx_path.stem}-font-",
            suffix=".docx",
            dir=docx_path.parent,
        )
        os.close(fd)
        try:
            with zipfile.ZipFile(temp_name, "w") as target:
                for item in source.infolist():
                    payload = (
                        updated_font_table
                        if item.filename == "word/fontTable.xml"
                        else source.read(item.filename)
                    )
                    target.writestr(item, payload)
            os.replace(temp_name, docx_path)
        finally:
            if os.path.exists(temp_name):
                os.unlink(temp_name)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    supplied_width = sum(widths)
    if supplied_width <= 0:
        raise ValueError(f"Table widths must be positive: {widths}")
    widths = [round(width * TABLE_WIDTH_DXA / supplied_width) for width in widths]
    widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_keep_with_next(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    keep = ppr.find(qn("w:keepNext"))
    if keep is None:
        ppr.append(OxmlElement("w:keepNext"))


def add_body(doc, text, indent=True, bold_lead=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.5
    if indent:
        pf.first_line_indent = Cm(0.74)
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_font(lead, bold=True)
        body = p.add_run(text[len(bold_lead) :])
        set_font(body)
    else:
        run = p.add_run(text)
        set_font(run)
    return p


def add_heading(doc, text, level=1):
    style = f"Heading {level}"
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    size = Pt(14 if level == 1 else 12)
    set_font(run, size=size, bold=True, color=INK)
    pf = p.paragraph_format
    pf.space_before = Pt(12 if level == 1 else 8)
    pf.space_after = Pt(6 if level == 1 else 4)
    pf.keep_with_next = True
    return p


def add_table(doc, headers, rows, widths, alignments=None, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(str(header))
        set_font(run, size=Pt(font_size), bold=True)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.alignment = (
                alignments[idx]
                if alignments
                else WD_ALIGN_PARAGRAPH.CENTER
            )
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(str(value))
            set_font(run, size=Pt(font_size))
    set_table_geometry(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_picture_path(doc, path, caption, width=5.8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    inline = run.add_picture(str(path), width=Inches(width))
    inline._inline.docPr.set("descr", caption)
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_after = Pt(12)
    cap = caption_p.add_run(caption)
    set_font(cap, size=Pt(9), color=MUTED)


def add_picture(doc, filename, caption, width=5.8):
    add_picture_path(doc, CHART_DIR / filename, caption, width)


def add_hyperlink(paragraph, text, url, font_size=10.5):
    part = paragraph.part
    relationship = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(color)
    rpr.append(underline)
    fonts = OxmlElement("w:rFonts")
    for key in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{key}"), FONT)
    rpr.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), str(int(font_size * 2)))
    rpr.append(size)
    run.append(rpr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def configure_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = BODY_SIZE
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for level, size, color in [(1, 14, INK), (2, 12, INK)]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.paragraph_format.space_before = Pt(12 if level == 1 else 8)
        style.paragraph_format.space_after = Pt(6 if level == 1 else 4)
        style.paragraph_format.keep_with_next = True

    for list_style in ["List Bullet", "List Number"]:
        style = doc.styles[list_style]
        style.font.name = FONT
        style.font.size = BODY_SIZE
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.5
    return doc


def add_list_item(doc, text, numbered=False):
    p = doc.add_paragraph(style="List Number" if numbered else "List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    set_font(r)
    return p


def pct(value):
    return f"{value * 100:.2f}%"


def number(value):
    return f"{value:,.0f}"


def load_inputs():
    metrics = json.loads(
        (DATA_DIR / "metadata" / "backtest_metrics.json").read_text(encoding="utf-8")
    )
    stress = json.loads(
        (DATA_DIR / "metadata" / "cost_stress_metrics.json").read_text(encoding="utf-8")
    )
    b_params = pd.read_csv(
        DATA_DIR / "processed" / "strategy_b_parameter_sensitivity.csv"
    )
    c_params = pd.read_csv(
        DATA_DIR / "processed" / "strategy_c_parameter_sensitivity.csv"
    )
    joinquant = json.loads(
        (JOINQUANT_DIR / "backtest_summary.json").read_text(encoding="utf-8")
    )
    dashboard = json.loads(
        (TASK_DIR / "dashboard" / "data" / "dashboard.json").read_text(
            encoding="utf-8"
        )
    )
    return metrics, stress, b_params, c_params, joinquant, dashboard


def write_report():
    metrics, stress, b_params, c_params, joinquant, dashboard = load_inputs()
    platform_results = {
        item["strategy"]: item for item in joinquant["backtests"]
    }
    short_results = {
        item["strategy"]: item for item in joinquant["short_window_backtests"]
    }
    short_settings = joinquant["short_window_settings"]
    doc = configure_document()

    add_heading(doc, "摘要", 1)
    add_body(
        doc,
        "本报告展示三套宽基ETF策略的设计、参数定型、长期回测、测试周表现与模拟盘部署状态。长期结果表明，策略C累计收益71.84%，是三者中唯一取得正超额收益的方案；策略B将最大回撤控制在12.31%，更接近防守型配置；策略A累计收益为-0.66%、最大回撤36.97%，说明简单高仓位择时在震荡和交易摩擦面前缺乏足够保护。综合收益与风险后，策略C作为主要观察对象，策略B作为防守型参照，策略A作为基础对照。",
    )
    add_body(
        doc,
        "核心判断：2026年7月20日至24日的五日测试期内，三套策略均保持现金，策略收益与回撤均为0%，同期沪深300上涨2.65%。这一结果显示三套趋势规则在快速上涨窗口中仍保持谨慎，降低了市场波动暴露，也产生了2.58%的相对机会成本。三组模拟盘于7月24日周五完成部署；7月25日为周六、A股休市，截至报告日期尚未产生部署后的运行数据。测试周采用冻结参数和T−1历史信息回放策略决策，属于模拟盘启动前的历史模拟；后续实际模拟净值、持仓和成交将在持续观察页面更新。",
        bold_lead="核心判断：",
    )

    add_heading(doc, "一、成果概览与研究范围", 1)
    add_heading(doc, "1.1 本次课程实践的核心成果", 2)
    add_body(
        doc,
        "本次实践形成了一套可以持续比较的三策略组合，评价范围同时覆盖历史收益、风险和市场暴露。三套方案分别代表简单择时、风险约束趋势跟随和跨ETF动量轮动，并在统一资金规模、基准和交易成本下接受检验。报告重点回答三个问题：不同规则带来了怎样的收益与市场暴露，回撤是否与收益相匹配，以及参数冻结后在测试周和模拟观察阶段表现如何。",
    )
    add_heading(doc, "1.2 聚宽平台的应用范围", 2)
    add_body(
        doc,
        "现有账号已完成注册并能正常登录，相关权限满足本次研究需要。实践覆盖聚宽的数据调用、策略编辑、历史回测、结果分析和模拟交易功能。长期回测用于评价策略表现，测试周用于观察冻结参数下的实际信号，三组模拟组合则承担后续连续验证。平台文档用于确认数据与交易规则，社区案例为策略C提供了思路；最终的标的范围、调仓频率和风险约束均根据本次研究目标重新设定，策略与组合保持私有。",
    )

    add_heading(doc, "二、三套策略代表不同的收益—风险取向", 1)
    add_body(
        doc,
        "三个策略使用500,000元初始资金，以沪深300作为共同基准。策略A代表简单、直接但高暴露的择时方案；策略B在趋势判断之外增加仓位预算和止损，更强调资金保护；策略C在三只宽基ETF之间选择动量最强者，并在趋势不足时持有现金，目标是在控制仓位的同时争取相对收益。",
    )
    add_table(
        doc,
        ["策略", "核心信号", "仓位与风控", "组合定位"],
        [
            ["A", "510300收盘价高于20日均线", "目标100%；无波动率止损", "基础对照"],
            ["B", "510300的10日均线高于30日均线", "风险预算、波动率止损、仓位上限50%", "防守型参照"],
            ["C", "三只宽基ETF的15日趋势强度排序", "每周选择最高正分ETF，仓位50%；弱势时持币", "主要观察对象"],
        ],
        [900, 2300, 3770, 2390],
        font_size=8.8,
    )
    add_body(
        doc,
        "策略C的标的池固定为沪深300ETF、中证500ETF和创业板ETF，覆盖大盘、中盘与成长风格。趋势评分同时考虑上涨速度和路径稳定性，因此短期急涨但走势反复的标的不会自动获得最高权重。周频调仓降低了日度噪声和换手，50%的仓位上限则为组合保留了现金缓冲。",
    )

    add_heading(doc, "三、比较口径与结论边界", 1)
    add_heading(doc, "3.1 时间划分与共同口径", 2)
    add_body(
        doc,
        "参数选择使用2016—2023年开发样本和2024—2025年验证样本，2026年只用于参数冻结后的观察。JoinQuant长期回测覆盖2019年1月1日至2026年7月17日，7月20日至24日另设五个交易日测试窗口，两段互不重叠。各策略使用相同初始资金、沪深300基准和ETF交易成本，信号以已完成的历史数据为依据，并在下一交易日执行。",
    )
    add_heading(doc, "3.2 如何理解报告中的两组结果", 2)
    add_body(
        doc,
        "策略排名与核心结论以JoinQuant回测结果为准；扩展样本分析用于进一步解释仓位、阶段表现和成本敏感性。由于行情复权、成交价格和成本处理可能存在差异，两组结果不要求数值完全一致，但应对收益方向和风险特征给出相互印证。扩展分析只用于解释结果，不用于事后重新选择更有利的参数。",
    )

    add_heading(doc, "四、参数定型强调跨阶段稳定性", 1)
    add_picture(
        doc,
        "figure2_parameter_sensitivity.png",
        "图1：策略B均线参数与策略C动量窗口的开发期、验证期累计收益",
        width=5.8,
    )
    add_body(
        doc,
        "图1显示，策略B的10/30组合在开发期和验证期分别取得2.00%和2.11%的正收益，是三组均线中跨阶段方向最一致的方案。5/20和20/60在开发期分别亏损4.31%和5.66%，进入验证期后才转正，说明均线信号对市场阶段较敏感。策略C的15日窗口在两阶段分别取得6.11%和36.29%；20日窗口验证期收益略高至36.77%，开发期却为-0.55%；30日窗口开发期达到9.21%，验证期降至8.02%。参数排名随阶段发生变化，单期峰值缺乏稳定性。10/30均线和15日动量窗口保留了较完整的跨阶段正向信号。",
    )

    add_heading(doc, "五、长期结果：策略C领先，策略B更稳健", 1)
    platform_rows = []
    for strategy in "ABC":
        item = platform_results[strategy]
        platform_rows.append(
            [
                strategy,
                pct(item["strategy_return"]),
                pct(item["annualized_return"]),
                pct(item["excess_return"]),
                pct(item["max_drawdown"]),
                f'{item["sharpe"]:.3f}',
                f'{item["beta"]:.3f}',
                pct(item["strategy_volatility"]),
            ]
        )
    add_table(
        doc,
        ["策略", "累计收益", "年化收益", "超额收益", "最大回撤", "夏普", "Beta", "波动率"],
        platform_rows,
        [900, 1270, 1270, 1270, 1270, 900, 900, 1580],
        font_size=8.5,
    )
    add_body(
        doc,
        "长期回测表明，策略A的累计收益为-0.66%，最大回撤却达到36.97%，收益与下行风险明显失衡。策略B把最大回撤降至12.31%，Beta降至0.164、波动率降至4.70%，防守特征清晰；其累计收益仅为2.91%，超额收益为-31.59%，说明低波动主要来自克制参与，趋势信号本身没有形成显著收益优势。",
    )
    add_body(
        doc,
        "策略C累计收益71.84%、年化收益7.69%、最大回撤16.91%，相对沪深300取得14.23%的累计超额，是三者中收益与回撤匹配度最好的方案。其Beta为0.426，说明50%仓位上限有效压低了市场方向暴露；夏普比率0.315和30.37%的相对最大回撤同时表明，超额收益的形成过程仍有较大波动。表中最重要的发现是，策略C的领先同时体现在收益和绝对回撤两个维度，策略B的优势则集中在防守。",
    )

    add_heading(doc, "六、收益来源与主要风险", 1)
    add_picture(
        doc,
        "figure1_nav_drawdown.png",
        "图2：扩展样本下三策略与沪深300的净值及回撤（2016—2026）",
        width=5.8,
    )
    add_body(
        doc,
        "图2显示三条策略净值具有明显不同的路径。策略A在2021年后持续走弱，回撤长期处于较深区间，说明简单MA20信号在反复震荡中容易连续承受假突破和交易摩擦。策略B的净值长期围绕初始水平窄幅波动，回撤显著较浅，其稳定性主要来自较低市场参与度。策略C在2016—2023年的领先并不突出，主要增值集中在2024年以后，并在期末超过沪深300。策略C的优势由特定阶段贡献，收益路径存在明显的市场风格依赖。",
    )
    add_picture(
        doc,
        "figure3_period_performance.png",
        "图3：开发期、验证期与参数冻结后的2026年样本外累计收益",
        width=5.8,
    )
    add_body(
        doc,
        "图3进一步定位了收益来源。策略A在开发期亏损34.06%，验证期反弹12.16%，参数冻结后的2026年又下跌6.41%，收益方向反复，缺少稳定信号。策略B在开发期和验证期仅分别上涨2.00%和2.11%，2026年下跌2.61%，呈现低波动、低收益的防守特征。策略C三个阶段分别上涨6.11%、36.29%和7.09%，在每个阶段都领先另外两套策略；其中验证期贡献最大，确认了长期领先对2024—2025年市场环境的依赖。2026年仍为正提供了一项样本外支持，当前样本长度尚不足以消除风格集中风险。",
    )

    full_rows = []
    for strategy in "ABC":
        m = metrics[strategy]["full"]
        full_rows.append(
            [
                strategy,
                pct(m["cumulative_return"]),
                pct(m["max_drawdown"]),
                f'{m["sharpe"]:.2f}',
                pct(m["average_exposure"]),
                f'{m["beta"]:.2f}',
                number(m["estimated_cost"]),
            ]
        )
    add_table(
        doc,
        ["策略", "累计收益", "最大回撤", "夏普", "平均仓位", "Beta", "估算成本/元"],
        full_rows,
        [900, 1390, 1390, 900, 1390, 900, 2490],
        font_size=8.8,
    )
    add_body(
        doc,
        "仓位和成本数据解释了三套策略的净值差异。策略A平均仓位52.92%，产生451笔订单和约16.94万元估算成本，较高参与度没有转化为正收益。策略B平均仓位仅18.95%，Beta约0.16，低回撤主要来自长期现金缓冲。策略C以34.07%的平均仓位取得56.93%的扩展样本收益，单位市场暴露的收益表现更强；其518笔订单和约12.30万元估算成本也显示轮动信号需要承担较高摩擦。三种策略在持仓时均集中于单一ETF，并可能跨夜持有，跳空和单一标的集中风险依然存在。",
    )

    add_heading(doc, "七、成本上升会明显压缩策略收益", 1)
    add_picture(
        doc,
        "figure4_cost_stress.png",
        "图4：佣金率与滑点同时翻倍后的全期累计收益变化",
        width=5.8,
    )
    add_body(
        doc,
        "图4显示成本冲击会直接改变策略结论。佣金和滑点同时翻倍后，策略A的累计收益由-30.38%降至-53.61%，减少23.23个百分点；策略B由2.16%降至-9.35%，原有微弱收益被完全吞噬；策略C由56.93%降至25.49%，减少31.44个百分点，降幅超过基础收益的一半。策略C在压力情景下仍保持三者第一，说明其信号具有一定收益缓冲；同时，最大的绝对收益损失揭示轮动换手是其关键风险来源。策略B的结果则表明，接近零的历史优势对交易成本缺乏容错空间。",
    )

    doc.add_page_break()
    add_heading(doc, "八、测试周揭示统一空仓信号，模拟盘尚待首日数据", 1)
    add_heading(doc, "8.1 冻结规则回放复现了当周可获得的信息", 2)
    add_body(
        doc,
        "长期回测在2026年7月17日结束，随后使用已经冻结的代码和参数回放7月20日至24日五个交易日。三套策略只读取T−1日及更早的已完成数据，交易指令在下一交易日执行，测试过程中没有使用未来价格或根据结果修改参数。因此，这段回测能够复现策略在当时信息条件下会产生的信号和目标仓位，并为尚未开始运行的模拟盘提供启动前观察。",
    )
    short_rows = []
    for strategy in "ABC":
        item = short_results[strategy]
        short_rows.append(
            [
                strategy,
                pct(item["strategy_return"]),
                pct(item["benchmark_return"]),
                pct(item["excess_return"]),
                pct(item["max_drawdown"]),
                f'{item["profitable_trades"]}/{item["losing_trades"]}',
                "保持现金",
            ]
        )
    add_table(
        doc,
        ["策略", "策略收益", "基准收益", "超额收益", "最大回撤", "盈利/亏损次数", "汇总状态"],
        short_rows,
        [900, 1300, 1300, 1300, 1300, 1650, 1600],
        font_size=8.5,
    )
    add_body(
        doc,
        "表中最直接的信号是三套策略的净值均保持不变，交易次数均为0，五个交易日全部处于现金状态。策略A的MA20入场条件没有触发；策略B的MA10持续低于MA30；策略C的三只候选ETF动量得分均为负。不同规则在同一窗口给出了统一的风险规避判断，说明当时的中短期趋势强度没有达到预设入场标准。",
    )

    add_heading(doc, "8.2 空仓规避了波动，也暴露出趋势信号的滞后", 2)
    add_body(
        doc,
        "沪深300在测试周上涨2.65%，三套策略收益均为0%，平台计算的超额收益均为-2.58%。市场价格已经上涨，滞后趋势指标仍未确认入场，反映出规则对快速反转或短促上涨行情的响应较慢。三套策略由此避开了价格波动、集中持仓和交易成本，也错过了当周上涨；测试周的主要风险表现为机会成本。",
    )
    add_body(
        doc,
        "0%回撤来自全程空仓，只能说明当周没有市场暴露，尚不能证明止损、仓位上限或轮动机制已经在持仓状态下有效控制风险。零成交也意味着该窗口没有形成可供分析的成交价格、滑点、拒单或部分成交样本。五日结果支持对信号行为和机会成本的判断，对执行质量及长期稳定性仍缺乏证据。",
    )

    add_heading(doc, "8.3 模拟盘已部署，部署后的实盘模拟尚未开始", 2)
    add_body(
        doc,
        "三个策略已于2026年7月24日周五分别建立独立JoinQuant模拟盘。7月25日为周六、A股休市，部署后尚未经历新的交易日，因此截至报告日期没有实际模拟订单、成交、持仓或净值变化。这里的“实盘模拟”指平台在真实交易时序下运行虚拟资金，不涉及真实资金。上一周回测复现了可获得信息和策略决策，仍缺少平台实时调度、成交价格与订单状态等执行证据，两类结果在报告中分开记录。",
    )

    add_heading(doc, "8.4 看板将持续更新实际模拟结果", 2)
    add_body(
        doc,
        "持续观察页面当前展示历史回测、冻结规则下的行情回放、成本敏感性和模拟盘等待状态。下一交易日起，页面将继续补充三套模拟盘的实际净值、持仓、订单、成交和风险指标，并与历史回测分区展示。这样可以直接识别理论信号与平台实际执行之间的偏差，也能避免把测试周回放误写成部署后的模拟业绩。",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    lead = p.add_run("持续观察页面：")
    set_font(lead, bold=True)
    add_hyperlink(p, "TASK7三策略表现与风险观察", DASHBOARD_URL)
    add_body(
        doc,
        "看板公开内容不包含账户访问方式。实际模拟数据形成后，将按相同日期区间比较三套策略，并保持历史回测、行情回放和平台模拟记录三类数据相互独立。",
    )

    add_heading(doc, "九、结论、建议与实践反思", 1)
    add_heading(doc, "9.1 下一阶段以策略C为主、策略B为风险参照", 2)
    add_body(
        doc,
        "策略C应作为下一阶段的主要观察对象，因为它在长期回测中取得最高收益和唯一正超额，同时保持中等回撤；策略B应作为防守型参照，用于判断更低仓位是否能在真实成交环境中持续降低波动；策略A不适合作为主要方案，但保留它有助于衡量新增风控与轮动机制究竟创造了多少改进。",
    )
    add_body(
        doc,
        "后续判断不应只看累计收益。三策略对比需要同时覆盖最大回撤、平均仓位、市场Beta、换手、成本和空仓机会成本。若策略C的收益继续领先但成本和回撤明显放大，应优先降低调仓频率或仓位，避免再次追逐更短参数；若策略B长期保持低回撤但资金使用率过低，则应评估其作为组合防守层的价值，单独评价收益容易低估其风险缓冲作用。",
    )
    add_heading(doc, "9.2 本次实践的经验与教训", 2)
    add_body(
        doc,
        "第一，参数优化应重点考察跨阶段稳定性，单一区间的最高收益仅作参考；开发期与验证期表现方向一致，比某个窗口的峰值更有参考价值。第二，风险控制不会自动提高收益，策略B的低回撤主要来自较低仓位，必须同时评价资金使用效率。第三，测试周没有成交仍然提供了有效信号证据：三套规则在相同信息下保持现金，同时暴露了趋势判断滞后、上涨机会成本和缺少执行样本三项局限。",
    )
    add_body(
        doc,
        "最重要的教训是，长期回测、短窗口测试和模拟组合回答的是不同问题。长期回测用于识别收益与回撤结构，短窗口用于检查冻结规则在指定日期的行为，模拟组合则用于观察真实交易时序、成交和持续风险。只有保持参数冻结、统一比较区间并持续记录，后续三策略对比才具有解释力。",
    )

    add_heading(doc, "十、资料来源", 1)
    sources = [
        ("JoinQuant新手指南", "https://www.joinquant.com/guide"),
        ("JoinQuant API文档", "https://cdn.joinquant.com/help/img/JoinQuantAPI.pdf"),
        (
            "JoinQuant技术支持：避免未来数据",
            "https://www.joinquant.com/community/post/detailMobile?postId=23804",
        ),
        (
            "聚宽社区：宽基ETF动量轮动钝化RSRS择时",
            "https://www.joinquant.com/community/post/detailMobile?postId=37120",
        ),
    ]
    for label, url in sources:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.1
        add_hyperlink(p, label, url, font_size=9)

    add_heading(doc, "附录：聚宽回测凭证", 1)
    add_body(
        doc,
        "以下六张截图分别对应三套策略的长期回测结果，以及2026年7月20日至24日测试周回测结果。",
    )
    evidence = [
        (
            "Screenshot 2026-07-25 at 00.17.55.png",
            "附图1：策略A长期回测完成页（2019-01-01至2026-07-17）",
        ),
        (
            "Screenshot 2026-07-25 at 00.16.52.png",
            "附图2：策略B长期回测完成页（2019-01-01至2026-07-17）",
        ),
        (
            "Screenshot 2026-07-25 at 00.16.39.png",
            "附图3：策略C长期回测完成页（2019-01-01至2026-07-17）",
        ),
        (
            "Screenshot 2026-07-25 at 00.19.23.png",
            "附图4：策略A测试周回测完成页（2026-07-20至2026-07-24）",
        ),
        (
            "Screenshot 2026-07-25 at 00.19.11.png",
            "附图5：策略B测试周回测完成页（2026-07-20至2026-07-24）",
        ),
        (
            "Screenshot 2026-07-25 at 00.19.00.png",
            "附图6：策略C测试周回测完成页（2026-07-20至2026-07-24）",
        ),
    ]
    for idx, (filename, caption) in enumerate(evidence):
        if idx:
            doc.add_page_break()
        add_picture_path(
            doc,
            PLATFORM_EVIDENCE_DIR / filename,
            caption,
            width=5.8,
        )

    doc.core_properties.title = "TASK7 三策略研究与模拟跟踪报告"
    doc.core_properties.subject = "JoinQuant三策略表现、风险与持续观察"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.save(OUTPUT)
    ensure_font_table_entry(OUTPUT)
    print(f"✅ Wrote {OUTPUT}")


if __name__ == "__main__":
    write_report()
