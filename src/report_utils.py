#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
报告生成共享模块 — python-docx 格式化辅助函数
供 Task1/Task2/.../TaskN 的 generate_report.py 复用

用法:
    from src.report_utils import (
        set_cjk_font, add_paragraph, add_heading_styled,
        add_picture_captioned, add_table,
    )
"""

import os
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ═══════════════════════════════════════════════════════════════
# 常量
# ═══════════════════════════════════════════════════════════════

FONT_NAME = '宋体'
FONT_SIZE = Pt(10.5)  # 五号


# ═══════════════════════════════════════════════════════════════
# CJK 字体 (通过 XML 操作)
# ═══════════════════════════════════════════════════════════════

def set_cjk_font(run, font_name=FONT_NAME):
    """为 run 设置中文字体"""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    run.font.name = font_name


# ═══════════════════════════════════════════════════════════════
# 段落
# ═══════════════════════════════════════════════════════════════

def add_paragraph(doc, text, bold=False, size=FONT_SIZE,
                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  space_before=Pt(0), space_after=Pt(0),
                  line_spacing=1.5, font_name=FONT_NAME,
                  color=None, first_line_indent=None):
    """添加格式化段落（宋体，1.5倍行距，0段间距，两端对齐）"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = space_before
    pf.space_after = space_after
    pf.line_spacing = line_spacing
    if first_line_indent:
        pf.first_line_indent = first_line_indent

    run = p.add_run(text)
    run.font.size = size
    run.bold = bold
    set_cjk_font(run, font_name)
    if color:
        run.font.color.rgb = color
    return p


# ═══════════════════════════════════════════════════════════════
# 标题
# ═══════════════════════════════════════════════════════════════

def add_heading_styled(doc, text, level=1):
    """添加标题 — Normal 段落 + 不同字号 Bold（不使用 Word Heading 样式）
    level=1: 14pt Bold（章标题）
    level=2: 12pt Bold（节标题）
    """
    if level == 1:
        return add_paragraph(doc, text, bold=True, size=Pt(14),
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            space_before=Pt(12), space_after=Pt(6))
    else:
        return add_paragraph(doc, text, bold=True, size=Pt(12),
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            space_before=Pt(8), space_after=Pt(4))


# ═══════════════════════════════════════════════════════════════
# 图片
# ═══════════════════════════════════════════════════════════════

def add_picture_captioned(doc, img_path, caption, width_inches=6.0):
    """添加带编号和图题的图片"""
    if not os.path.exists(img_path):
        add_paragraph(doc, f'[图表缺失: {os.path.basename(img_path)}]',
                     color=RGBColor(0xCC, 0x00, 0x00),
                     alignment=WD_ALIGN_PARAGRAPH.CENTER)
        return

    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_img.add_run()
    run.add_picture(img_path, width=Inches(width_inches))

    add_paragraph(doc, caption, bold=False, size=Pt(9),
                 alignment=WD_ALIGN_PARAGRAPH.CENTER,
                 space_after=Pt(12), font_name='宋体')


# ═══════════════════════════════════════════════════════════════
# 表格
# ═══════════════════════════════════════════════════════════════

def add_table(doc, headers, rows, col_widths=None):
    """添加格式化表格（Table Grid 样式，灰色表头，9pt 宋体）"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers),
                         style='Table Grid')
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 表头
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        set_cjk_font(run, '宋体')
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D9D9D9')
        cell._element.get_or_add_tcPr().append(shading)

    # 数据行
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            set_cjk_font(run, '宋体')

    doc.add_paragraph()  # 表后空行
    return table
