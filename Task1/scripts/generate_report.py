#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成 Task1 .docx 报告
格式: 宋体五号(10.5pt), 1.5倍行距, 0段间距, 两端对齐
"""

import os
import sys
import pandas as pd
import numpy as np
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '..'))
from src.report_utils import (
    set_cjk_font, add_paragraph, add_heading_styled,
    add_picture_captioned, add_table, FONT_SIZE,
)

# ═══════════════════════════════════════════════════════════════
# 路径
# ═══════════════════════════════════════════════════════════════

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TASK_DIR = os.path.join(BASE_DIR, '..')
ROOT_DIR = os.path.join(BASE_DIR, '..', '..')
CHART_DIR = os.path.join(ROOT_DIR, 'data', 'charts', 'task1')
CSV_DIR = os.path.join(ROOT_DIR, 'data', 'csv')
OUTPUT_FILE = os.path.join(TASK_DIR, 'Rebecca+Task1.docx')

STOCK_CODE = "000001"
STOCK_NAME = "平安银行"
MARKET = "A股"
CHART_CLOSE = os.path.join(CHART_DIR, f"{STOCK_CODE}_{STOCK_NAME}_close_price.png")
CHART_CANDLE = os.path.join(CHART_DIR, f"{STOCK_CODE}_{STOCK_NAME}_candle.png")
CSV_FILE = os.path.join(CSV_DIR, f"{STOCK_CODE}_{STOCK_NAME}_{MARKET}_daily.csv")
MULTI_CHART = os.path.join(CHART_DIR, "multi_stock_compare.png")
AH_CHART = os.path.join(CHART_DIR, "AH_comparison.png")

# 多股汇总与AH汇总
MULTI_CSV = os.path.join(CSV_DIR, "multi_stock_summary.csv")
AH_CSV = os.path.join(CSV_DIR, "AH_summary.csv")


# ═══════════════════════════════════════════════════════════════
# 本地辅助函数
# ═══════════════════════════════════════════════════════════════

def add_code_block(doc, code_text):
    """添加代码块（Courier New 9pt，浅灰背景）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing = 1.2
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    set_cjk_font(run, 'Courier New')
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
    p._element.pPr.append(shading)
    return p


def _read_csv_safe(path):
    """读取 CSV，自动处理 BOM 和编码"""
    if not os.path.exists(path):
        return None
    return pd.read_csv(path, encoding='utf-8-sig')


def _col(df, *candidates):
    """返回第一个存在的列名"""
    for c in candidates:
        if c in df.columns:
            return c
    return candidates[0]


# ═══════════════════════════════════════════════════════════════
# 章节
# ═══════════════════════════════════════════════════════════════

def write_intro(doc):
    """引言"""
    add_heading_styled(doc, '引言', level=1)
    add_paragraph(doc,
        '量化交易是指利用计算机程序、数学模型和统计分析工具，自动化地执行交易策略的一种交易方式。'
        '随着大数据和人工智能技术的发展，量化交易已成为现代金融市场中不可或缺的组成部分。'
        '本报告是北京大学光华管理学院"量化交易：AI大模型辅助的金融交易策略"课程的第一项任务成果，'
        '围绕以下内容展开：（1）量化交易相较于传统手工操作的优势分析；'
        '（2）K线、基本面分析和技术分析三个核心概念的详细解释；'
        '（3）基于AKShare/Tushare Pro双数据源的Python编程实践，涵盖单股数据获取、多股跨行业对比、'
        'A股+港股跨市场K线分析，以及自包含交互式看板的构建。',
        first_line_indent=Cm(0.74))


def write_chapter1_advantages(doc):
    """第一章: 量化交易的优势"""
    add_heading_styled(doc, '一、量化交易相较于传统手工操作交易的优势', level=1)

    add_heading_styled(doc, '1.1 处理更多数据', level=2)
    add_paragraph(doc,
        '传统手工交易受限于人的精力和认知能力，交易员通常只能同时关注少数几个市场或品种。'
        '量化交易系统可以借助计算机的强大算力，在毫秒级别内处理海量市场数据——涵盖数百只股票、'
        '期货合约、外汇货币对等多市场、多品种的实时行情数据。此外，量化系统还可以融合宏观经济'
        '指标、财务报表、新闻舆情、社交媒体情绪等非结构化数据，从而构建更全面、多维度的投资决策依据。'
        '这种数据处理的广度和深度，是人工交易难以企及的。',
        first_line_indent=Cm(0.74))

    add_heading_styled(doc, '1.2 执行更一致', level=2)
    add_paragraph(doc,
        '人类交易者在决策过程中不可避免地受到情绪因素（如恐惧、贪婪、过度自信、损失厌恶等）的影响，'
        '导致在实际操作中偏离既定策略。量化交易则将交易规则固化为程序代码，'
        '计算机严格执行每一条买卖指令，不受任何情绪干扰。无论在何种市场环境下，'
        '量化策略都能保持纪律性和一致性，确保交易行为与策略设计完全吻合。',
        first_line_indent=Cm(0.74))

    add_heading_styled(doc, '1.3 能历史检验（回测）', level=2)
    add_paragraph(doc,
        '量化交易策略在投入实盘之前，可以利用历史数据进行回测（Backtesting），'
        '即在历史行情中模拟策略的运行效果。通过回测，交易者可以评估策略的收益率、'
        '最大回撤、夏普比率、胜率、盈亏比等关键绩效指标，验证策略在不同市场环境下的表现。'
        '这种先验证、后执行的模式，使得量化策略在进入实战前就经过了严格的统计检验，'
        '大大降低了策略失效的风险。',
        first_line_indent=Cm(0.74))

    add_heading_styled(doc, '1.4 风险规则可写入系统', level=2)
    add_paragraph(doc,
        '风险管理是交易中至关重要的一环。量化交易可以将复杂的风险控制规则直接编码进交易系统：'
        '单笔交易最大亏损限制、每日最大回撤熔断机制、投资组合层面的风险敞口控制、'
        '波动率自适应仓位管理等。一旦市场条件触发预设的风控阈值，系统将自动减仓或停止交易，'
        '无需人工干预，有效避免了人工风控中可能出现的犹豫、侥幸心理或操作延迟。',
        first_line_indent=Cm(0.74))

    add_paragraph(doc,
        '综合以上四个维度，量化交易通过数据驱动、纪律执行、历史验证和系统化风控，'
        '弥补了传统手工交易在信息处理能力、执行一致性和风险管理效率方面的不足。',
        first_line_indent=Cm(0.74))

    add_table(doc,
        ['维度', '传统手工交易', '量化交易'],
        [
            ['数据处理', '关注少数品种，信息维度有限', '毫秒级处理数百品种，融合多源异构数据'],
            ['执行纪律', '受情绪影响，易偏离策略', '程序严格执行，不受情绪干扰'],
            ['策略验证', '依赖主观经验，缺乏客观评估', '历史回测，统计检验，量化绩效指标'],
            ['风险管理', '人工判断，反应滞后', '程序化风控，自动熔断，实时响应'],
        ])


def write_chapter2_concepts(doc):
    """第二章: 基本概念解释"""
    add_heading_styled(doc, '二、基本概念解释', level=1)

    # 2.1 K线
    add_heading_styled(doc, '2.1 K线（K-Line / Candlestick）', level=2)
    add_paragraph(doc,
        'K线，又称蜡烛图（Candlestick Chart）或日本线，起源于18世纪日本的大米期货市场，'
        '由日本商人本间宗久发明，是目前金融市场中最常用的价格图表形式之一。',
        first_line_indent=Cm(0.74))

    add_paragraph(doc,
        '一根标准的K线由四个价格要素构成：开盘价（Open）、收盘价（Close）、'
        '最高价（High）和最低价（Low）。K线的实体部分表示开盘价与收盘价之间的价格区间：'
        '当收盘价高于开盘价时，实体通常用红色或空心绘制，称为阳线，表示价格上涨；'
        '当收盘价低于开盘价时，实体通常用绿色或实心绘制，称为阴线，表示价格下跌。'
        '实体上方的细线称为上影线，其顶端代表当日最高价；实体下方的细线称为下影线，'
        '其底端代表当日最低价。',
        first_line_indent=Cm(0.74))

    add_paragraph(doc,
        'K线图不仅可以展示单一周期的价格信息，更重要的价值在于通过多根K线的组合形态'
        '（如锤子线、吞没形态、十字星、启明星等）来研判市场多空力量对比和潜在的趋势转折。'
        '常见K线周期包括日K线、周K线、月K线，以及5分钟、15分钟、60分钟等短期K线。',
        first_line_indent=Cm(0.74))

    add_table(doc,
        ['要素', '英文', '含义'],
        [
            ['开盘价', 'Open', '该周期第一笔成交价'],
            ['收盘价', 'Close', '该周期最后一笔成交价'],
            ['最高价', 'High', '该周期最高成交价'],
            ['最低价', 'Low', '该周期最低成交价'],
            ['阳线', 'Bullish Candle', '收盘价 > 开盘价，通常红色/空心，表示上涨'],
            ['阴线', 'Bearish Candle', '收盘价 < 开盘价，通常绿色/实心，表示下跌'],
            ['上影线', 'Upper Shadow', '实体上方细线，顶端 = 最高价'],
            ['下影线', 'Lower Shadow', '实体下方细线，底端 = 最低价'],
        ])

    # 2.2 基本面
    add_heading_styled(doc, '2.2 基本面分析（Fundamental Analysis）', level=2)
    add_paragraph(doc,
        '基本面分析是一种通过评估经济、行业和公司层面的内在价值来判断证券投资价值的方法论。'
        '其核心思想是：每一只证券都有一个内在价值，市场价格围绕内在价值波动；'
        '当市场价格低于内在价值时，是买入时机；当市场价格高于内在价值时，则是卖出时机。',
        first_line_indent=Cm(0.74))

    add_paragraph(doc,
        '基本面分析通常涵盖三个层次：（1）宏观经济分析——考察GDP增长率、通货膨胀率、'
        '利率水平、货币政策、财政政策、就业数据等宏观经济指标，判断整体经济周期所处的阶段；'
        '（2）行业分析——研究行业生命周期、竞争格局、监管政策、技术变革趋势等，'
        '评估特定行业的发展前景和投资吸引力；'
        '（3）公司分析——深入考察企业的财务报表（资产负债表、利润表、现金流量表），'
        '计算市盈率（P/E）、市净率（P/B）、净资产收益率（ROE）、负债率、自由现金流等财务指标，'
        '评估公司的盈利能力、成长性、偿债能力和运营效率。',
        first_line_indent=Cm(0.74))

    add_paragraph(doc,
        '基本面分析的代表人物包括本杰明·格雷厄姆（Benjamin Graham）和沃伦·巴菲特（Warren Buffett）。'
        '基本面分析更适合中长期投资者，其优势在于能够发现被市场低估的优质资产，'
        '但缺点是信息获取和分析过程耗时较长，且市场可能长期偏离基本面判断。',
        first_line_indent=Cm(0.74))

    add_table(doc,
        ['分析层次', '考察内容', '典型指标'],
        [
            ['宏观经济', '经济周期、货币政策、财政政策', 'GDP增长率、CPI、利率、PMI、失业率'],
            ['行业分析', '行业生命周期、竞争格局、监管环境', '市场规模、集中度、技术变革趋势'],
            ['公司分析', '财务报表、盈利能力、成长性', 'P/E、P/B、ROE、负债率、自由现金流'],
        ])

    # 2.3 技术面
    add_heading_styled(doc, '2.3 技术分析（Technical Analysis）', level=2)
    add_paragraph(doc,
        '技术分析是一种通过研究市场交易数据（主要是价格和成交量）来预测未来价格走势的分析方法。'
        '技术分析建立在三个基本假设之上：（1）市场行为包容消化一切信息——所有影响价格的因素'
        '（基本面、政策、心理等）都已反映在价格中；（2）价格以趋势方式演变——价格运动具有惯性，'
        '一旦形成趋势，更可能继续而非反转；（3）历史会重演——人性不变，'
        '过去有效的价格形态和技术指标在未来仍将发挥作用。',
        first_line_indent=Cm(0.74))

    add_paragraph(doc,
        '技术分析的工具和方法主要包括：'
        '（1）图表形态分析——识别头肩顶/底、双重顶/底、三角形整理、旗形等经典价格形态；'
        '（2）技术指标——包括趋势跟踪指标（如移动平均线MA、MACD）、'
        '震荡指标（如RSI相对强弱指数、KDJ随机指标）、'
        '成交量指标（如VWAP、OBV）等；'
        '（3）道氏理论、波浪理论和江恩理论等经典分析框架。',
        first_line_indent=Cm(0.74))

    add_paragraph(doc,
        '技术分析与基本面分析并非互斥，而是互补的。很多成功的投资者会将两者结合：'
        '用基本面分析筛选投资标的（决定买什么），用技术分析选择入场和出场时机'
        '（决定何时买卖）。在量化交易领域，技术指标往往被量化为具体的数学模型，'
        '成为交易策略的核心信号来源。',
        first_line_indent=Cm(0.74))

    add_table(doc,
        ['维度', '基本面分析', '技术分析'],
        [
            ['核心问题', '买什么——筛选投资标的', '何时买卖——选择交易时机'],
            ['数据来源', '财务报表、宏观经济指标、行业报告', '价格、成交量、技术指标'],
            ['时间框架', '中长期（月/年）', '短期至中期（日/周/月）'],
            ['理论基础', '内在价值回归', '市场行为包容一切、趋势延续、历史重演'],
            ['代表人物', '格雷厄姆、巴菲特', '道氏、江恩、Wilder'],
            ['局限', '市场可能长期偏离基本面', '强趋势中指标可能钝化'],
        ])


def write_chapter3_practice(doc):
    """第三章: 数据获取与Python编程实践"""
    add_heading_styled(doc, '三、数据获取与Python编程实践', level=1)

    # 3.1 数据源
    add_heading_styled(doc, '3.1 数据源架构与平台注册', level=2)
    add_paragraph(doc,
        '本任务采用双数据源架构：主力使用AKShare（免费、无需注册、覆盖A股+港股的金融数据接口），'
        '补充使用Tushare Pro（需注册Token，数据质量高但对港股有频率限制）。'
        '两者均支持前复权（qfq）价格——以当前股价为基准向前调整历史价格，适合技术分析的连续性要求。',
        first_line_indent=Cm(0.74))

    add_paragraph(doc,
        'Tushare Pro（https://www.tushare.pro/）是国内知名的免费金融数据开放平台，'
        '提供股票、基金、期货、期权、指数、宏观经济等多类金融数据接口。注册步骤：'
        '访问官网→填写手机号/邮箱/密码完成注册→登录后在"个人主页→Token管理"复制API Token→'
        '在Python代码中通过ts.set_token()配置后即可调用。新注册用户拥有基础积分，'
        '可免费访问大部分常用行情数据接口。',
        first_line_indent=Cm(0.74))

    # 3.2 单股数据获取
    add_heading_styled(doc, '3.2 单股数据获取——以平安银行（000001）为例', level=2)
    add_paragraph(doc,
        '以下Python程序以平安银行（000001.SZ）为示例，展示了三个核心功能：'
        '（1）通过AKShare获取过去一年（约241个交易日）的前复权日线数据；'
        '（2）绘制含MA20/MA60移动平均线的收盘价曲线图；'
        '（3）将数据保存为CSV格式文件。AKShare不可用时自动回退到Tushare Pro。',
        first_line_indent=Cm(0.74))

    code_snippet = '''import akshare as ak
import pandas as pd
import matplotlib.pyplot as plt

# 1. 通过 AKShare 获取前复权日线数据
df = ak.stock_zh_a_hist(
    symbol="000001",              # 平安银行
    period="daily",
    start_date="20250628",
    end_date="20260628",
    adjust="qfq"                  # 前复权
)

# 2. 绘制收盘价曲线图（含 MA20/MA60）
df["MA20"] = df["收盘"].rolling(20).mean()
df["MA60"] = df["收盘"].rolling(60).mean()

plt.figure(figsize=(14, 7))
plt.plot(df["日期"], df["收盘"],
         color="#1f77b4", linewidth=1.5, label="每日收盘价")
plt.plot(df["日期"], df["MA20"],
         color="orange", linewidth=1.0, linestyle="--", label="MA20")
plt.plot(df["日期"], df["MA60"],
         color="green", linewidth=1.0, linestyle="--", label="MA60")
plt.title("平安银行（000001）每日收盘价走势图")
plt.xlabel("日期"); plt.ylabel("收盘价（元）")
plt.legend(); plt.grid(True, alpha=0.3)
plt.savefig("close_price.png", dpi=150)

# 3. 保存为 CSV 文件
df.to_csv("000001_平安银行_A股_daily.csv",
          index=False, encoding="utf-8-sig")
print("数据已保存，共", len(df), "条记录")'''

    add_code_block(doc, code_snippet)

    add_paragraph(doc,
        f'完整脚本见 Task1/scripts/update_data.py。用户修改 STOCK_CODE 和 STOCK_NAME '
        '两个变量即可获取任意沪深A股的历史交易数据。',
        first_line_indent=Cm(0.74))

    # 3.2.1 平安银行走势图
    add_heading_styled(doc, '3.2.1 平安银行收盘价走势', level=2)

    add_picture_captioned(doc, CHART_CLOSE,
        f'图1：{STOCK_NAME}（{STOCK_CODE}）过去一年每日收盘价及MA20/MA60走势图',
        width_inches=5.3)

    if os.path.exists(CSV_FILE):
        df = _read_csv_safe(CSV_FILE)
        if df is not None:
            dc = _col(df, 'Date', '日期')
            close_c = _col(df, 'Close', '收盘')
            high_c = _col(df, 'High', '最高')
            low_c = _col(df, 'Low', '最低')
            vol_c = _col(df, 'Volume', '成交量')
            amount_c = _col(df, 'Amount', '成交额')

            df[dc] = pd.to_datetime(df[dc])
            close = df[close_c]
            high = df[high_c]
            low = df[low_c]
            vol = df[vol_c]
            amount = df[amount_c]

            start_date = df[dc].min().strftime('%Y年%m月%d日')
            end_date = df[dc].max().strftime('%Y年%m月%d日')
            days = len(df)
            change_pct = (close.iloc[-1] - close.iloc[0]) / close.iloc[0] * 100
            trend = "上涨" if change_pct > 0 else "下跌"
            max_close = close.max()
            min_close = close.min()
            max_date = df.loc[close.idxmax(), dc].strftime('%Y年%m月%d日')
            min_date = df.loc[close.idxmin(), dc].strftime('%Y年%m月%d日')

            add_paragraph(doc,
                f'图1展示了{STOCK_NAME}（{STOCK_CODE}）在{start_date}至{end_date}期间'
                f'（共{days}个交易日）的每日收盘价走势。整体价格呈{trend}趋势，'
                f'期初收盘价{close.iloc[0]:.2f}元，期末{close.iloc[-1]:.2f}元，'
                f'区间累计涨跌幅{change_pct:+.2f}%。最高收盘价出现于{max_date}（{max_close:.2f}元），'
                f'最低出现于{min_date}（{min_close:.2f}元）。',
                first_line_indent=Cm(0.74))

            add_table(doc,
                ['指标', '数值', '说明'],
                [
                    ['数据区间', f'{start_date} 至 {end_date}', f'共{days}个交易日'],
                    ['期初收盘价', f'{close.iloc[0]:.2f} 元', f'{start_date} 收盘价'],
                    ['期末收盘价', f'{close.iloc[-1]:.2f} 元', f'{end_date} 收盘价'],
                    ['区间涨跌幅', f'{change_pct:+.2f}%', f'整体呈{trend}趋势'],
                    ['最高收盘价', f'{max_close:.2f} 元', f'出现在{max_date}'],
                    ['最低收盘价', f'{min_close:.2f} 元', f'出现在{min_date}'],
                    ['收盘价均值', f'{close.mean():.2f} 元', '241个交易日的算术平均'],
                    ['收盘价标准差', f'{close.std():.2f} 元', '衡量价格离散程度'],
                ])

    # 3.2.2 K线图
    add_heading_styled(doc, '3.2.2 平安银行K线图', level=2)
    add_paragraph(doc,
        '与单纯收盘价曲线相比，K线图在一张图中同时呈现开盘价、收盘价、最高价和最低价四个维度的信息，'
        '能够更全面地反映每日多空博弈的过程。下图为平安银行过去一年的日K线图（含MA20/MA60均线叠加）。',
        first_line_indent=Cm(0.74))

    add_picture_captioned(doc, CHART_CANDLE,
        f'图2：{STOCK_NAME}（{STOCK_CODE}）日K线图（含MA20/MA60）',
        width_inches=5.3)

    # 3.3 多股对比
    add_heading_styled(doc, '3.3 多股跨行业对比分析', level=2)
    add_paragraph(doc,
        '单只股票的分析只能反映个别公司的表现。为了更全面地理解A股市场的结构性特征，'
        '本任务选取了5只分别代表银行、白酒消费、新能源汽车、保险和新能源电池五个不同行业的'
        '沪深A股标的，进行了跨行业对比分析。',
        first_line_indent=Cm(0.74))

    multi = _read_csv_safe(MULTI_CSV)
    if multi is not None and len(multi) >= 5:
        rows = []
        for _, r in multi.iterrows():
            rows.append([
                f'{r["股票名称"]}（{r["代码"]}）',
                r['行业'],
                f'{r["区间涨跌幅(%)"]:+.2f}%',
                f'{r["年化波动率(%)"]:.2f}%',
                f'{r["最大回撤(%)"]:.2f}%',
                f'{r["日均成交量(万手)"]:.0f}',
            ])
        add_table(doc,
            ['股票（代码）', '行业', '区间涨跌幅', '年化波动率', '最大回撤', '日均成交量(万手)'],
            rows)

        # Analysis
        chg_col = multi['区间涨跌幅(%)']
        best = multi.loc[chg_col.idxmax()]
        worst = multi.loc[chg_col.idxmin()]
        vol_col = multi['年化波动率(%)']
        most_vol = multi.loc[vol_col.idxmax()]
        least_vol = multi.loc[vol_col.idxmin()]

        add_paragraph(doc,
            f'从表中可以看出五只股票在过去一年中呈现显著分化。'
            f'涨幅最大的是{best["股票名称"]}（{best["区间涨跌幅(%)"]:+.2f}%），'
            f'跌幅最大的是{worst["股票名称"]}（{worst["区间涨跌幅(%)"]:+.2f}%）。'
            f'波动率方面，{most_vol["股票名称"]}的年化波动率最高（{most_vol["年化波动率(%)"]:.2f}%），'
            f'{least_vol["股票名称"]}最低（{least_vol["年化波动率(%)"]:.2f}%），'
            f'两者相差超过一倍，反映出不同行业板块的风险特征差异显著。'
            f'金融类股票（银行、保险）波动率相对较低，而新能源和高科技相关标的波动更大。'
            f'这种行业间差异正是投资组合分散化配置的理论基础——通过配置低相关性资产来降低整体组合风险。',
            first_line_indent=Cm(0.74))

    add_picture_captioned(doc, MULTI_CHART,
        '图3：5只A股跨行业对比——标准化累计收益率曲线',
        width_inches=5.3)

    # 3.4 A+H跨市场
    add_heading_styled(doc, '3.4 A股+港股跨市场K线分析', level=2)
    add_paragraph(doc,
        '为进一步拓展分析维度，本任务还获取了5只港股权重标的（恒生成分股），与5只A股一起构成'
        '覆盖沪深和香港两大交易市场的10只标的池。港股选取了腾讯控股、阿里巴巴、美团、'
        '香港交易所和中国移动，涵盖了互联网、电商、本地生活、交易所运营和电信等不同行业。',
        first_line_indent=Cm(0.74))

    ah = _read_csv_safe(AH_CSV)
    if ah is not None:
        ah_rows = []
        for _, r in ah.iterrows():
            ah_rows.append([
                f'{r["名称"]}（{r["市场"]}）',
                f'{r["交易日"]}天',
                f'{r["期初"]:.2f}',
                f'{r["期末"]:.2f}',
                f'{r["涨跌幅(%)"]:+.2f}%',
                f'{r["年化波动(%)"]:.2f}%',
                f'{r["最大回撤(%)"]:.2f}%',
            ])
        add_table(doc,
            ['标的（市场）', '交易日', '期初价', '期末价', '涨跌幅', '年化波动率', '最大回撤'],
            ah_rows)

        a_share = ah[ah['市场'] == 'A股']
        hk_share = ah[ah['市场'] == '港股']
        a_avg_chg = a_share['涨跌幅(%)'].mean()
        hk_avg_chg = hk_share['涨跌幅(%)'].mean()
        a_avg_vol = a_share['年化波动(%)'].mean()
        hk_avg_vol = hk_share['年化波动(%)'].mean()

        add_paragraph(doc,
            f'跨市场对比显示，A股5标的平均涨跌幅为{a_avg_chg:+.2f}%，平均年化波动率{a_avg_vol:.2f}%；'
            f'港股5标的平均涨跌幅为{hk_avg_chg:+.2f}%，平均年化波动率{hk_avg_vol:.2f}%。'
            f'港股整体波动率更高，这与港股市场没有涨跌停板限制、机构投资者占比高、'
            f'受国际市场资金流动影响更大等结构性特征有关。'
            f'A股方面，宁德时代以+55.95%的涨幅和39.13%的年化波动率表现最为突出，体现了新能源板块的高成长高波动特性。',
            first_line_indent=Cm(0.74))

    add_picture_captioned(doc, AH_CHART,
        '图4：5只A股+5只港股K线矩阵对比（前复权价格）',
        width_inches=5.3)

    # 3.5 CSV
    add_heading_styled(doc, '3.5 数据保存与后续使用', level=2)
    add_paragraph(doc,
        f'所有10只标的的交易数据已保存为CSV格式文件，存放于 data/csv/ 目录。'
        f'CSV（Comma-Separated Values）是一种通用、轻量级的数据交换格式，'
        f'可被Excel、Python（pandas）、R、MATLAB等几乎所有数据分析工具直接读取。'
        f'每条记录包含日期(Date)、股票代码、开盘价(Open)、最高价(High)、最低价(Low)、'
        f'收盘价(Close)、成交量(Volume)、成交额(Amount)、涨跌幅(PctChg)、换手率(Turnover)等字段。',
        first_line_indent=Cm(0.74))

    add_paragraph(doc,
        '这些数据为后续任务提供了坚实的基础数据支撑，可直接用于：'
        '计算各类技术指标（如MACD、RSI、布林带等）；'
        '构建和回测交易策略（如均线交叉策略、动量策略、均值回归策略等）；'
        '进行风险管理分析（如计算VaR、最大回撤优化等）；'
        '开发机器学习预测模型（如LSTM股价预测、随机森林涨跌分类等）。',
        first_line_indent=Cm(0.74))


def write_chapter4_dashboard(doc):
    """第四章: 交互看板"""
    add_heading_styled(doc, '四、交互式数据看板', level=1)

    add_heading_styled(doc, '4.1 看板概述', level=2)
    add_paragraph(doc,
        '为直观展示和分析10只标的的行情数据，本任务构建了一个自包含的交互式Web看板'
        '（Task1/dashboard/index.html），无需后端服务器即可在浏览器中直接运行，'
        '并已部署到GitHub Pages供在线访问。看板采用ECharts 5.5渲染，所有CSV数据序列化为JSON嵌入页面，'
        '实现零依赖、零构建的纯前端解决方案。',
        first_line_indent=Cm(0.74))

    add_heading_styled(doc, '4.2 看板功能', level=2)

    add_table(doc,
        ['功能模块', '说明', '技术实现'],
        [
            ['标的切换', '左侧边栏按A股/港股分组，支持下拉选择和快捷按钮两种方式切换标的',
             'JavaScript动态数据绑定'],
            ['交互K线图', 'ECharts蜡烛图 + MA20/MA60均线叠加 + 成交量柱状图，支持悬停十字光标查看OHLC详情',
             'ECharts candlestick + line + bar 组合图'],
            ['日收益分布', '直方图展示每日收益率的频率分布，标注均值/标准差/偏度',
             'ECharts bar chart + 统计计算'],
            ['累计收益曲线', '展示区间累计收益率走势，正收益区域填充红色，负收益填充绿色',
             'ECharts line + area chart'],
            ['关键指标', '实时计算并展示区间涨跌幅、年化波动率、最大回撤、胜率四项核心指标',
             'JavaScript统计函数'],
            ['数据导出', '一键将当前标的数据下载为CSV文件',
             'Blob + URL.createObjectURL'],
            ['响应式布局', '桌面端左侧边栏+右侧图表双栏布局，移动端自动切换为上下堆叠',
             'CSS Flexbox + @media query'],
        ])

    add_heading_styled(doc, '4.3 技术架构', level=2)
    add_paragraph(doc,
        '看板的整体技术架构分为三层：数据层——10只标的的CSV数据通过Python脚本（build_dashboard.py）'
        '读取并序列化为JSON对象，直接嵌入HTML的<script>标签中；渲染层——ECharts 5.5库'
        '（通过CDN加载）渲染K线图、直方图和累计收益曲线三张图表；'
        '交互层——纯JavaScript实现标的切换、指标计算、数据导出和图表自适应缩放。'
        '所有图表共享同一个ECharts实例管理器，切换标的时自动释放旧实例并创建新实例，避免内存泄漏。',
        first_line_indent=Cm(0.74))

    add_paragraph(doc,
        '此外，Task1还包含一个Flask后端版本（dashboard/server.py），通过REST API实时从AKShare'
        '拉取最新数据，适用于需要盘中实时更新的场景。静态HTML版本和Flask版本共享同一套前端代码，'
        '仅在数据注入方式上有所不同（嵌入JSON vs. API请求）。',
        first_line_indent=Cm(0.74))

    add_heading_styled(doc, '4.4 部署与访问', level=2)
    add_paragraph(doc,
        '静态看板已部署到GitHub Pages：https://rebekahlllimx.github.io/quant-trading-ai/Task1/dashboard/index.html。'
        '本地使用可直接双击打开HTML文件（file://协议兼容），或运行python server.py启动Flask后端（端口8765）。'
        'GitHub Pages自动同步main分支的最新提交，每次push后自动部署。',
        first_line_indent=Cm(0.74))


def write_chapter5_conclusion(doc):
    """总结"""
    add_heading_styled(doc, '五、总结', level=1)
    add_paragraph(doc,
        '本报告完成了量化交易基础阶段的多项核心任务：'
        '第一，从数据处理能力、执行纪律性、历史回测验证和系统化风控四个维度，'
        '系统阐述了量化交易相较于传统手工交易的核心优势；'
        '第二，对K线、基本面分析和技术分析三个基础概念进行了详细解释，'
        '包括定义、构成要素、分析方法和应用场景；'
        '第三，基于AKShare/Tushare Pro双数据源架构，完成了10只标的（5只A股+5只港股）的数据获取，'
        '覆盖银行、白酒、新能源、保险、互联网、电商等多个行业，构建了从数据获取到可视化的完整数据管道；'
        '第四，通过多股跨行业对比和A+H跨市场分析，展示了不同行业板块和不同市场之间的结构性差异，'
        '为后续的量化策略开发提供了多维度、高质量的数据基础；'
        '第五，构建了自包含交互式看板和Flask后端两种部署方案，'
        '兼顾了GitHub Pages的零成本静态部署和实时数据更新的灵活性。',
        first_line_indent=Cm(0.74))

    add_paragraph(doc,
        f'所有代码、数据和报告均已托管至GitHub仓库（https://github.com/RebekahLllimx/quant-trading-ai），'
        f'遵循统一的目录结构和SOP工作流规范。数据统一使用前复权（qfq）价格，确保技术分析的连续性；'
        f'所有脚本遵循命名规范（update_data.py / plot_*.py / build_dashboard.py / generate_report.py）；'
        f'共享模块（src/indicators.py / src/report_utils.py）跨任务复用，避免代码重复。',
        first_line_indent=Cm(0.74))

    add_paragraph(doc,
        f'（报告生成日期：{datetime.now().strftime("%Y年%m月%d日")}）',
        first_line_indent=Cm(0.74))


# ═══════════════════════════════════════════════════════════════
# 构建文档
# ═══════════════════════════════════════════════════════════════

def build_document():
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # 默认样式
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = FONT_SIZE
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 各章节
    write_intro(doc)
    write_chapter1_advantages(doc)
    write_chapter2_concepts(doc)
    write_chapter3_practice(doc)
    write_chapter4_dashboard(doc)
    write_chapter5_conclusion(doc)

    # 保存
    doc.save(OUTPUT_FILE)
    print(f'文档已保存至: {OUTPUT_FILE}')
    print('完成！')


if __name__ == '__main__':
    build_document()
