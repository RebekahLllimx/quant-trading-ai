#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成 TASK8 中文专业学习报告 Word 文档。"""

from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
TASK8 = ROOT / "Task8"
FIG_DIR = ROOT / "artifacts" / "charts" / "task8"
OUT = TASK8 / "从数据到执行_量化交易策略与机器学习应用综合实践报告.docx"

NAVY = "000000"
ORANGE = "E07A3F"
GREEN = "2E7D65"
GRAY = "666666"
LIGHT_BLUE = "F2F2F2"
LIGHT_GRAY = "F7F7F7"
MID_GRAY = "D9D9D9"
RED = "A94442"
BODY_FONT = "STSong"
HEADING_FONT = "STSong"
LATIN_FONT = "Times New Roman"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        edge_data = kwargs[edge]
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key in ["val", "sz", "space", "color"]:
            if key in edge_data:
                element.set(qn(f"w:{key}"), str(edge_data[key]))


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, cn=BODY_FONT, size=10.5, bold=None, color=None) -> None:
    # Word stores Latin and East Asian fonts separately.  Using one Chinese
    # font for every slot causes digits, punctuation and field results to be
    # substituted inconsistently across Word and LibreOffice.
    run.font.name = LATIN_FONT
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), LATIN_FONT)
    r_fonts.set(qn("w:hAnsi"), LATIN_FONT)
    r_fonts.set(qn("w:eastAsia"), cn)
    r_fonts.set(qn("w:cs"), LATIN_FONT)
    r_fonts.set(qn("w:hint"), "eastAsia")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_format(paragraph, *, indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=1.5, before=0, after=0) -> None:
    pf = paragraph.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line
    if indent:
        pf.first_line_indent = Cm(0.74)


def add_field(paragraph, instruction: str, *, placeholder: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(paragraph, "PAGE", placeholder="1")


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = LATIN_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:cs"), LATIN_FONT)
    normal.font.size = Pt(10.5)
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(0.74)

    heading_specs = [
        ("Title", HEADING_FONT, 16, NAVY, 0, 8),
        ("Subtitle", BODY_FONT, 12, NAVY, 0, 6),
        ("Heading 1", HEADING_FONT, 14, NAVY, 12, 6),
        ("Heading 2", HEADING_FONT, 12, NAVY, 8, 4),
        ("Heading 3", HEADING_FONT, 10.5, NAVY, 6, 3),
    ]
    for name, cn, size, color, before, after in heading_specs:
        st = styles[name]
        st.font.name = LATIN_FONT
        st._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
        st._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), cn)
        st._element.rPr.rFonts.set(qn("w:cs"), LATIN_FONT)
        st.font.size = Pt(size)
        st.font.bold = True if name != "Subtitle" else False
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.5
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.first_line_indent = Cm(0)

    if "图表题注" not in [s.name for s in styles]:
        cap = styles.add_style("图表题注", WD_STYLE_TYPE.PARAGRAPH)
    else:
        cap = styles["图表题注"]
    cap.font.name = LATIN_FONT
    cap._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    cap._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    cap._element.rPr.rFonts.set(qn("w:cs"), LATIN_FONT)
    cap.font.size = Pt(9)
    cap.font.bold = False
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(4)
    cap.paragraph_format.line_spacing = 1.0
    cap.paragraph_format.keep_with_next = True

    if "摘要正文" not in [s.name for s in styles]:
        abstract = styles.add_style("摘要正文", WD_STYLE_TYPE.PARAGRAPH)
    else:
        abstract = styles["摘要正文"]
    abstract.font.name = LATIN_FONT
    abstract._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    abstract._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    abstract._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    abstract._element.rPr.rFonts.set(qn("w:cs"), LATIN_FONT)
    abstract.font.size = Pt(10.5)
    abstract.paragraph_format.line_spacing = 1.5
    abstract.paragraph_format.first_line_indent = Cm(0.74)
    abstract.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abstract.paragraph_format.space_after = Pt(0)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    section.header_distance = Cm(1.27)
    section.footer_distance = Cm(1.27)
    section.different_first_page_header_footer = True

    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""

    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)
    for run in fp.runs:
        set_run_font(run, cn=BODY_FONT, size=8.5, color=NAVY)

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(115)
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("从数据到执行")
    set_run_font(r, cn=HEADING_FONT, size=24, bold=True, color=NAVY)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(14)
    r = p2.add_run("量化交易策略与机器学习应用的综合实践报告")
    set_run_font(r, cn=HEADING_FONT, size=16, bold=True, color=NAVY)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(76)
    r = p3.add_run("任务一至任务七学习成果总结")
    set_run_font(r, cn=BODY_FONT, size=12, color=NAVY)

    info = doc.add_table(rows=4, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.autofit = False
    info.columns[0].width = Cm(3.0)
    info.columns[1].width = Cm(6.5)
    rows = [
        ("报告作者", "李沐晓"),
        ("课程名称", "量化交易工作坊"),
        ("报告性质", "专业学习成果报告"),
        ("完成日期", "2026年7月25日"),
    ]
    for i, (k, v) in enumerate(rows):
        info.cell(i, 0).text = k
        info.cell(i, 1).text = v
        for j in range(2):
            cell = info.cell(i, j)
            set_cell_margins(cell, top=120, bottom=120)
            set_cell_border(
                cell,
                bottom={"val": "nil"},
                top={"val": "nil"},
                left={"val": "nil"},
                right={"val": "nil"},
            )
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.first_line_indent = Cm(0)
            set_run_font(para.runs[0], cn=BODY_FONT, size=10.5, bold=(j == 0), color=NAVY)
    info.rows[0].height = Cm(0.9)
    doc.add_page_break()


def add_toc_line(doc: Document, title: str, page: int, *, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.65 if level == 2 else 0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.tab_stops.add_tab_stop(Cm(15.3), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r = p.add_run(title)
    set_run_font(r, cn=BODY_FONT, size=9.2 if level == 1 else 8.8, bold=(level == 1), color=NAVY)
    r2 = p.add_run(f"\t{page}")
    set_run_font(r2, size=8.8, color=NAVY)


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run("目录")
    set_run_font(r, cn=BODY_FONT, size=16, bold=True, color=NAVY)
    toc_entries = [
        (1, "摘要", 3),
        (1, "一、量化交易核心概念", 4),
        (2, "1.1　从算法到交易系统", 4),
        (2, "1.2　数据、指标、信号与交易结果", 4),
        (2, "1.3　量化交易的核心价值", 5),
        (2, "1.4　统一评价口径", 6),
        (1, "二、量化交易策略综合分析", 7),
        (2, "2.1　规则策略的共同比较口径", 7),
        (2, "2.2　双均线与海龟策略的结果和原因", 7),
        (2, "2.3　市场环境、参数与执行风险", 9),
        (2, "2.4　从单标的回测到宽基指数基金策略", 10),
        (2, "2.5　成本压力、短窗回放与模拟盘状态", 11),
        (2, "2.6　多策略系统的组合方法", 13),
        (1, "三、机器学习在量化交易中的应用总结", 14),
        (2, "3.1　低信噪比改变了机器学习问题", 14),
        (2, "3.2　数据预处理、特征工程与时间验证", 14),
        (2, "3.3　模型选择与评价：任务五", 15),
        (2, "3.4　从预测到交易：任务六", 16),
        (2, "3.5　机器学习的优势与局限", 18),
        (2, "3.6　深度学习与非结构化数据的研究方向", 18),
        (1, "四、结论与展望", 19),
        (2, "4.1　主要收获与认识变化", 19),
        (2, "4.2　任务六失败带来的方法转变", 19),
        (2, "4.3　机构与个人的评价目标", 19),
        (2, "4.4　下一阶段计划", 20),
        (1, "附录　基于前期作业复盘的改进建议", 21),
        (1, "参考资料", 22),
    ]
    for level, title, page in toc_entries:
        add_toc_line(doc, title, page, level=level)
    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1, *, page_break=False) -> None:
    p = doc.add_paragraph(text, style=f"Heading {level}")
    if page_break:
        p.paragraph_format.page_break_before = True


def add_paragraph(doc: Document, text: str, *, bold_prefix: str | None = None, indent=True, keep=False) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, indent=indent)
    p.paragraph_format.keep_together = keep
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True, color=NAVY)
        r2 = p.add_run(text[len(bold_prefix) :])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)


def add_callout(doc: Document, title: str, text: str, *, color=NAVY) -> None:
    add_paragraph(doc, f"{title}：{text}", bold_prefix=f"{title}：")


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Cm(-0.38)
        p.paragraph_format.line_spacing = 1.35
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(item)
        set_run_font(r)


def add_figure(doc: Document, filename: str, caption: str, width=5.72) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(5)
    run = p.add_run()
    run.add_picture(str(FIG_DIR / filename), width=Inches(width))
    cp = doc.add_paragraph(caption, style="图表题注")
    cp.paragraph_format.keep_together = True
    for r in cp.runs:
        set_run_font(r, size=9, color=NAVY)


def add_table(
    doc: Document,
    number: int,
    title: str,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    widths: Sequence[float] | None = None,
    *,
    font_size=8.8,
    page_break_before=False,
) -> None:
    if page_break_before:
        doc.add_page_break()
    cap = doc.add_paragraph(f"表{number}　{title}", style="图表题注")
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(4)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if widths is None:
        total = 14.4
        widths = [total / len(headers)] * len(headers)
    else:
        scale = 14.4 / sum(widths)
        widths = [w * scale for w in widths]
    for i, w in enumerate(widths):
        table.columns[i].width = Cm(w)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    prevent_row_split(hdr)
    for j, h in enumerate(headers):
        c = hdr.cells[j]
        c.text = h
        set_cell_shading(c, MID_GRAY)
        set_cell_margins(c, top=80, bottom=80, start=120, end=120)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.0
        for r in p.runs:
            set_run_font(r, cn=BODY_FONT, size=font_size, bold=True, color=NAVY)
        set_cell_border(
            c,
            top={"val": "single", "sz": "4", "color": "000000"},
            bottom={"val": "single", "sz": "4", "color": "000000"},
            left={"val": "single", "sz": "4", "color": "000000"},
            right={"val": "single", "sz": "4", "color": "000000"},
        )
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for j, value in enumerate(row):
            c = cells[j]
            c.text = str(value)
            set_cell_margins(c, top=80, bottom=80, start=120, end=120)
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 or len(str(value)) > 18 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                set_run_font(r, cn=BODY_FONT, size=font_size)
            set_cell_border(
                c,
                bottom={"val": "single", "sz": "4", "color": "000000"},
                top={"val": "single", "sz": "4", "color": "000000"},
                left={"val": "single", "sz": "4", "color": "000000"},
                right={"val": "single", "sz": "4", "color": "000000"},
            )
    # ``cantSplit`` protects individual rows only.  Chaining the paragraphs
    # through every row except the final one keeps the complete table together
    # on one page, as required by the report brief.
    for row_idx, row in enumerate(table.rows):
        keep_with_next = row_idx < len(table.rows) - 1
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_with_next = keep_with_next
                paragraph.paragraph_format.keep_together = True
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_abstract(doc: Document) -> None:
    p = doc.add_paragraph("摘要", style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    text = (
        "围绕量化交易从研究到执行的完整链条，报告总结前七项任务形成的数据处理、规则策略、机器学习与模拟交易实践。研究采用复权行情清洗、统一基准回测、时间顺序划分、样本外评价、回撤分析和成本压力测试。双均线与海龟策略在下跌样本中减少相对损失，但绝对收益仍为负；任务六模型测试曲线下面积为0.571，策略收益低于简单基准；任务七策略C平台长期累计收益为71.84%，成本翻倍后优势明显收窄，模拟盘尚无连续数据。正文四部分形成从概念界定与评价口径，到规则策略比较和机器学习交易检验，再到依据前述证据总结认识与制定研究计划的递进关系。结果表明，策略价值取决于样本外稳定性、风险、成本与执行的共同约束；后续将以严格验证研究多模态和非结构化数据的增量信息。"
    )
    ap = doc.add_paragraph(text, style="摘要正文")
    ap.paragraph_format.space_after = Pt(10)
    kp = doc.add_paragraph()
    kp.paragraph_format.first_line_indent = Cm(0)
    kp.paragraph_format.line_spacing = 1.3
    r1 = kp.add_run("关键词：")
    set_run_font(r1, cn=HEADING_FONT, size=10.5, bold=True, color=NAVY)
    r2 = kp.add_run("量化交易；趋势跟踪；机器学习；样本外验证；风险管理")
    set_run_font(r2)


def chapter1(doc: Document) -> None:
    add_heading(doc, "一、量化交易核心概念与学习认识", 1)
    add_heading(doc, "1.1　学习目标", 2)
    add_paragraph(
        doc,
        "本次学习以“把交易想法转化为可检验、可执行和可复盘的规则”为主线。传统主观交易能够吸收经验和情境判断，但在信息处理规模、纪律一致性、事后归因和团队复核方面存在天然限制。量化交易通过数据、模型和程序把决策条件显式化，使策略可以在历史样本中检验，也可以在执行阶段记录偏差；然而，它并不保证盈利，更不能消除市场不确定性。",
    )
    add_paragraph(
        doc,
        "学习前，我更倾向于把量化交易理解为算法，类似解决一道条件已经给出的数学题：只要模型足够好，就应当得到确定答案。前七项任务逐步改变了这一理解。金融市场的数据并不完整，变量含义会随时间和制度变化，模型还需要在未知环境中接受检验。即使找到了信号，也要经过仓位、成本、成交和监控才能成为收益。因此，我现在更愿意把量化交易理解为一个持续面对信息不对称的研究与交易系统。",
    )
    add_heading(doc, "1.2　量化交易的基本概念", 2)
    add_paragraph(
        doc,
        "量化交易是利用可重复的数据处理和明确规则形成投资决策的方法。回测是按照预先定义的信号和成交假设在历史数据上模拟交易；样本外验证是在模型或参数确定后，用未参与选择的数据检验其迁移能力；最大回撤衡量资产净值从历史高点跌至后续低点的最大幅度；滑点则表示理论成交价与实际可成交价格之间的偏差。四者共同提醒研究者：预测分数不是交易结果，历史高分也不是未来承诺。",
    )
    add_figure(doc, "图01_量化研究闭环.png", "图1　量化研究与执行的持续闭环")
    add_paragraph(
        doc,
        "图1把前七项任务重新组织为一条直线型逻辑：先定义问题和证据标准，再治理数据、构建特征与模型，随后完成组合决策和执行监控，最后通过预期与实现差异进入下一轮迭代。该图不是结果证据，而是本报告解释各项成果的系统框架。其关键含义是，任何一个环节的误差都可能传导到最终收益，单独优化算法无法替代对数据和执行的控制。",
    )
    add_heading(doc, "1.3　量化交易的核心价值", 2)
    add_paragraph(
        doc,
        "量化交易的第一项价值是规则透明：研究者必须说明何时买入、何时退出、使用多少仓位以及在何种条件下停止。第二项价值是证据纪律：结果必须同时说明样本、基准、成本和选择过程。第三项价值是风险前置：在下单前定义可接受损失，而不是在亏损后临时解释。第四项价值是可扩展与可复盘：相同流程可以迁移到更多标的和周期，并留下可核对记录。",
    )
    add_table(
        doc,
        1,
        "任务一至任务七的学习成果与证据状态",
        ["任务", "主要成果", "在系统中的位置", "证据状态"],
        [
            ["任务一", "多市场行情获取、复权与质量检查", "数据治理", "可复算教学样本"],
            ["任务二", "技术指标计算与信息分类", "特征构建", "方法与图形验证"],
            ["任务三", "双均线策略、成本回测与参数扫描", "规则基线", "单标的短样本"],
            ["任务四", "突破、波动率止损与风险预算", "风险约束策略", "单标的短样本"],
            ["任务五", "分类流程、时间划分与泛化检验", "预测模型", "含独立测试结果"],
            ["任务六", "横截面排序、组合构建与方向实验", "预测到决策", "测试期仅三个季度"],
            ["任务七", "三策略平台回测、本地复算与部署", "执行验证", "模拟盘尚无运行数据"],
        ],
        widths=[1.7, 5.6, 3.4, 4.6],
        font_size=8.6,
    )


def chapter2(doc: Document) -> None:
    add_heading(doc, "二、数据基础与特征构建", 1)
    add_heading(doc, "2.1　任务一的数据获取与质量控制", 2)
    add_paragraph(
        doc,
        "任务一使用公开数据接口获取五只A股和五只港股的日线开盘价、最高价、最低价、收盘价与成交量，并将数据获取、字段标准化、日期排序和图表输出连接为可重复管线。前复权处理用于降低除权除息对历史价格序列和技术指标的机械冲击。质量检查至少覆盖日期重复、缺失记录、字段类型、异常价格关系和数据区间，使后续计算建立在明确版本上。",
    )
    add_paragraph(
        doc,
        "这些处理满足了教学样本需要，但不能被理解为完整的数据治理。公开接口可能调整字段、频率和历史覆盖；前复权不解决幸存者偏差、停牌、涨跌停、公告时点和指数成分历史变化；当天收盘后才能获得的数据也不能假设在收盘前已经可用。因此，后续研究应建立“数据来源—下载时间—复权口径—修订记录”的版本表，并对每一特征标注最早可用时点，具体改进见附录A-01。",
    )
    add_heading(doc, "2.2　任务二的技术指标构建", 2)
    add_paragraph(
        doc,
        "任务二计算并解释了移动平均线、指数平滑异同移动平均线、相对强弱指标、随机指标、平均真实波幅、布林带、平均趋向指数和商品通道指数。与逐一背诵公式相比，更重要的是判断指标提供了什么信息以及是否与其他指标重复。表2按交易作用将指标压缩为四类，使指标选择与后续信号、仓位和风险控制直接对应。",
    )
    add_table(
        doc,
        2,
        "技术指标分类及其系统作用",
        ["信息类别", "代表指标", "主要作用", "常见误用"],
        [
            ["趋势方向", "移动平均线、指数平滑异同移动平均线", "描述价格方向与平滑后的趋势结构", "把滞后指标当作提前预测"],
            ["动量速度", "相对强弱指标、随机指标", "衡量涨跌速度和短期极端位置", "机械使用固定超买超卖阈值"],
            ["波动程度", "平均真实波幅、布林带", "设置止损距离、仓位尺度和波动状态", "忽略跳空和流动性导致的成交偏差"],
            ["趋势强度与偏离", "平均趋向指数、商品通道指数", "辅助判断状态和价格相对位置", "把更多指标误认为更多独立信息"],
        ],
        widths=[2.6, 4.5, 4.4, 3.8],
        font_size=8.5,
    )
    add_paragraph(
        doc,
        "技术指标本质上是对历史价量信息的变换。不同平滑方法会改变响应速度，参数也不存在跨市场、跨周期固定最优值。若多个指标由相同价格窗口计算，它们可能高度相关，变量数量增加不代表信息维度增加。指标只有被转化为明确的入场、退出和仓位规则后，才能通过回测和样本外检验判断是否具有交易价值。",
    )
    add_heading(doc, "2.3　对数据与指标的认识变化", 2)
    add_paragraph(
        doc,
        "任务一使我第一次意识到，下载一份行情和建立一条可复现的数据管线并不是同一件事；任务二则修正了“指标越多越全面”的直觉。多个技术指标可能只是对同一价格历史的不同表达，真正稀缺的是及时、独立且与交易目标相关的信息。Jump Trading课程材料关于低信噪比和研究系统的讨论，帮助我重新理解这一点：更丰富的数据只有在时间对齐、偏差控制和可交易目标明确的前提下才有价值。",
    )
    add_callout(
        doc,
        "本章小结",
        "数据和指标只是研究材料。只有进一步写成入场、退出、仓位和成本规则，才能通过回测比较不同策略在同一市场样本中的收益与风险。",
        color=ORANGE,
    )


def chapter3(doc: Document) -> None:
    add_heading(doc, "三、规则量化策略综合分析", 1)
    add_heading(doc, "3.1　比较口径与结论边界", 2)
    add_paragraph(
        doc,
        "任务三和任务四均以贵州茅台为标的，核心比较区间为2025年5月29日至2026年7月3日，初始资金100万元，手续费按买卖各0.03%计入，买卖滑点各0.01%，共同约束为仅做多，并以同期买入持有作为基准。相同口径使策略差异主要来自信号和风险管理，但样本只有约十三个月且主要处于下跌环境，因此属于教学性比较，不能代表长期、跨资产或跨市场表现。",
    )
    add_table(
        doc,
        3,
        "双均线、海龟策略与买入持有的主要结果",
        ["指标", "双均线策略", "海龟策略", "买入持有"],
        [
            ["累计收益率", "-5.67%", "-0.74%", "-18.24%"],
            ["年化收益率", "-3.61%", "-0.47%", "-11.63%"],
            ["最大回撤", "-15.20%", "-9.72%", "-21.35%"],
            ["夏普比率", "-0.58", "-0.29", "-1.12"],
            ["交易胜率", "35.71%", "40.00%", "不适用"],
            ["完整交易次数", "14", "5", "不适用"],
        ],
        widths=[4.1, 3.7, 3.7, 3.7],
        font_size=9.0,
    )
    add_figure(doc, "图02_规则策略收益与回撤.png", "图2　任务三与任务四规则策略的收益和回撤比较")
    add_paragraph(
        doc,
        "从图2和表3可以看出，双均线策略相对买入持有少亏12.57个百分点，最大回撤低6.15个百分点；海龟策略相对买入持有少亏17.50个百分点，最大回撤低11.63个百分点。在本样本中，海龟策略的绝对损失和回撤也低于双均线。但两项策略的绝对收益和夏普比率均为负，因此这里只能说明风险暴露有所降低，不能说明策略已经能够稳定盈利。",
    )
    add_heading(doc, "3.2　双均线策略的特点", 2)
    add_paragraph(
        doc,
        "双均线策略在五日均线上穿十五日均线时买入，在下穿时退出。它只依赖收盘价、两个窗口和一组清晰规则，优点是容易复核、参数空间较小，并能完整展示数据、信号、成本、交易记录和绩效评价的闭环。局限在于信号滞后、震荡期反复交易、固定全仓导致风险随市场波动变化，而且没有独立止损和市场状态过滤。",
    )
    add_figure(doc, "图03_双均线参数敏感性.png", "图3　双均线参数组合的累计收益率敏感性")
    add_paragraph(
        doc,
        "图3对六个短周期和六个长周期进行了扫描。大部分组合收益为负，正收益集中在短周期3日、长周期10日的孤立区域，相邻参数并没有形成连续的正收益带。参数热力图不能只看最高值，还要观察较优区域是否平滑、连续，以及能否在未参与选择的时期保留。这里的孤立高值更可能受到样本影响，不能直接称为最优参数。",
    )
    add_heading(doc, "3.3　海龟策略的风险控制", 2)
    add_paragraph(
        doc,
        "海龟策略使用唐奇安通道突破作为入场和退出条件，以平均真实波幅衡量市场波动，并据此设置止损距离和风险预算。与双均线相比，其主要进步并不是声称突破预测更准确，而是把“每次判断错误可能损失多少”放到下单前。两倍平均真实波幅在当前候选参数中表现相对平衡，但理论风险预算只有在止损可成交、没有严重跳空且流动性充足时才近似成立。",
    )
    add_table(
        doc,
        4,
        "规则策略的适用场景与失效条件",
        ["市场状态", "双均线策略", "海龟策略", "共同风险"],
        [
            ["持续上涨趋势", "能够跟随方向，但入场和退出均滞后", "突破入场并按波动率管理风险", "趋势反转时利润回吐"],
            ["区间震荡", "均线反复交叉，形成小额亏损", "假突破后连续止损", "成本与滑点累积"],
            ["单边下跌、仅做多", "可通过空仓减少部分损失", "风险预算可能压低回撤", "难以获得正绝对收益"],
            ["跳空或流动性不足", "理论成交价偏乐观", "止损价格可能无法实现", "回测与执行偏差扩大"],
        ],
        widths=[3.0, 4.2, 4.2, 3.9],
        font_size=8.5,
    )
    add_heading(doc, "3.4　两项策略的关联与互补", 2)
    add_paragraph(
        doc,
        "两项策略都属于趋势跟踪，收益来源可能高度相关。双均线强调平滑后的方向变化，海龟强调价格突破和波动率尺度；这种机制差异不足以自动形成分散化。多策略系统还需要加入不同资产、不同时间尺度或不同收益来源，并在组合层限制总风险，改进方向见附录A-05和A-06。",
    )
    add_paragraph(
        doc,
        "任务三对我的意义不是证明五日和十五日均线可以盈利，而是第一次把完整回测闭环建立起来。任务四进一步把注意力从“入场是否押对”转向“错误发生时损失是否受控”。两项策略绝对收益均为负，使我不再把研究目标设成寻找一组漂亮参数，而是先问策略为什么可能有效、在哪些环境下失效、风险是否可以承受，以及结论能否在新样本中被否证。",
    )


def chapter4(doc: Document) -> None:
    add_heading(doc, "四、机器学习在量化交易中的应用", 1)
    add_heading(doc, "4.1　机器学习研究流程", 2)
    add_paragraph(
        doc,
        "机器学习在量化交易中的作用，是在高维数据中估计条件关系、排序标的或辅助识别市场状态。完整流程包括目标定义、时间戳对齐、缺失和异常处理、特征工程、训练—验证—测试划分、模型选择、概率或排序评价、成本化回测以及上线监控。若目标和可交易决策没有连接，即使分类分数提高，也可能无法形成更好的收益。",
    )
    add_table(
        doc,
        5,
        "机器学习研究流程、评价重点与常见风险",
        ["环节", "关键问题", "主要评价", "常见风险"],
        [
            ["数据预处理", "样本时点、缺失、异常和复权是否一致", "覆盖率、缺失率、标签平衡", "未来信息泄漏、幸存者偏差"],
            ["特征工程", "特征是否及时、独立且可解释", "稳定性、共线性、跨期分布", "变量堆叠、重复信息"],
            ["模型训练", "复杂度是否与有效样本匹配", "验证期曲线下面积、排序相关", "对验证集过度调参"],
            ["样本外评价", "是否使用未参与选择的数据", "区间、分期结果、简单基线", "看过测试集后继续选择"],
            ["交易转化", "分数如何变为仓位与订单", "净收益、回撤、换手和容量", "预测改善但收益不改善"],
            ["部署监控", "数据和关系是否漂移", "信号、成交、风险与停用指标", "模型失效后仍机械运行"],
        ],
        widths=[2.2, 5.0, 4.0, 4.1],
        font_size=8.3,
    )
    add_heading(doc, "4.2　任务五：教学数据与金融数据的差异", 2)
    add_paragraph(
        doc,
        "任务五先在乳腺癌教学数据上比较逻辑回归、决策树、随机森林和梯度提升，再将同样的数据治理和评价框架迁移到金融分类。医疗教学数据的逻辑回归与随机森林测试曲线下面积约为0.994，说明在结构清晰、标签相对稳定的样本中，模型可以获得很强区分能力；但这一结果不能外推到金融市场，因为两类问题的数据生成过程、信噪比和非平稳性完全不同。",
    )
    add_figure(doc, "图04_TASK5样本外AUC区间.png", "图4　任务五金融分类模型的样本外判别能力及不确定区间")
    add_paragraph(
        doc,
        "图4以2025年测试样本为横轴，0.500表示随机排序基准。逻辑回归、决策树和随机森林的曲线下面积分别为0.509、0.477和0.533，重采样区间均覆盖0.500。即使随机森林的点估计最高，也只能说明它在这一测试样本中略高，不能说明模型已经具有稳定预测能力。区间较宽也说明，有限而相关的金融样本不足以支持精确判断，增加模型复杂度并不会自动增加信息。",
    )
    add_figure(doc, "图05_TASK5跨期AUC.png", "图5　任务五模型判别能力在验证、开发和测试时期的变化")
    add_paragraph(
        doc,
        "图5分别展示2023年验证期、2024年开发期和2025年测试期的结果。三个模型在验证期的表现较高，进入开发期后下降，测试期进一步接近随机水平。由于每个时期只有一个聚合结果，不能据此估计长期衰减速度；但这种跨期变化已经提示分布漂移和验证集选择偏差。失败结果的价值在于迫使我检查标签平移、随机标签和分期评价，而不是只保留漂亮分数。",
    )
    add_heading(doc, "4.3　任务六：横截面排序与组合构建", 2)
    add_paragraph(
        doc,
        "任务六以季度横截面数据构建二十三个特征，使用七个季度训练、三个季度测试，对线性回归、岭回归、逻辑回归、决策树、随机森林和梯度提升进行比较。线性回归按预设验证规则入选，验证期平均排序相关系数为0.129，测试期为0.273；每季选择预测排名前30只股票并等权持有，扣除单边20个基点成本后，三个测试季度累计净收益为41.46%，同期全市场等权为-15.91%。",
    )
    add_figure(doc, "图06_TASK6季度组合收益.png", "图6　任务六测试期线性回归前30组合与全市场等权的季度收益")
    add_paragraph(
        doc,
        "图6显示策略在三个测试季度均取得正收益，而全市场等权在其中两个季度为负。这一结果与模型存在一定横截面排序信息相一致，但测试期只有三个季度，独立市场状态数量远少于股票行数，不能用横截面样本量夸大时间证据。随机森林测试累计净收益达到50.68%，但该结果是在查看测试数据后才知道，不能反过来替换预先选择的线性回归并称为独立最优模型。平均换手率83.33%也提示收益对交易成本和换仓时点敏感。",
    )
    add_heading(doc, "4.4　方向预测实验", 2)
    add_figure(doc, "图07_TASK6验证与测试AUC.png", "图7　任务六三轮方向模型的验证与测试表现")
    add_paragraph(
        doc,
        "图7连接同一轮实验的验证和测试结果。第一轮静态随机森林验证曲线下面积为0.756，测试仅0.335；第二轮静态逻辑回归从0.692降至0.520；第三轮使用180日滚动训练后，验证为0.631、测试为0.571，迁移差距缩小。第三轮结果表明滚动更新可能更适应时变关系，但仍没有区间估计，而且每轮设计在前一轮结果之后调整，属于探索性迭代，不能把最终数值视为完全独立的确认性证据。",
    )
    add_paragraph(
        doc,
        "更重要的是，第三轮弱预测优势没有自动形成更高交易收益。机器学习择时策略测试累计收益仅0.50%，同期买入持有为9.93%，均线策略为9.68%。原因包括阈值把弱概率差异离散化、模型长期低仓位、换手成本和预测目标与收益目标不一致。这个结果把“模型评价”与“策略评价”明确分开：曲线下面积衡量排序能力，交易收益还取决于方向幅度、仓位函数、成本和市场路径。",
    )
    add_table(
        doc,
        6,
        "任务五与任务六主要结果及证据边界",
        ["研究对象", "主要结果", "能够支持的判断", "不能支持的判断"],
        [
            ["任务五医疗教学数据", "逻辑回归、随机森林曲线下面积约0.994", "流程能够识别结构清晰样本", "不能外推金融盈利能力"],
            ["任务五金融分类", "测试曲线下面积0.477至0.533，区间覆盖0.500", "未发现稳定判别优势", "不能称随机森林显著有效"],
            ["任务六横截面排序", "线性回归测试排序相关0.273；组合累计净收益41.46%", "本组三季度中存在排序与组合效果", "不能证明长期稳定或容量充足"],
            ["任务六方向模型", "滚动逻辑回归测试曲线下面积0.571；策略收益0.50%", "预测与交易目标需要分开评价", "不能用分数改善替代收益证据"],
        ],
        widths=[3.2, 4.4, 4.2, 3.5],
        font_size=8.3,
        page_break_before=False,
    )
    add_heading(doc, "4.5　对机器学习应用的反思", 2)
    add_paragraph(
        doc,
        "任务六的失败最触动我。作为初学者，我在模型调参时很容易把注意力放在“下一组参数能否更高”，却难以判断变化来自真实结构、样本偶然还是对验证集的适应。验证高分没有稳定迁移，滚动模型的弱预测优势也没有转化为收益，使我感性地理解了金融问题的低信噪比：交易需要市场知识、经验和判断，调参不能替代对资产、制度、流动性和参与者行为的理解。",
    )
    add_paragraph(
        doc,
        "Jump Trading课程材料提出“真实信号还是结构化噪声”的问题，我进一步追问：为什么应当先验地相信金融市场一定存在可持续、可发现且扣除成本后仍可盈利的信号？数学题通常有明确公理和可计算目标，市场却不断受到参与者学习和制度变化影响。因此，以后每个模型都应被当作可被否证的假设；“尚未发现稳定信号”也是有效结果，而不是必须通过更复杂模型掩盖的失败。更严格的时间验证和嵌套选择见附录A-04。",
    )


def chapter5(doc: Document) -> None:
    add_heading(doc, "五、任务七：从回测到模拟交易", 1)
    add_heading(doc, "5.1　当前证据范围", 2)
    add_paragraph(
        doc,
        "任务七把研究从单次回测推进到统一平台、本地复算、成本压力测试和模拟盘部署。为了避免把不同结果混在一起，本报告分别说明JoinQuant平台长期回测、本地扩展样本复算、冻结参数后的五日历史回放，以及已部署但尚无运行记录的私有模拟盘。这四部分回答的问题不同，不能拼接成一条看似连续的实盘净值。",
    )
    add_table(
        doc,
        8,
        "任务七四层证据及当前状态",
        ["证据层级", "时间范围", "用途", "当前状态"],
        [
            ["平台长期回测", "2019-01-01至2026-07-17", "统一平台比较三策略历史绩效", "已完成，可报告历史指标"],
            ["本地扩展复算", "2016-01-04至2026-07-24", "分析路径、暴露、参数与成本敏感性", "已完成，不替代平台结果"],
            ["冻结参数五日回放", "2026-07-20至2026-07-24", "观察信号、仓位与机会成本", "三策略均空仓，收益0.00%"],
            ["私有模拟盘", "2026-07-24部署", "检验订单、成交、持仓和执行偏差", "截至7月25日尚无运行数据"],
        ],
        widths=[3.1, 3.9, 5.0, 3.3],
        font_size=8.4,
    )
    add_callout(
        doc,
        "说明",
        "截至2026年7月25日，三个私有模拟盘已经部署，但尚未形成可分析的净值、订单、成交与持仓记录。因此本报告不对模拟交易或真实资金收益作定量结论。",
        color=RED,
    )
    add_heading(doc, "5.2　三套策略与长期回测结果", 2)
    add_paragraph(
        doc,
        "策略A是300日均线全仓基线，用于观察长期趋势过滤的最简实现；策略B在长短均线信号上加入平均真实波幅止损、波动率仓位和冷却期，目标是降低回撤；策略C在三只宽基交易型开放式指数基金中进行动量选择，并设置单标的仓位上限和现金缓冲，目标是扩展资产选择同时限制集中风险。三者分别代表基线、风险约束和资产轮动三个层次。",
    )
    add_table(
        doc,
        7,
        "任务七三策略JoinQuant平台长期回测指标",
        ["指标", "策略A", "策略B", "策略C"],
        [
            ["累计收益率", "-0.66%", "2.91%", "71.84%"],
            ["年化收益率", "-0.09%", "0.39%", "7.69%"],
            ["超额收益率", "-33.97%", "-31.59%", "14.23%"],
            ["最大回撤", "-36.97%", "-12.31%", "-16.91%"],
            ["夏普比率", "-0.291", "-0.769", "0.315"],
            ["市场敏感度", "0.521", "0.164", "0.426"],
            ["年化波动率", "14.10%", "4.70%", "11.70%"],
        ],
        widths=[4.5, 3.6, 3.6, 3.6],
        font_size=8.9,
        page_break_before=False,
    )
    add_paragraph(
        doc,
        "平台结果中，策略C累计收益率和超额收益率最高，最大回撤低于策略A，但相对基准最大回撤仍达到-30.37%，说明“跑赢基准”不等于相对净值路径平稳。策略B最大回撤最低，市场敏感度和波动率也最低，但累计收益只有2.91%，其稳健性部分来自较低仓位和较长空仓。策略A长期收益接近零且回撤最大，表明单一超长期均线在该口径下无法形成足够补偿。",
    )
    add_heading(doc, "5.3　本地复算与成本压力测试", 2)
    add_figure(doc, "图08_TASK7本地扩展净值与回撤.png", "图8　任务七三策略与沪深300的本地扩展样本净值和回撤")
    add_paragraph(
        doc,
        "图8使用2016年1月4日至2026年7月24日的本地结构化数据，将各策略和沪深300净值归一化，并在下方面板展示回撤。策略A的主要问题不是某一次急跌，而是2021年后回撤长期扩大，期末累计收益-30.38%、最大回撤-42.91%。策略B净值长期接近初始值，最大回撤-11.82%，但平均仓位仅18.95%，低回撤与低市场参与度不可分开。策略C的优势主要在2025年后形成，期末累计收益56.93%、最大回撤-16.72%、平均仓位34.07%，仍需要更多状态验证。",
    )
    add_paragraph(
        doc,
        "本地复算与平台结果在收益幅度上并不完全一致，原因可能包括数据源、复权、交易日、成交假设、平台函数和版本差异。这里不能选择更好看的口径，而应分别报告并记录差异来源。平台结果用于统一环境比较，本地复算用于拆解净值、仓位和成本，两者可以相互核对，但不能互相替代。",
    )
    add_figure(doc, "图09_TASK7成本压力测试.png", "图9　任务七本地扩展样本的成本翻倍压力测试")
    add_paragraph(
        doc,
        "图9在相同纵轴下比较基础成本和成本翻倍。策略A累计收益从-30.38%下降至-53.61%，变化-23.23个百分点；策略B从2.16%降至-9.35%，变化-11.51个百分点；策略C从56.93%降至25.49%，变化-31.44个百分点。策略C在压力情景下仍为正，但利润被大幅压缩，说明高收益路径并不等于对成本稳健。策略B绝对变化较小，却由微利转为亏损，反映其收益缓冲不足。改进方向见附录A-03。",
    )
    add_heading(doc, "5.4　五日回放与模拟盘进展", 2)
    add_paragraph(
        doc,
        "冻结参数后的五日历史回放覆盖2026年7月20日至24日，三策略均未产生买卖，累计收益和最大回撤均为0.00%，同期沪深300上涨2.65%，平台口径超额约-2.58%。短窗空仓说明策略按规则没有信号，不等于风险已经消失：它同时产生机会成本，也不能检验成交、滑点、停牌和订单拒绝。五日结果只用于验证状态和信号一致性，不能称为模拟交易。",
    )
    add_paragraph(
        doc,
        "任务七使“量化交易是系统”不再只是抽象说法。数据版本决定信号，信号经过仓位和成本才成为订单，订单还要面对成交和监控。过去我更关注算法输出，现在会同时追问：数据何时可得、策略是否按版本运行、空仓是否符合预期、成本是否侵蚀优势、出现多大偏差时应停用。模拟盘下一步需要持续记录净值、订单、成交、持仓和拒单原因，至少跨越若干市场状态后再讨论可执行性，见附录A-08。",
    )


def chapter6(doc: Document) -> None:
    add_heading(doc, "六、多策略系统与风险分析", 1)
    add_heading(doc, "6.1　预测、决策与执行", 2)
    add_paragraph(
        doc,
        "综合前述任务，多策略量化系统应至少分为预测层、决策层和执行层。预测层输出方向、收益或排序信息；决策层把弱信号转化为资产选择、仓位和风险预算；执行层负责订单、成交、成本和监控。层级分开可以避免用一个模型同时承担所有目标，也便于定位问题究竟来自数据、预测、组合还是执行。",
    )
    add_table(
        doc,
        9,
        "多策略系统的层级、输入输出与控制措施",
        ["系统层级", "主要输入", "主要输出", "控制措施"],
        [
            ["预测层", "时点对齐的数据、技术与非结构化特征", "方向概率、收益估计、横截面排序", "滚动验证、简单基线、漂移监控"],
            ["决策层", "预测分数、相关性、波动率和成本", "标的选择、仓位、风险预算与现金比例", "组合回撤、集中度、换手和情景约束"],
            ["执行层", "目标仓位、市场流动性和订单规则", "成交、实际成本、持仓和偏差记录", "限价、成交率、拒单处理和停用规则"],
            ["复盘层", "预期与实际结果、异常事件和版本记录", "归因、修订计划和下一轮假设", "变更审批、可复现报告和审计追踪"],
        ],
        widths=[2.6, 4.7, 4.4, 3.6],
        font_size=8.4,
    )
    add_heading(doc, "6.2　多策略组合思路", 2)
    add_paragraph(
        doc,
        "规则策略提供透明基线和趋势暴露，机器学习模型可用于排序、状态识别或动态调整，资产轮动扩展收益来源。构建组合时不应仅按历史收益高低分配，而应先估计策略相关性、平均仓位、回撤同步性和成本。对高度相关的趋势策略，需要限制共同方向暴露；对低仓位策略，需要把现金贡献和机会成本单独归因；对高换手模型，需要在优化中直接扣除成本而不是事后解释。",
    )
    add_paragraph(
        doc,
        "机构与个人的目标函数并不相同。若作为服务金融机构的从业者，我会在治理、风险预算和审查约束内更进取地关注收益，同时重视模型可解释性，以便沟通和执行；作为个人投资者，我更看重稳定性、回撤和能否长期坚持。两种角色都不能忽略风险，只是收益目标、容量、问责和承受路径不同。组合设计因此应先写清使用场景，再决定优化指标。",
    )
    add_heading(doc, "6.3　主要风险", 2)
    add_paragraph(
        doc,
        "第一类是数据风险，包括历史修订、复权差异、时点不可得和样本选择偏差；第二类是模型风险，包括过拟合、概念漂移、低有效样本和解释失真；第三类是市场与组合风险，包括相关性上升、波动突变、流动性收缩和集中暴露；第四类是执行与治理风险，包括成交偏差、系统故障、版本不一致和缺少停用规则。任何一类风险都可能使回测优势在部署后消失。",
    )
    add_table(
        doc,
        10,
        "主要限制、影响与对应改进建议",
        ["主要限制", "可能影响", "当前处理", "改进编号"],
        [
            ["数据时点与版本记录不完整", "产生泄漏或难以复算", "统一来源文件并披露口径", "A-01"],
            ["规则策略为单标的短样本", "结论受市场状态影响", "限定为教学性结果", "A-02、A-05"],
            ["成本与可交易性假设简化", "高估净收益和止损效果", "任务七加入成本翻倍情景", "A-03"],
            ["机器学习验证状态数量有限", "验证高分难以迁移", "按时间划分并保留测试集", "A-04"],
            ["多策略相关性与风险预算未完整实现", "表面分散、实际同向暴露", "提出三层架构与角色目标", "A-06"],
            ["模拟盘尚无连续运行记录", "无法评价订单与成交偏差", "只报告部署事实", "A-07、A-08"],
        ],
        widths=[4.2, 4.2, 4.4, 2.5],
        font_size=8.4,
    )
    add_heading(doc, "6.4　从单次模型到持续研究", 2)
    add_paragraph(
        doc,
        "Jump Trading课程材料强调市场不会静止，竞争力来自透明、可复现和持续迭代的研究系统。结合本项目，这意味着每次实验都应保存数据版本、特征定义、参数、训练窗口、成本假设和结果；部署后监控输入漂移、信号分布、换手、成交率、回撤和基准偏离；达到预设阈值时降仓或停用。模型不应因为曾经高分而获得永久信任。",
    )
    add_callout(
        doc,
        "综合来看",
        "策略组合不是把多个收益曲线简单相加，而是把不同信息、风险和执行约束放在同一套规则中。系统优势首先来自可复核的证据和对错误的控制，其次才是模型复杂度。",
    )


def chapter7(doc: Document) -> None:
    add_heading(doc, "七、总结与展望", 1)
    add_heading(doc, "7.1　主要学习收获", 2)
    add_paragraph(
        doc,
        "第一，形成了从数据获取、质量检查、指标构建到规则回测的基础能力。任务一至任务四使每个交易想法都需要写成可复核条件，并与基准、成本和回撤共同评价。第二，理解了机器学习评价的层级：训练高分、验证高分、测试迁移和交易收益是不同问题。任务五和任务六显示，失败结果和不稳定结果比单次高分更能暴露研究边界。",
    )
    add_paragraph(
        doc,
        "第三，建立了更具体的风险意识。最大回撤不仅是报表指标，还代表策略在真实路径中能否被坚持；低回撤可能来自低仓位，高收益可能依赖高换手，止损也可能因跳空而失效。第四，任务七把版本、平台、本地复算、成本和部署连接起来，使我理解执行证据必须独立于历史回测积累。",
    )
    add_heading(doc, "7.2　核心结论", 2)
    add_bullets(
        doc,
        [
            "量化交易的核心不是单一算法，而是数据、特征、模型、组合、执行和监控共同构成的系统。",
            "任务三和任务四在当前下跌样本中降低了相对损失和回撤，但绝对收益为负，只能支持风险管理层面的学习结论。",
            "任务五表明医疗教学数据上的高曲线下面积不能外推金融市场；金融弱预测结果说明模型不能创造数据中不存在的信息。",
            "任务六发现一定横截面排序信息，但测试期只有三个季度；方向模型的分数改善没有自动转化为更高交易收益。",
            "任务七完成了从历史回测到部署的流程，但模拟盘尚无运行数据，现阶段不能评价真实成交或资金绩效。",
            "在低信噪比和市场变化下，不预设可盈利信号必然存在，把每个策略当作可否证假设，是比追求漂亮结果更重要的研究纪律。",
        ],
    )
    add_heading(doc, "7.3　后续研究计划", 2)
    add_paragraph(
        doc,
        "短期内，我会优先扩大标的和时间范围，使用滚动样本外和成本情景检验规则策略；对机器学习模型采用嵌套时间验证、稳定特征和简单基线，并把换手、回撤和容量直接纳入评价。任务七模拟盘将按日保存订单、成交、持仓和净值，按周复核信号与实现偏差，只有在形成足够连续记录后才讨论执行质量。",
    )
    add_paragraph(
        doc,
        "中长期我希望继续研究深度学习、多模态模型和非结构化数据。我的直觉仍倾向于认为复杂现象可能存在可解释结构，文本、图像、公告、新闻和产业链数据也许能够扩展纯价量信息集；但这是一项待检验假设，而不是“一切都能被模型解释”的结论。领域专用模型的价值应来自更好的问题定义、时间对齐和信息抽取，而不是用更大模型掩盖弱标签、数据泄漏或样本不足。",
    )
    add_paragraph(
        doc,
        "因此，探索复杂模型与反对复杂度崇拜并不矛盾。未来每个多模态实验都应与简单模型比较，说明新增信息何时可得，进行严格样本外验证，并在交易成本和执行约束下评价增量价值。如果复杂模型不能稳定超过简单基线，最诚实的结论仍然是尚未发现可交易优势。",
    )
    add_callout(
        doc,
        "总体认识",
        "学习前我把量化交易看作算法；学习后我把它看作一个持续面对信息不对称、证据不足和执行摩擦的系统。专业性不只表现为找到答案，也表现为知道哪些结论现在还不能成立。",
        color=GREEN,
    )


def final_chapter1(doc: Document) -> None:
    add_heading(doc, "一、量化交易核心概念", 1, page_break=True)
    add_heading(doc, "1.1　从算法到交易系统", 2)
    add_paragraph(
        doc,
        "学习开始时，我把量化交易主要理解为算法问题。这种理解隐含两个前提：数据已经准备完毕，预测目标也已经确定。在前七项任务中，这两个前提都无法直接成立。复权口径、信息披露时间、参数适用时期和可成交价格都会改变策略结果，因此算法只是交易系统中的一个环节。",
    )
    add_paragraph(
        doc,
        "量化交易可以定义为一套交易系统。它把对收益和风险的判断转化为可复现的数据、明确的规则和可执行的订单，并根据历史检验、样本外结果和实际成交持续修正。算法在其中生成信号，信号是由数据得到的买卖方向或资产排序；组合决策再把信号转换为具体标的和仓位；执行环节负责下单、成交和成本记录。市场参与者、制度和流动性持续变化，因此系统还必须允许关系失效、模型降权和策略停用。",
    )
    add_figure(doc, "图01_量化研究闭环.png", "图1　量化研究与交易执行的循环关系")
    add_paragraph(
        doc,
        "图1把上述定义展开为研究问题、数据、特征与模型、组合决策、执行监控和复盘。实际成交、成本和风险会通过复盘返回研究端，形成下一轮检验。各环节之间存在误差传导：数据时间戳错位可能制造虚假特征，有效信号也可能被过高换手和滑点抵消。这个循环解释了为什么策略评价必须同时检查数据、模型和执行。",
    )

    add_heading(doc, "1.2　数据、指标、信号与交易结果", 2)
    add_paragraph(
        doc,
        "把量化交易理解为系统以后，首先要区分原始数据与决策时可用的信息。任务一从五只A股和五只港股开始，统一日期、字段和前复权口径，并检查缺失、重复和异常价格关系。前复权会按照后续除权除息比例调整历史价格，使公司行动前后的价格序列保持可比。它仍然无法处理停牌、涨跌停和公告披露时点，也无法自动消除幸存者偏差。幸存者偏差是指样本只保留当前仍然存在的标的，从而遗漏已经退市或停止交易的对象。数据版本还需要记录下载日期、字段变化和最早可用时间，具体做法见附录A-01。",
    )
    add_paragraph(
        doc,
        "数据口径确定后，任务二把历史价量数据转换为技术指标。技术指标是对既有价格和成交量的统计变换，它不能创造新的事实信息。移动平均线描述平滑后的价格方向，平均真实波幅衡量包含跳空在内的近期波动，相对强弱指标比较一段时间内上涨和下跌的速度。多个指标即使名称不同，也可能重复使用同一组价格。指标选择因此需要先说明度量对象，再检验该指标是否增加了独立信息。",
    )
    add_table(
        doc,
        1,
        "量化研究各环节及其需要回答的问题",
        ["环节", "主要内容", "需要回答的问题"],
        [
            ["数据", "行情、财务、文本及其时间戳", "决策时是否真实可得，口径是否一致"],
            ["特征", "趋势、动量、波动和基本面度量", "是否重复、稳定并与目标有关"],
            ["信号", "规则条件、预测概率或横截面排序", "是否经过未参与选择的数据检验"],
            ["组合", "资产选择、仓位、风险预算和现金比例", "弱信号如何转化为可承受的暴露"],
            ["执行", "订单、成交、滑点、费用和异常", "理论收益在交易后保留了多少"],
            ["复盘", "归因、漂移、版本和停用条件", "结果变化来自市场还是系统实现"],
        ],
        widths=[2.0, 5.2, 7.1],
        font_size=8.8,
    )
    add_table(
        doc,
        2,
        "技术指标的信息类别和使用限制",
        ["信息类别", "代表指标", "在交易系统中的作用", "使用限制"],
        [
            ["趋势方向", "移动平均线、指数平滑异同移动平均线", "识别平滑后的价格方向", "趋势确认滞后，震荡期容易反复"],
            ["动量速度", "相对强弱指标、随机指标", "观察涨跌速度和短期极端位置", "固定阈值可能随市场状态失效"],
            ["波动程度", "平均真实波幅、布林带", "设置仓位和止损的风险尺度", "跳空时理论止损价格未必可成交"],
            ["趋势强度与偏离", "平均趋向指数、商品通道指数", "辅助区分趋势和震荡状态", "状态阈值仍需样本外检验"],
        ],
        widths=[2.6, 4.5, 4.4, 3.8],
        font_size=8.5,
    )

    add_heading(doc, "1.3　量化交易的核心价值", 2)
    add_paragraph(
        doc,
        "数据、指标和信号的关系明确后，可以进一步判断量化方法为交易提供了什么价值。第一项价值是规则透明。入场、退出、仓位和停用条件需要提前写清，策略结果才能追溯到具体规则。第二项价值是证据可复核。策略需要与同一时期的基准比较，并同时报告收益、回撤、换手和成本。基准是投资者在相同期间可以采用的替代方案，本次规则策略比较使用同期买入持有。量化方法仍然包含研究者的判断，但数据和规则会为判断留下可检查的依据。",
    )
    add_paragraph(
        doc,
        "可复核的比较还需要事前风险控制。任务三使用固定仓位，任务四则使用波动率风险预算。风险预算是先规定单笔交易可承受的资金损失比例，再根据止损距离反推仓位。这个差异说明，规则明确以后仍需统一风险尺度。持续复盘构成另一项价值。保存数据、参数、订单和成交记录以后，研究者才能区分模型失效、市场变化和实现错误。量化方法提高了检验与执行的一致性，但盈利仍取决于信号质量、市场条件和交易成本。",
    )

    add_heading(doc, "1.4　统一评价口径", 2)
    add_paragraph(
        doc,
        "上述价值需要通过统一指标检验。累计收益率等于期末净值除以初始净值再减1，它衡量整个区间的财富变化。最大回撤是净值相对此前历史高点的最大跌幅，它反映投资者在持有过程中经历的最严重亏损。夏普比率以年化超额收益除以年化波动率，用来比较每承担一单位波动获得的收益；各任务保留原作业设定的无风险利率口径。",
    )
    add_paragraph(
        doc,
        "机器学习结果还需要使用排序指标。曲线下面积表示随机抽取一个正类样本和一个负类样本时，模型把正类排在前面的概率，0.500对应随机排序。秩相关系数比较预测排名与实际收益排名，取值范围为-1至1，正值表示两组排名方向一致。滑点是理论成交价与实际可成交价格之间的差额。只有把这些指标与样本、基准和成本放在一起，分数才具有明确的交易含义。",
    )
    add_paragraph(
        doc,
        "指标口径统一后，还要区分证据来自哪一种检验。回测是在历史数据上按照预定规则模拟交易。样本外检验要求先确定模型和参数，再观察未参与选择的数据。平台模拟盘按照真实交易日顺序运行虚拟资金，可以记录订单、成交和持仓，但它仍不包含真实资金交易的全部压力。截至2026年7月25日，任务七的三个模拟盘已经部署，尚未形成可分析的订单、成交和净值记录。第二章因此先使用已有回测和短窗回放比较规则策略，再讨论执行证据的空缺。",
    )


def final_chapter2(doc: Document) -> None:
    add_heading(doc, "二、量化交易策略综合分析", 1, page_break=True)
    add_heading(doc, "2.1　规则策略的共同比较口径", 2)
    add_paragraph(
        doc,
        "第一章的评价口径只有在样本和成本一致时才能支持策略比较。双均线和海龟策略都在贵州茅台上接受检验，区间为2025年5月29日至2026年7月3日，初始资金100万元，买卖手续费各0.03%，买卖滑点各0.01%。两套策略都只做多，即仓位可以为正或为零，但不建立空头仓位。同期买入持有作为基准。统一标的、区间和费用后，结果差异主要来自信号和风险控制。这个样本只有约十三个月，而且整体下跌，因此只能用于解释当前条件下的机制，长期和跨资产表现仍需另行检验。",
    )
    add_table(
        doc,
        3,
        "双均线、海龟策略与买入持有的主要结果",
        ["指标", "双均线策略", "海龟策略", "买入持有"],
        [
            ["累计收益率", "-5.67%", "-0.74%", "-18.24%"],
            ["年化收益率", "-3.61%", "-0.47%", "-11.63%"],
            ["最大回撤", "-15.20%", "-9.72%", "-21.35%"],
            ["夏普比率", "-0.58", "-0.29", "-1.12"],
            ["交易胜率", "35.71%", "40.00%", "不适用"],
            ["完整交易次数", "14", "5", "不适用"],
        ],
        widths=[4.1, 3.7, 3.7, 3.7],
        font_size=9.0,
    )

    add_heading(doc, "2.2　双均线与海龟策略的结果和原因", 2)
    add_paragraph(
        doc,
        "比较条件固定后，可以先分析两套策略如何生成信号。图2把累计收益率和最大回撤放在同一页，纵轴都以百分比表示，零线用于区分盈利和亏损。双均线策略在五日均线上穿十五日均线时买入，在五日均线下穿十五日均线时退出。海龟策略使用唐奇安通道判断突破，通道上轨和下轨分别由过去一段时间的最高价和最低价构成；策略再用平均真实波幅设置止损距离，并按照风险预算计算仓位。",
    )
    add_figure(doc, "图02_规则策略收益与回撤.png", "图2　规则策略与买入持有的收益和回撤比较")
    add_paragraph(
        doc,
        "图2显示，双均线的累计收益比买入持有高12.57个百分点，最大回撤低6.15个百分点。海龟策略的累计收益比买入持有高17.50个百分点，最大回撤低11.63个百分点；它的累计损失和回撤也低于双均线。较少的交易次数、波动率止损和仓位预算可以解释这一差异，但单一样本无法确认这些机制在其他时期仍会得到相同结果。两项策略的绝对收益和夏普比率均为负，因此当前证据只支持相对减损结论。",
    )
    add_paragraph(
        doc,
        "进一步比较两套机制，双均线的规则较短，参数也较少，每次交易都可以直接追溯。均线需要等待价格趋势在平均值中体现，因此信号存在确认滞后。固定仓位还会使同一次交叉在不同波动环境下承担不同风险。海龟策略通过风险预算控制单次判断错误的预期损失，但两倍平均真实波幅只在当前候选参数和样本中相对平衡。跳空或流动性不足时，实际成交价可能越过止损价，真实损失也会超过预算。",
    )
    add_paragraph(
        doc,
        "机制比较之后还需要检查参数稳定性。图3扫描不同的短期和长期均线组合，横轴是长期均线周期，纵轴是短期均线周期，格内数字为累计收益率。若较高收益来自稳定关系，相邻参数通常会形成连续的高值区域；孤立最大值更可能依赖某一段特殊行情。",
    )
    add_figure(doc, "图03_双均线参数敏感性.png", "图3　双均线参数组合的累计收益率敏感性")
    add_paragraph(
        doc,
        "扫描结果显示，多数参数组合仍为负收益。正收益集中在短周期3日、长周期10日附近的孤立格点，相邻组合没有形成连续的正收益区域。这个分布缺少参数稳定性证据，因此3日和10日只能视为当前样本中的局部高值。后续检验需要扩大标的和时间范围，并采用滚动样本外窗口观察整个参数邻域，见附录A-02。",
    )

    add_heading(doc, "2.3　市场环境、参数与执行风险", 2)
    add_table(
        doc,
        4,
        "规则策略的适用环境和主要失效方式",
        ["市场环境", "双均线策略", "海龟策略", "需要同时检查的风险"],
        [
            ["持续趋势", "跟随方向，但确认和退出滞后", "突破后持有，并按波动调整风险", "趋势反转后的利润回吐"],
            ["区间震荡", "均线反复交叉", "假突破后连续止损", "小额亏损和交易成本累积"],
            ["单边下跌且只做多", "空仓可以减少部分损失", "风险预算可以压低暴露", "难以获得正绝对收益"],
            ["跳空或流动性收缩", "理论成交价偏乐观", "止损价格可能无法实现", "回测与实际执行偏差"],
        ],
        widths=[3.0, 4.2, 4.2, 3.9],
        font_size=8.5,
    )
    add_paragraph(
        doc,
        "参数稳定性仍不足以说明策略适用于所有市场。表4把结果放回持续趋势、区间震荡、单边下跌和流动性收缩四种环境。双均线和海龟策略都依赖趋势延续，因此两者可能具有较高的收益相关性。收益相关性衡量两套策略的收益是否同向变化；相关性较高时，组合持有也难以获得充分分散。风险分析还需要检查回撤持续时间、回撤形成环境、邻近参数结果以及成本和流动性冲击，后续检验方案见附录A-03和A-05。",
    )
    add_paragraph(
        doc,
        "对市场环境和失败方式的分析改变了前后两项任务的关系。任务三首次把信号、成本、基准和绩效连成完整回测。任务四进一步把注意力从入场判断转向损失控制。两个策略的绝对收益均为负，但回测仍识别出各自的适用条件、风险来源和后续检验需求。",
    )

    add_heading(doc, "2.4　从单标的回测到宽基指数基金策略", 2)
    add_paragraph(
        doc,
        "单标的回测揭示了信号和风险控制的机制，任务七进一步把研究对象扩展到三套宽基指数基金策略。宽基指数基金跟踪覆盖较多股票的市场指数，可以降低单一公司的特有风险。策略A使用高仓位均线规则作为简单基线。策略B在趋势信号上加入波动率风险预算、止损和50%仓位上限。策略C比较沪深300、中证500和创业板三只基金过去十五日的收益率，每周选择动量最高的基金，最高持有50%；当三只基金的动量都不为正时，资金保持现金。三套策略使用50万元初始资金和相同平台成本口径，长期回测覆盖2019年1月1日至2026年7月17日。",
    )
    add_table(
        doc,
        5,
        "三套宽基指数基金策略的平台长期回测结果",
        ["指标", "策略A", "策略B", "策略C"],
        [
            ["累计收益率", "-0.66%", "2.91%", "71.84%"],
            ["年化收益率", "-0.09%", "0.39%", "7.69%"],
            ["超额收益率", "-33.97%", "-31.59%", "14.23%"],
            ["最大回撤", "-36.97%", "-12.31%", "-16.91%"],
            ["夏普比率", "-0.29", "-0.77", "0.32"],
            ["市场敏感度", "0.521", "0.164", "0.426"],
            ["年化波动率", "14.10%", "4.70%", "11.70%"],
        ],
        widths=[4.5, 3.6, 3.6, 3.6],
        font_size=8.9,
    )
    add_paragraph(
        doc,
        "平台结果显示，策略C的累计收益率和超额收益率最高，最大回撤低于策略A。超额收益率是策略收益相对于平台基准的差额。策略B把最大回撤降至12.31%，累计收益为2.91%，市场敏感度为0.164。市场敏感度衡量策略收益随基准变化的程度，较低数值通常对应较低的市场暴露。因此，策略B的防守表现有相当部分来自较低仓位。策略A的长期收益接近零，却承受36.97%的最大回撤，收益不足以补偿持有期间的路径风险。",
    )
    add_paragraph(
        doc,
        "统一平台指标说明了期末结果，净值路径则用于判断优势在何时形成。图4使用2016年1月4日至2026年7月24日的本地扩展样本。归一化净值是把各序列的起始值统一设为1，因此不同策略可以在同一尺度上比较。上方面板显示归一化净值，下方面板显示各策略相对于自身历史高点的回撤。本地复算与平台结果使用不同来源和实现，数值差异需要单独解释。",
    )
    add_figure(doc, "图08_TASK7本地扩展净值与回撤.png", "图4　三套宽基指数基金策略的本地扩展样本净值和回撤")
    add_paragraph(
        doc,
        "沿时间路径观察，策略A在2021年后进入持续较深的回撤，期末累计收益为-30.38%，最大回撤为-42.91%。策略B长期围绕初始净值波动，最大回撤为-11.82%，平均仓位只有18.95%。平均仓位是各交易日已投资资金占总资产比例的均值，因此策略B的较浅回撤需要与较低资金使用率一起理解。策略C的领先主要在2025年后形成，期末累计收益为56.93%，最大回撤为-16.72%，平均仓位为34.07%。收益集中在较短阶段会增加结论对市场状态的依赖。",
    )
    add_paragraph(
        doc,
        "平台与本地结果的差异可能来自数据源、复权、成交价格、订单撮合和成本实现。同一个策略名称只有在数据版本、参数和成交规则一致时才应得到可比结果。研究系统需要保存这些信息，并对差异进行归因。选择单一较高结果会掩盖实现风险。",
    )

    add_heading(doc, "2.5　成本压力、短窗回放与模拟盘状态", 2)
    add_paragraph(
        doc,
        "长期回测给出历史路径，执行可行性还取决于交易成本。成本压力测试是在其他规则保持不变时提高费用假设，用来衡量结果对交易摩擦的敏感程度。图5在同一纵轴上比较基础成本和佣金、滑点同时翻倍的结果，零线表示收益由正转负。",
    )
    add_figure(doc, "图09_TASK7成本压力测试.png", "图5　三套宽基指数基金策略的成本翻倍压力测试")
    add_paragraph(
        doc,
        "压力测试显示，成本翻倍后，策略A由-30.38%降至-53.61%，策略B由2.16%降至-9.35%，策略C由56.93%降至25.49%。策略C仍保持正收益，但减少31.44个百分点，说明轮动换手会明显侵蚀优势。策略B的历史微利被成本完全吞噬。当前测试只衡量成本敏感性，策略的可交易性还取决于最低佣金、涨跌停、停牌、成交量和冲击成本，见附录A-03。",
    )
    add_table(
        doc,
        6,
        "任务七不同结果的阅读方式",
        ["结果来源", "时间范围", "适合回答的问题", "当前结果"],
        [
            ["平台长期回测", "2019-01-01至2026-07-17", "统一平台下的历史收益和风险", "三策略历史指标已经形成"],
            ["本地扩展复算", "2016-01-04至2026-07-24", "净值路径、仓位和成本敏感性", "用于解释平台结果，不替代平台口径"],
            ["冻结规则五日回放", "2026-07-20至2026-07-24", "当时信息下的信号和目标仓位", "三策略均空仓，收益0.00%"],
            ["平台模拟盘", "2026-07-24部署", "订单、成交、持仓和执行偏差", "截至7月25日尚无运行数据"],
        ],
        widths=[3.2, 3.8, 5.0, 3.3],
        font_size=8.4,
    )
    add_paragraph(
        doc,
        "成本压力之外，五日回放用于检验冻结规则在新窗口中的行为。冻结规则是指进入观察窗口前已经确定数据、参数和交易条件，观察期间不再根据结果修改。三套策略都没有交易，沪深300同期上涨2.65%，平台口径超额收益约为-2.58%。0.00%回撤来自全程空仓，它说明当周没有市场暴露，也记录了趋势规则错过上涨的机会成本。由于没有订单和成交，这五天无法检验滑点、拒单和止损。平台模拟盘需要从新的交易日开始独立积累记录，历史回放不计入模拟业绩，后续记录方案见附录A-08。",
    )

    add_heading(doc, "2.6　多策略系统的组合方法", 2)
    add_paragraph(
        doc,
        "明确单策略的市场条件、成本和执行边界以后，才适合讨论多策略组合。趋势暴露是策略收益对持续上涨或下跌行情的共同依赖。双均线、海龟和宽基基金轮动都含有趋势暴露，因此市场反转时可能同时失效。组合前需要统一数据时点和成本，计算策略收益相关性、平均仓位、回撤同步性和换手，再设置单策略权重、总仓位和现金缓冲。",
    )
    add_table(
        doc,
        7,
        "多策略交易系统的层级和控制内容",
        ["层级", "主要输入", "主要输出", "需要控制的风险"],
        [
            ["预测层", "规则信号、概率或横截面排序", "方向、收益估计和资产排名", "泄漏、过拟合和关系漂移"],
            ["决策层", "预测分数、波动率、相关性和成本", "标的选择、仓位和风险预算", "集中度、共同暴露和过度换手"],
            ["执行层", "目标仓位、流动性和订单规则", "成交、实际成本和持仓", "延迟、拒单、滑点和跳空"],
            ["复盘层", "预期与实际结果、异常和版本", "归因、降权、停用或再训练", "错误重复和失效策略继续运行"],
        ],
        widths=[2.4, 4.7, 4.3, 3.9],
        font_size=8.4,
    )
    add_paragraph(
        doc,
        "表7把多策略系统拆成预测、决策、执行和复盘四层。预测层产生方向、概率或资产排序。决策层把预测与风险、相关性和成本结合，形成目标仓位。执行层把目标仓位转化为订单，并记录实际成交。复盘层比较预期与实现，根据差异决定归因、降权或停用。最大回撤、市场状态、参数敏感性和压力情景需要在组合层统一控制，见附录A-06和A-07。",
    )
    add_paragraph(
        doc,
        "规则策略提供了透明基线，但固定规则难以表达复杂的非线性关系和多资产排序。机器学习可以扩展信号形式，同时增加时间泄漏、过拟合和关系漂移的风险。第三章继续沿用相同的样本外和成本口径，分析模型如何从数据进入交易决策。",
    )


def final_chapter3(doc: Document) -> None:
    add_heading(doc, "三、机器学习在量化交易中的应用总结", 1, page_break=True)
    add_heading(doc, "3.1　低信噪比改变了机器学习问题", 2)
    add_paragraph(
        doc,
        "机器学习模型需要先定义标签。标签是模型要预测的结果，金融任务可以把它设为未来涨跌、未来收益或横截面排名。金融收益通常具有低信噪比，即稳定且可预测的部分相对于随机波动较小。真实关系还会随着波动、流动性、监管和参与者行为变化。重叠标签会重复计算相近的未来价格区间，时间泄漏则会把决策时尚未出现的信息带入训练。研究因此需要先说明标签能否转化为交易，并确认每项输入在决策时已经可得，算法选择应在这些条件明确以后进行。",
    )
    add_paragraph(
        doc,
        "这一判断也在另一门课程中得到呼应。在学校修读《量化金融专题》期间，课程邀请Jump Trading进行客座讲座。听完后，我发现他们对低信噪比、预测期限和有效样本的经验，与前述任务中的实际体验十分接近。有效样本量是考虑样本重叠和相关性后，真正提供独立信息的观察数量。长期预测的独立时间样本较少，复杂模型容易缺少训练证据；短期数据行数更多，却对推断速度、交易摩擦和订单撮合等市场微观结构更敏感。因此，模型复杂度需要与有效样本量和执行条件共同设计。",
    )

    add_heading(doc, "3.2　数据预处理、特征工程与时间验证", 2)
    add_paragraph(
        doc,
        "低信噪比使研究流程本身成为结果可信度的一部分。数据预处理负责核对标签、缺失值、异常值和时间顺序。特征工程把原始数据转换为模型输入，并检查变量的经济含义、相关性和跨期稳定性。时间验证按照先后顺序划分训练、验证和测试窗口，使模型只能使用当时已经可得的信息。表8按照这三个环节继续展开模型选择、样本外评价、交易映射和部署监控。",
    )
    add_table(
        doc,
        8,
        "机器学习研究流程及容易出现的偏差",
        ["环节", "当前采用的方法", "主要检查", "常见偏差"],
        [
            ["数据预处理", "核对标签、缺失、异常和时间顺序", "样本覆盖、类别比例、时间戳", "使用未来信息或修订后数据"],
            ["特征工程", "只在训练集内筛选、去相关和检验稳定性", "共线性、跨期分布和可解释性", "用测试集反复挑变量"],
            ["模型选择", "训练期内部交叉验证或扩展窗口", "简单基线、参数邻域和验证波动", "追逐单次最高分"],
            ["样本外评价", "模型冻结后观察独立测试窗口", "区间、分期结果和随机基准", "看过测试结果后继续改模型"],
            ["交易映射", "将概率或排序转成阈值、仓位和组合", "净收益、回撤、换手和成本", "用预测分数代替策略收益"],
            ["部署监控", "记录输入、信号、成交和版本", "漂移、成交偏差和停用条件", "关系失效后仍按原模型运行"],
        ],
        widths=[2.2, 5.0, 4.0, 4.1],
        font_size=8.2,
    )
    add_paragraph(
        doc,
        "任务五的乳腺癌教学案例把569例相互独立的患者样本分成455例训练和114例测试，特征筛选只读取训练集。最终保留六项变量，再用同一测试集比较四类模型。逻辑回归用特征的线性加权结果估计分类概率；决策树按照特征阈值逐层划分样本；随机森林汇总多棵树的结果；梯度提升逐轮拟合前一轮的预测误差。患者记录没有时间先后依赖，因此可以随机划分。金融样本具有明确时序，随机打散会把未来市场状态带入训练，所以金融任务必须按照时间顺序划分。",
    )
    add_paragraph(
        doc,
        "完成时间划分以后，特征工程需要控制信息重复。价量指标大多来自相同的价格和成交量序列，二十个高度相关的指标可能只重复表达少数信息。财务和文本数据还需要记录实际披露时间，因为报告期结束日通常早于投资者能够读取信息的日期。后续研究需要把时点可得性和历史成分保存为数据规则，见附录A-01和A-04。",
    )

    add_heading(doc, "3.3　模型选择与评价：任务五", 2)
    add_paragraph(
        doc,
        "研究流程确定后，任务五先用结构较清晰的教学数据检验分类管线。乳腺癌案例把恶性病例编码为正类，逻辑回归和随机森林的测试曲线下面积均约为0.994，梯度提升约为0.992，决策树约为0.972。简单逻辑回归已经接近复杂模型，说明当前标签与细胞核形态特征之间具有较强的可分关系。这些数值只评价该教学样本中的分类流程，金融标签仍需单独检验。",
    )
    add_paragraph(
        doc,
        "金融案例把未来60个交易日收益位于同期前30%的股票记为正类，位于后30%的股票记为负类。图6展示2025年测试样本的曲线下面积，竖线0.500表示随机排序。点两侧的区间来自时间块重采样，即重复抽取连续时间块来保留相邻样本的相关性。若区间覆盖0.500，当前样本仍无法区分模型优势和随机波动。",
    )
    add_figure(doc, "图04_TASK5样本外AUC区间.png", "图6　金融分类模型的测试曲线下面积及不确定区间")
    add_paragraph(
        doc,
        "逻辑回归、决策树和随机森林的测试值分别为0.509、0.477和0.533，三组区间都覆盖0.500。随机森林的点估计最高，但区间宽度说明有效时间状态仍然有限。随机标签检验使分数回到随机水平，人工构造的可预测标签则得到高分，这两个对照没有发现评价管线失效。接近随机的真实标签结果因此指向目标定义、信息覆盖和市场变化，继续扩大参数搜索缺少独立证据。",
    )
    add_paragraph(
        doc,
        "单个测试窗口只能描述一个时期，跨期比较才能检查关系是否迁移。图7把同一研究中的验证、开发和测试时期分开。每个点是一个时期的聚合结果，三个点提供的时间状态仍然有限。",
    )
    add_figure(doc, "图05_TASK5跨期AUC.png", "图7　金融分类模型在不同时期的判别能力")
    add_paragraph(
        doc,
        "跨期结果显示，模型在2023年验证期的表现较高，进入2024年和2025年后普遍下降并接近随机水平。三个时期不足以估计衰减速度，但已经表明训练关系没有稳定迁移。分布漂移是特征分布或特征与标签关系随时间发生变化，它与验证窗口偶然性、信号衰减都可能造成这种结果。模型需要在滚动时间窗口和不同市场状态下继续评价，一次验证高分无法提供长期有效性证据。",
    )

    add_heading(doc, "3.4　从预测到交易：任务六", 2)
    add_paragraph(
        doc,
        "任务五主要评价分类排序，任务六进一步把预测排名连接到投资组合。主任务使用股票与季度构成的面板数据，即每一季度同时包含多只股票的特征和后续收益。前七个季度用于训练，后三个季度用于测试。模型以季度秩相关系数评价预测排名，再选择排名前30只股票等权持有。等权持有表示每只入选股票获得相同比例的初始权重。线性回归按训练期内部规则入选，验证期平均秩相关系数为0.129，测试期为0.273。扣除单边20个基点成本后，三个测试季度累计收益为41.46%，同期全市场等权为-15.91%。",
    )
    add_paragraph(
        doc,
        "为了把排序分数连接到实际收益，图8以季度为横轴，柱形比较线性回归前30名组合和全市场等权收益，零线区分正负。季度结果应逐期阅读，不应把三个柱形延伸成长期趋势。",
    )
    add_figure(doc, "图06_TASK6季度组合收益.png", "图8　线性回归前30名组合与全市场等权的测试季度收益")
    add_paragraph(
        doc,
        "三个测试季度中，前30名组合收益均为正，全市场等权有两个季度为负，这与测试期存在横截面排序信息相符。时间证据仍然只有三个季度，增加同一季度的股票行数不会增加独立市场状态。平均单边换手率为83.33%，这里的单边换手率等于相邻两期组合权重变化绝对值之和的一半。较高换手会放大费用和换仓时点的影响。随机森林的测试累计收益更高，但预先规定的模型选择规则没有选中它，因此测试结果不能用于事后更换模型。",
    )
    add_paragraph(
        doc,
        "横截面组合检验同一时期内的股票排序，方向预测附加实验则判断未来三日涨跌。图9连接每一轮实验的验证分数和测试分数，线段越陡，说明两个时期的迁移差距越大。",
    )
    add_figure(doc, "图07_TASK6验证与测试AUC.png", "图9　三轮方向模型的验证与测试表现")
    add_paragraph(
        doc,
        "三轮结果显示，第一轮随机森林从验证0.756降至测试0.335，第二轮逻辑回归从0.692降至0.520。第三轮使用180日滚动逻辑回归，每个交易日只用最近180日数据重新训练，验证值为0.631，测试值为0.571。滚动训练缩小了迁移差距，但后续方案受到前一轮结果影响，因此0.571属于探索性发现。机器学习择时收益只有0.50%，低于同期买入持有的9.93%和均线策略的9.68%。曲线下面积只衡量排序，策略收益还取决于涨跌幅度、交易阈值、仓位、持有期和成本。",
    )
    add_table(
        doc,
        9,
        "机器学习结果应如何解释",
        ["研究内容", "主要结果", "合理解释", "仍需补充的检验"],
        [
            ["医疗教学分类", "测试曲线下面积0.972至0.994", "流程能识别结构较清晰的标签", "不外推金融收益"],
            ["金融分类", "测试值0.477至0.533，区间覆盖0.500", "尚未发现稳定判别优势", "扩大时间状态并滚动检验"],
            ["季度横截面排序", "测试秩相关0.273，三季度组合收益41.46%", "当前三季度存在排序效果", "历史成分、披露时点、成本和更长样本"],
            ["三日方向模型", "滚动模型测试0.571，策略收益0.50%", "弱排序没有转成更高收益", "独立新窗口和交易映射优化"],
        ],
        widths=[3.1, 4.3, 4.3, 3.6],
        font_size=8.2,
    )

    add_heading(doc, "3.5　机器学习的优势与局限", 2)
    add_paragraph(
        doc,
        "任务五和任务六共同界定了机器学习的适用范围。机器学习可以处理多特征、非线性关系和横截面排序，并输出连续概率作为动态仓位的输入。统一训练管线也便于重复运行和部署监控。金融标签噪声、有效样本不足、关系漂移和交易成本会限制这些优势。曲线下面积评价分类排序，秩相关评价横截面排名，策略收益还包含仓位和成本，三者回答不同问题。特征重要性只说明模型依赖，不能证明因果关系。",
    )
    add_paragraph(
        doc,
        "这些局限在任务六中表现得最直接，也构成整个学习过程中最强的个人触动。作为初学者，调参时很难判断分数变化来自真实结构、样本偶然，还是模型对验证窗口的适应。验证高分没有稳定迁移，测试分数改善也没有带来更高收益，这使我直观感受到量化金融的困难。模型仍然需要市场知识、交易经验和判断来定义目标、选择信息并解释失效，单纯调参无法完成这些工作。",
    )

    add_heading(doc, "3.6　深度学习与非结构化数据的研究方向", 2)
    add_paragraph(
        doc,
        "前述局限也确定了下一阶段的研究方向。深度学习使用多层神经网络学习复杂关系；多模态模型联合处理结构化价量数据、文本和图像；非结构化数据包括公告、新闻、电话会议记录和产业链资料。它们可能补充纯价量序列缺少的公司与行业信息，也可以针对特定金融任务进行专门化。研究目的在于扩展信息来源，模型规模本身不能保证发现盈利信号。",
    )
    add_paragraph(
        doc,
        "每个新实验必须满足三个条件。信息在决策时已经可得；新模型相对于简单基准产生稳定的样本外增量；增量在扣除成本、延迟和风险约束后仍有交易价值。直觉上仍倾向于相信复杂现象可能存在可解释结构，但信号的存在与可发现性都需要独立证据。因此，后续研究应先补充信息和验证，再增加模型复杂度。",
    )


def final_chapter4(doc: Document) -> None:
    add_heading(doc, "四、结论与展望", 1, page_break=True)
    add_heading(doc, "4.1　主要收获与认识变化", 2)
    add_paragraph(
        doc,
        "前七项任务覆盖了数据获取、指标计算、规则回测、机器学习训练、组合评价和平台部署，评价方式也在实践中发生变化。早期分析主要关注收益率，现在会同时检查样本、基准、回撤、换手和成本。早期参数选择集中于单个最高值，现在更关心相邻参数和新时期能否得到相近结果。任务七又把订单、成交和持续监控纳入研究，使回测成为执行验证的起点。",
    )
    add_paragraph(
        doc,
        "这些变化来自前述实证结果。规则策略在下跌样本中减少了相对损失，但没有取得正绝对收益。机器学习在教学数据上得到高分，进入金融测试以后接近随机水平，较弱的排序优势也没有稳定转化为收益。任务七进一步说明，平台回测、本地复算和模拟运行分别回答历史表现、实现差异和执行过程问题。现阶段最可靠的成果是一套能够说明结果范围、证据限制和后续检验的方法，未来收益仍需新的独立数据验证。",
    )

    add_heading(doc, "4.2　任务六失败带来的方法转变", 2)
    add_paragraph(
        doc,
        "在这些变化中，任务六的失败对研究方法影响最直接。数学题通常有明确条件和可计算理论，金融市场则可能根本不存在稳定、可发现且扣除成本后仍能盈利的信号。那场客座讲座提出的“真实信号还是结构化噪声”，与任务六的失败形成了直接呼应。信号存在本身因此成为需要检验的假设。若独立测试没有发现稳定增量信息，这个结果也构成一项完整结论。",
    )
    add_paragraph(
        doc,
        "这一认识仍然允许使用复杂模型，但研究顺序需要调整。模型训练前应说明金融机制和简单基准，模型选择期间应保护独立测试窗口，参数搜索也需要预设停止条件。若复杂模型只提高验证分数，成本后的收益没有稳定改善，研究就应保留负结果并结束当前搜索。",
    )

    add_heading(doc, "4.3　机构与个人的评价目标", 2)
    add_paragraph(
        doc,
        "方法转变还需要与使用者的目标相匹配。在机构从业角色下，策略会在治理和风险预算内更积极地追求收益，同时重视可解释性，因为模型需要经过审查、沟通并进入执行系统。机构还要考虑策略容量、客户约束和团队分工。在个人投资角色下，评价更重视稳定性、回撤和长期可执行性。这里的稳定性是指收益路径和风险水平与资金用途、流动性需求及心理承受能力相匹配。",
    )
    add_paragraph(
        doc,
        "两种角色的目标不同，但期末收益都不足以单独完成评价。策略B的低回撤部分来自低仓位，策略C的高收益又受到成本和阶段集中的影响。只有先写明使用者目标，收益、回撤、资金使用率、成本和可解释性才能获得合适的权重。",
    )

    add_heading(doc, "4.4　下一阶段计划", 2)
    add_table(
        doc,
        10,
        "当前研究限制和下一步改进",
        ["当前限制", "优先行动", "附录编号"],
        [
            ["数据时点和历史版本记录不足", "保存披露日、历史成分、复权和字段版本", "A-01"],
            ["规则策略集中在单标的和短样本", "扩展资产与市场状态，采用滚动样本外窗口", "A-02、A-05"],
            ["成本和成交假设仍较简化", "加入涨跌停、停牌、冲击成本和成交失败", "A-03"],
            ["模型选择受到有限验证窗口影响", "采用嵌套时间验证并保护新测试窗口", "A-04"],
            ["多策略共同暴露尚未量化", "测量相关性和回撤贡献，设置组合风险预算", "A-06"],
            ["模拟盘尚无连续运行数据", "保存订单、成交、持仓、异常和版本记录", "A-07、A-08"],
        ],
        widths=[5.0, 7.8, 2.5],
        font_size=8.5,
    )
    add_paragraph(
        doc,
        "角色目标明确以后，可以把方法要求转化为下一阶段计划。根据表10，短期工作应先积累任务七的模拟盘数据，核对冻结规则是否按预期生成订单，并记录实际成本和异常。规则策略将扩大到更多标的和市场状态。机器学习部分采用滚动或嵌套时间验证；嵌套时间验证在内层窗口选择特征和参数，在外层未触碰窗口评价最终模型。数据和验证基础稳定后，再研究多策略风险配置和小资金验证。",
    )
    add_paragraph(
        doc,
        "在短期证据稳定之后，中期研究会把深度学习和非结构化数据纳入同一套评价纪律。新增文本或多模态信息必须说明时间戳，并与简单模型比较。模型上线后还需要监控输入分布、预测分布、换手、成交和回撤。当实时输入分布明显偏离训练期时，系统需要触发复核、降权或停用。这样的监控可以减少把关系漂移和结构化噪声误写成稳定结论的概率。",
    )
    add_paragraph(
        doc,
        "回到最初的认识，学习前我把量化交易看成算法；现在则把它理解为持续面对信息不对称、市场变化和执行摩擦的系统。下一步的专业成长包括继续学习模型，也包括理解市场、接受不确定性，并判断一项结论需要哪些数据和检验才能成立。",
    )


def appendix(doc: Document) -> None:
    add_heading(doc, "附录　基于前期作业复盘的改进建议", 1, page_break=True)
    add_paragraph(
        doc,
        "改进顺序以证据可靠性为先。数据时点、样本外验证、交易成本和执行记录达到基本要求以后，再增加模型复杂度。",
        indent=False,
    )
    items = [
        (
            "A-01　建立时点可得的数据治理",
            "为每个数据集记录来源、下载时间、复权方式、字段变更和缺失处理；为公告、财务和非结构化特征记录最早可用时间；保留原始快照与处理后版本的对应关系。验收标准是任一结果均可由冻结版本复算，且不存在使用未来可得信息的字段。",
        ),
        (
            "A-02　扩展标的、周期和滚动样本外检验",
            "规则策略应覆盖不同资产、牛熊和震荡状态，采用滚动训练、验证和测试窗口，报告各窗口收益、回撤、换手和失效比例。参数选择以稳定区域和样本外分布为主，不以单一最高值为目标。",
        ),
        (
            "A-03　加入可交易性和成本压力测试",
            "在手续费和固定滑点之外，加入成交量约束、涨跌停、停牌、冲击成本和止损跳空情景。至少设置基础、双倍和极端成本三档，并报告收益下降的百分点、盈亏符号是否变化以及容量上限。",
        ),
        (
            "A-04　使用更严格的机器学习验证",
            "采用按时间嵌套的模型选择：内层只用于特征和参数选择，外层保留为未触碰测试；同时报告简单基线、区间估计、分期结果和校准。看过测试结果后提出的新方案必须进入下一独立窗口，原测试期结论保持不变。",
        ),
        (
            "A-05　建立市场状态识别",
            "使用波动率、趋势强度、流动性和相关性定义可解释的市场状态，比较双均线、突破和轮动策略在不同状态下的条件表现。状态模型只用于降低不适配暴露，不能用事后标签解释所有失败。",
        ),
        (
            "A-06　构建组合层风险预算",
            "根据策略相关性、回撤同步性、平均仓位和成本分配风险。单独按照历史收益加权会放大样本依赖。组合还需要设置资产集中度、共同方向暴露、波动和回撤上限。机构与个人版本分别定义收益目标、风险承受和解释要求。",
        ),
        (
            "A-07　建立模型监控和停用规则",
            "部署后监控输入缺失、特征漂移、预测分布、换手、成交率、成本偏差和回撤。当关键指标连续越过阈值时触发人工复核、降仓或停用；任何版本变更都需保存原因和生效日期。",
        ),
        (
            "A-08　持续积累任务七模拟交易证据链",
            "按日保存目标仓位、订单、成交、持仓、净值、现金和拒单记录，按周比较回测预期与模拟实现。先完成至少三个月连续观察，并尽量覆盖不同波动状态，再评价执行偏差和策略可用性；模拟结果仍不能替代真实资金风险评估。",
        ),
    ]
    for heading, text in items:
        add_heading(doc, heading, 2)
        add_paragraph(doc, text)


def references(doc: Document) -> None:
    add_heading(doc, "参考资料", 1)
    refs = [
        "[1] 量化交易工作坊任务一至任务八课程任务书与课程讲义，2026。",
        "[2] 前序任务成果文件：任务一至任务七报告、结构化数据、图表、回测摘要与仪表板数据，2026。",
        "[3] 《量化金融专题》课程Jump Trading客座讲座材料：低信噪比、研究系统、市场变化与人工智能研究方法，课程提供资料，2026。",
        "[4] J. Welles Wilder Jr. New Concepts in Technical Trading Systems. Trend Research, 1978。",
        "[5] The Original Turtle Trading Rules，课程策略实现参考资料。",
        "[6] AKShare与Tushare Pro数据接口说明，任务一数据获取参考。",
        "[7] JoinQuant研究与模拟交易平台说明，任务七平台回测和部署参考。",
        "[8] scikit-learn模型评价与时间序列验证方法说明，任务五和任务六建模参考。",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        set_paragraph_format(p, indent=False, line=1.2)
        p.paragraph_format.left_indent = Cm(0.65)
        p.paragraph_format.first_line_indent = Cm(-0.65)
        r = p.add_run(ref)
        set_run_font(r, size=9.2)
    p = doc.add_paragraph()
    set_paragraph_format(p, indent=False, line=1.2)
    r = p.add_run(
        "说明：文中数值均来自工作区内已冻结的前序任务材料。任务七平台长期回测结束日统一采用正式报告与结构化摘要中的2026年7月17日；"
        "本地扩展复算结束于2026年7月24日。"
    )
    set_run_font(r, size=9.2)


def add_doc_properties(doc: Document) -> None:
    props = doc.core_properties
    props.title = "从数据到执行：量化交易策略与机器学习应用的综合实践报告"
    props.subject = "量化交易工作坊任务八专业学习成果报告"
    props.author = "李沐晓"
    props.keywords = "量化交易, 趋势跟踪, 机器学习, 样本外验证, 风险管理"
    props.comments = "基于任务一至任务七成果的综合总结"


def main() -> None:
    doc = Document()
    configure_styles(doc)
    configure_document(doc)
    add_doc_properties(doc)
    add_cover(doc)
    add_toc(doc)
    add_abstract(doc)
    final_chapter1(doc)
    final_chapter2(doc)
    final_chapter3(doc)
    final_chapter4(doc)
    appendix(doc)
    references(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
